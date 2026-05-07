"""Generate the Claude Code local-LLM setup docs (user + admin) in docx + md.

Style mirrors the in-app guide at web/templates/claude_code_setup.html:
brand blue gradient table headers, alternating rows, status pills, callout
boxes, KPI cards. Two output documents:

    docs/CLAUDE_CODE_USER_SETUP.{docx,md}
    docs/CLAUDE_CODE_ADMIN_SETUP.{docx,md}

Run:
    .venv/bin/python misc_scripts/build_claude_code_docs.py
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
DOCS = REPO / "docs"

sys.path.insert(0, str(REPO))
from my_config import get_config

LLM_PUBLIC_URL = get_config().local_llm_public_url or "https://<your-llm-host>/local-llm"

# --- Brand palette (matches web/templates/claude_code_setup.html) -----------
BLUE = "0046AD"
BLUE_LIGHT = "1D4ED8"
GREEN = "00A651"
AMBER = "F59E0B"
RED = "DC2626"
SLATE = "334155"
SLATE_DEEP = "0F172A"
SLATE_MUTED = "64748B"
ROW_ALT = "F2F7FF"      # light blue alternating row
CODE_BG = "F1F5F9"      # code block background
CALLOUT_TIP = "DCFCE7"
CALLOUT_NOTE = "DBEAFE"
CALLOUT_WARN = "FEF3C7"
CALLOUT_IMPORTANT = "FEE2E2"
KPI_BG = "EAF2FB"

MONO = "JetBrains Mono"
SANS = "Calibri"

# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------

def _shade(cell, hex_fill: str) -> None:
    """Set cell background fill colour."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def _set_cell_borders(cell, color: str = "C7D2DE", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _left_accent(cell, color: str, size: str = "32") -> None:
    """Thick coloured left border (used for callouts)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:color"), color)
    borders.append(left)
    for side in ("top", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def _run(p, text, *, bold=False, color=None, size=None, mono=False, italic=False):
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if size:
        r.font.size = Pt(size)
    r.font.name = MONO if mono else SANS
    if mono:
        rPr = r._element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
            rFonts.set(qn(f"w:{attr}"), MONO)
    return r


def _para(doc_or_cell, text="", *, style=None, align=None, space_after=None,
          bold=False, color=None, size=None, mono=False):
    p = doc_or_cell.add_paragraph(style=style) if style else doc_or_cell.add_paragraph()
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        _run(p, text, bold=bold, color=color, size=size, mono=mono)
    return p


def _hr(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CBD5E1")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# High-level building blocks
# ---------------------------------------------------------------------------

def add_hero(doc, title: str, subtitle_pills: list[str], blurb: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    _shade(cell, BLUE)
    _set_cell_borders(cell, color=BLUE, size="4")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _run(p, title, bold=True, color="FFFFFF", size=22)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _run(p, "  •  ".join(subtitle_pills), color="DBEAFE", size=10, italic=True)

    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _run(p, blurb, color="FFFFFF", size=11)
    doc.add_paragraph()


def add_kpi_grid(doc, cards: list[tuple[str, str, str]]) -> None:
    """cards: [(emoji, title, body), ...] — laid out in a 1×N table."""
    table = doc.add_table(rows=1, cols=len(cards))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    col_w = Cm(16 / len(cards))
    for col in table.columns:
        col.width = col_w
    for cell, (emoji, title, body) in zip(table.rows[0].cells, cards):
        cell.width = col_w
        _shade(cell, KPI_BG)
        _left_accent(cell, BLUE, size="24")
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, emoji, size=18)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, title, bold=True, color=SLATE_DEEP, size=11)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _run(p, body, color=SLATE, size=9.5)
    doc.add_paragraph()


def add_section_heading(doc, emoji: str, title: str, *, level: int = 1) -> None:
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    _run(h, f"{emoji}  {title}", bold=True, color=BLUE, size=16 if level == 1 else 13)


def add_callout(doc, kind: str, title: str, body: str) -> None:
    palette = {
        "tip":      (CALLOUT_TIP, GREEN, "💡 TIP"),
        "note":     (CALLOUT_NOTE, BLUE, "📌 NOTE"),
        "warning":  (CALLOUT_WARN, AMBER, "⚠️ WARNING"),
        "important":(CALLOUT_IMPORTANT, RED, "❗ IMPORTANT"),
    }
    fill, accent, default_title = palette[kind]
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    _shade(cell, fill)
    _left_accent(cell, accent, size="28")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    _run(p, title or default_title, bold=True, color=accent, size=9)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _run(p, body, color=SLATE_DEEP, size=10.5)
    doc.add_paragraph()


def add_code_block(doc, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Cm(16)
    cell = table.rows[0].cells[0]
    cell.width = Cm(16)
    _shade(cell, CODE_BG)
    _set_cell_borders(cell, color="D6DEE8", size="4")
    cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
    for line in code.split("\n"):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        _run(p, line if line else " ", mono=True, size=9.5, color=SLATE_DEEP)
    doc.add_paragraph()


def add_table(doc, headers: list[str], rows: list[list[str]],
              *, col_widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.autofit = False
    if col_widths_cm:
        for col, w in zip(table.columns, col_widths_cm):
            col.width = Cm(w)
    # Header row
    for cell, h in zip(table.rows[0].cells, headers):
        if col_widths_cm:
            cell.width = Cm(col_widths_cm[headers.index(h)])
        _shade(cell, BLUE)
        _set_cell_borders(cell, color=BLUE_LIGHT, size="4")
        cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _run(p, h, bold=True, color="FFFFFF", size=10)
    # Body
    for ri, row in enumerate(rows):
        for cell, val in zip(table.rows[ri + 1].cells, row):
            if col_widths_cm:
                cell.width = Cm(col_widths_cm[row.index(val) if val in row else 0])
            if ri % 2 == 1:
                _shade(cell, ROW_ALT)
            _set_cell_borders(cell, color="DDE5EE", size="2")
            cell.paragraphs[0]._p.getparent().remove(cell.paragraphs[0]._p)
            # Render value: support inline `code` markup with backticks
            p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            parts = val.split("`")
            for j, part in enumerate(parts):
                if not part:
                    continue
                if j % 2 == 1:
                    _run(p, part, mono=True, size=9.5, color=BLUE)
                else:
                    _run(p, part, color=SLATE, size=10)
    doc.add_paragraph()


def add_bullets(doc, items: list[str], style: str = "List Bullet") -> None:
    for it in items:
        p = doc.add_paragraph(style=style)
        parts = it.split("`")
        for j, part in enumerate(parts):
            if not part:
                continue
            if j % 2 == 1:
                _run(p, part, mono=True, size=10, color=BLUE)
            else:
                _run(p, part, color=SLATE, size=11)


def add_para(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    parts = text.split("`")
    for j, part in enumerate(parts):
        if not part:
            continue
        if j % 2 == 1:
            _run(p, part, mono=True, size=10, color=BLUE)
        else:
            _run(p, part, color=SLATE, size=11)


def set_default_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = SANS
    style.font.size = Pt(11)


# ---------------------------------------------------------------------------
# User doc
# ---------------------------------------------------------------------------

USER_TOC = [
    ("🧠", "What is Claude Code?"),
    ("✨", "The headline"),
    ("1️⃣", "Install Node.js"),
    ("2️⃣", "Install Claude Code"),
    ("3️⃣", "Configure five env vars"),
    ("4️⃣", "Take it for a spin"),
    ("5️⃣", "Your first real task"),
    ("🛡️", "Permission model"),
    ("⌨️", "Keyboard shortcuts"),
    ("📖", "Recipe gallery"),
    ("📝", "The CLAUDE.md trick"),
    ("🆚", "VS Code extension"),
    ("🔄", "Switching back to real Claude"),
    ("⚠️", "Caveats"),
    ("🌐", "Network gotchas"),
    ("❓", "FAQ"),
    ("🆘", "Where to get help"),
    ("📚", "References"),
]


def build_user_doc() -> Document:
    doc = Document()
    set_default_styles(doc)

    add_hero(
        doc,
        "🤖  Claude Code → Our Internal LLM",
        ["status: live", "backend: internal", "corp network only", "cost: $0"],
        "Run Anthropic's Claude Code CLI against our self-hosted models. "
        "No API key. No usage caps. No per-token billing. Just point and go. ⚡",
    )

    add_kpi_grid(doc, [
        ("🔒", "100% Local",
         "Every prompt, every response, every byte stays on hardware we own."),
        ("💰", "$0 Cost",
         "No API key, no per-token billing, no usage caps."),
        ("🏢", "Our Infra",
         "Runs on Macs we own, fronted by lab-vm1 on the corp network."),
    ])

    add_section_heading(doc, "📑", "Table of contents")
    add_para(doc, "Click any entry below to jump to that section. "
                  "For an always-on outline sidebar in Word: View → Navigation Pane (Windows) or View → Sidebar → Navigation (Mac).")
    for emoji, title in USER_TOC:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, f"→  {emoji}  {title}", color=BLUE_LIGHT, size=10.5)
    _hr(doc)

    add_section_heading(doc, "🧠", "What is Claude Code?")
    add_para(doc,
        "Claude Code is a terminal AI pair programmer. You run `claude` in a project directory, "
        "type what you want in plain English, and it reads files, writes code, runs commands, "
        "and shows you a diff before saving. Think \"a teammate who pair-programs in your repo,\" "
        "not \"a chatbot in a browser tab.\"")
    add_para(doc,
        "With this setup, all of that runs against our self-hosted models instead of Anthropic's "
        "cloud. Same CLI, same workflow — just our hardware doing the thinking.")
    _hr(doc)

    add_section_heading(doc, "✨", "The headline")
    add_para(doc,
        "Three local models are exposed through one endpoint.")
    add_table(doc,
        ["Model id", "Best for", "Notes"],
        [
            ["glm-4.7-flash",
             "Default — coding, tool use, chat",
             "Quickest first token. The Opus and Sonnet picker tiers both point here."],
            ["qwen2.5-coder-32b",
             "Coding-heavy sessions with lots of tool calls",
             "Code-tuned 32B; alternative if GLM-Flash misbehaves on your task."],
            ["laguna",
             "Long-form prose, summaries",
             "Runs via Ollama; slower first hit (cold load). Wired to Haiku tier."],
        ],
        col_widths_cm=[5.5, 5.5, 5.0],
    )
    _hr(doc)

    add_section_heading(doc, "⏱️", "Speed & latency — what to expect")
    add_callout(doc, "important", "❗ READ THIS BEFORE YOU JUDGE",
        "Every turn re-prefills the entire conversation. There is no prompt cache locally today. "
        "Plan for 1–3 minutes per turn, not the sub-second feel of anthropic.com. "
        "Your first \"Hi\" can take 2+ minutes — that's the floor, not a bug.")
    add_para(doc, "Measured on studio1 (Apple Silicon, GLM-4.7-Flash 8-bit) with Claude Code's stock system prompt + ~60 tools:")
    add_table(doc,
        ["Turn", "Typical latency", "Why"],
        [
            ["First \"Hi\" in a fresh session",
             "≈ 90–150 s",
             "Prefill of system prompt + tool definitions (~9K tokens) at ~90 tok/s."],
            ["Tool-using turn (read a file, suggest an edit)",
             "≈ 90–180 s",
             "Same prefill plus the file you just attached + the prior turns."],
            ["Long session (10+ turns, large files in context)",
             "Grows turn-over-turn",
             "Conversation keeps re-prefilling. Use `/compact` and `/clear` proactively."],
        ],
        col_widths_cm=[5.5, 3.5, 7.0],
    )
    add_section_heading(doc, "🤔", "Why so much slower than anthropic.com?", level=2)
    add_bullets(doc, [
        "Prompt caching — Anthropic caches your system prompt and tools server-side, so a repeat turn skips prefill entirely (sub-second first token). Our stack doesn't have this yet (the underlying flag is broken in our current mlx-lm version).",
        "Hardware — a Mac Studio is not a datacenter GPU. Cloud Claude runs on accelerator clusters with orders-of-magnitude more memory bandwidth.",
        "Model size — GLM-4.7-Flash is ~30 GB on disk; Opus / Sonnet are far larger and run on far bigger machines. Smaller model partly compensates for the slower hardware, but only partly.",
    ])
    add_section_heading(doc, "✅", "Good fit for", level=2)
    add_bullets(doc, [
        "Learning Claude Code's workflow without burning Anthropic credits.",
        "Single-file edits, code review, explanations of opaque code.",
        "Pre-PR self-review on small diffs.",
        "Anything where data shouldn't leave the LAN.",
    ])
    add_section_heading(doc, "🚫", "Less good fit for", level=2)
    add_bullets(doc, [
        "Tight iterative loops on large files (each turn pays full prefill).",
        "Multi-file refactors that need long agentic chains.",
        "Anything where you'd notice a 60-second wait every turn.",
    ])
    add_para(doc, "For those, switch back to real Claude — see \"Switching back to real Claude\" below.")
    _hr(doc)

    add_section_heading(doc, "1️⃣", "Install Node.js")
    add_para(doc, "You need Node 18+ (LTS recommended).")

    add_section_heading(doc, "🪟", "Windows", level=2)
    add_callout(doc, "tip", "🏢 CORP-MANAGED LAPTOP — TRY SOFTWARE CENTER FIRST",
        "Open Software Center (Start menu → 'Software Center'), search for 'Claude', click Install. "
        "If it's there, you're done — no admin prompt, no PATH wrangling. Skip ahead to Step 3 (env vars). "
        "The rollout is still in progress though, so most laptops don't have it yet — if Claude isn't "
        "listed for you, use the winget steps below to install Node, then continue to Step 2 for the "
        "npm install one-liner.")
    add_para(doc, "If Software Center doesn't list Claude — install Node manually. "
                  "The flags below work without admin rights and on corp Wi-Fi:")
    add_code_block(doc, "winget install OpenJS.NodeJS.LTS --source winget --scope user")
    add_callout(doc, "note", "📘 WHY THOSE TWO FLAGS",
        "--source winget pins the Microsoft 'winget' source instead of the default 'msstore', which "
        "fails with 0x8a15005e (server certificate did not match) on corp Wi-Fi because SSL inspection "
        "breaks the Microsoft Store source. --scope user installs Node into your profile only and "
        "modifies your user PATH; no admin elevation needed. Close and reopen your terminal afterwards "
        "so the new PATH takes effect.")

    add_section_heading(doc, "🍎", "macOS", level=2)
    add_code_block(doc, "brew install node")

    add_section_heading(doc, "🐧", "Linux (Ubuntu / Debian)", level=2)
    add_code_block(doc, "sudo apt install -y nodejs npm")

    add_para(doc, "Sanity check:")
    add_code_block(doc, "node -v && npm -v")
    _hr(doc)

    add_section_heading(doc, "2️⃣", "Install Claude Code")
    add_callout(doc, "tip", "🏢 INSTALLED VIA SOFTWARE CENTER ON WINDOWS?",
        "You're done — SC installs Claude Code alongside Node. Skip to Step 3 (env vars).")
    add_para(doc, "Otherwise (Software Center doesn't list Claude yet, or you're on Mac/Linux), "
                  "same one-liner everywhere:")
    add_code_block(doc, "npm install -g @anthropic-ai/claude-code")
    add_para(doc, "Verify:")
    add_code_block(doc, "claude --version")
    _hr(doc)

    add_section_heading(doc, "3️⃣", "Configure five env vars")
    add_para(doc, "These point Claude Code at our internal endpoint and pick which local model "
                  "each tier resolves to. Get the API key from the team lead.")

    add_section_heading(doc, "🍎🐧", "macOS / Linux — make it permanent", level=2)
    add_para(doc, "Open `~/.zshrc` (or `~/.bashrc` on bash) in your editor — pick whichever you have:")
    add_code_block(doc,
        "subl ~/.zshrc      # Sublime Text\n"
        "code ~/.zshrc      # VS Code\n"
        "nano ~/.zshrc      # nano (no install needed)")
    add_para(doc, "Append these five lines and save:")
    add_code_block(doc,
        f"export ANTHROPIC_BASE_URL={LLM_PUBLIC_URL}\n"
        "export ANTHROPIC_AUTH_TOKEN=<your-bearer-token>\n"
        "export ANTHROPIC_DEFAULT_OPUS_MODEL=glm-4.7-flash\n"
        "export ANTHROPIC_DEFAULT_SONNET_MODEL=glm-4.7-flash\n"
        "export ANTHROPIC_DEFAULT_HAIKU_MODEL=laguna")
    add_para(doc, "Reload:")
    add_code_block(doc, "source ~/.zshrc")

    add_section_heading(doc, "🪟", "Windows — PowerShell, persistent (user-level)", level=2)
    add_code_block(doc,
        "[System.Environment]::SetEnvironmentVariable("
        f"'ANTHROPIC_BASE_URL', '{LLM_PUBLIC_URL}', 'User')\n"
        "[System.Environment]::SetEnvironmentVariable("
        "'ANTHROPIC_AUTH_TOKEN', '<your-bearer-token>', 'User')\n"
        "[System.Environment]::SetEnvironmentVariable("
        "'ANTHROPIC_DEFAULT_OPUS_MODEL', 'glm-4.7-flash', 'User')\n"
        "[System.Environment]::SetEnvironmentVariable("
        "'ANTHROPIC_DEFAULT_SONNET_MODEL', 'glm-4.7-flash', 'User')\n"
        "[System.Environment]::SetEnvironmentVariable("
        "'ANTHROPIC_DEFAULT_HAIKU_MODEL', 'laguna', 'User')")
    add_para(doc, "Then close and reopen the terminal.")

    add_section_heading(doc, "🩺", "Verify — confirm the values are set", level=2)
    add_para(doc, "Open a fresh terminal (so it picks up the new vars), then run the line for "
                  "your shell. All five values should print non-empty — if any are blank, the vars "
                  "didn't persist and `claude` will fall back to api.anthropic.com.")
    add_para(doc, "🪟 PowerShell")
    add_code_block(doc,
        "$env:ANTHROPIC_BASE_URL; $env:ANTHROPIC_AUTH_TOKEN.Substring(0,8) + \"...\"; "
        "$env:ANTHROPIC_DEFAULT_OPUS_MODEL; $env:ANTHROPIC_DEFAULT_SONNET_MODEL; $env:ANTHROPIC_DEFAULT_HAIKU_MODEL")
    add_para(doc, "🪟 Windows CMD")
    add_code_block(doc,
        "echo %ANTHROPIC_BASE_URL% & echo %ANTHROPIC_AUTH_TOKEN:~0,8%... & "
        "echo %ANTHROPIC_DEFAULT_OPUS_MODEL% & echo %ANTHROPIC_DEFAULT_SONNET_MODEL% & echo %ANTHROPIC_DEFAULT_HAIKU_MODEL%")
    add_para(doc, "🍎🐧 macOS / Linux")
    add_code_block(doc,
        "echo \"$ANTHROPIC_BASE_URL\"; echo \"${ANTHROPIC_AUTH_TOKEN:0:8}...\"; "
        "echo \"$ANTHROPIC_DEFAULT_OPUS_MODEL\"; echo \"$ANTHROPIC_DEFAULT_SONNET_MODEL\"; echo \"$ANTHROPIC_DEFAULT_HAIKU_MODEL\"")
    add_para(doc, "The token is truncated to its first 8 characters so you can sanity-check it's "
                  "set without echoing the full secret to your terminal scrollback.")
    _hr(doc)

    add_section_heading(doc, "4️⃣", "Take it for a spin")
    add_code_block(doc, "cd ~/some/repo\nclaude")
    add_para(doc, f"Inside the prompt, type `/status` — confirm `ANTHROPIC_BASE_URL` shows {LLM_PUBLIC_URL}.")
    add_para(doc, "Then say hi:")
    add_code_block(doc, "> hi, what model are you?")
    _hr(doc)

    add_section_heading(doc, "🎛️", "Switch models on the fly")
    add_para(doc,
        "The five env vars in step 3 set your defaults. To try a different model "
        "for one session without editing your shell config, override at launch.")

    add_section_heading(doc, "①", "Env-var prefix (one session)", level=2)
    add_code_block(doc, "ANTHROPIC_MODEL=qwen2.5-coder-32b claude")
    add_para(doc,
        "`ANTHROPIC_MODEL` takes precedence over the per-tier vars. Banner will read "
        "`qwen2.5-coder-32b[1m]` instead of the default. Closes the override when the "
        "session ends.")

    add_section_heading(doc, "②", "CLI flag (one session)", level=2)
    add_code_block(doc, "claude --model qwen2.5-coder-32b")
    add_para(doc, "Same effect as the env-var prefix; pick whichever feels natural.")

    add_section_heading(doc, "③", "/model picker (mid-session)", level=2)
    add_para(doc,
        "Inside Claude Code, `/model` switches between the Opus / Sonnet / Haiku tiers. "
        "Each tier resolves to whichever id you set in `ANTHROPIC_DEFAULT_{OPUS,SONNET,HAIKU}_MODEL`. "
        "Useful if you wired Sonnet to `qwen2.5-coder-32b` — `/model` then becomes a "
        "GLM ↔ Qwen toggle without restarting.")

    add_callout(doc, "tip", "💡 WHEN TO REACH FOR EACH MODEL",
        "Default to glm-4.7-flash. If a coding turn comes back empty, drops a tool call, or "
        "the model talks about code instead of writing it, retry with qwen2.5-coder-32b — "
        "it's code-tuned and doesn't have a thinking-mode prefix that can swallow short answers.")
    _hr(doc)

    add_section_heading(doc, "5️⃣", "Your first real task — a 2-minute tutorial")
    add_section_heading(doc, "①", "Open a project (or start fresh)", level=2)
    add_para(doc, "`cd` into an existing repo and run `claude` there — it'll see your code and edit files in place.")
    add_section_heading(doc, "②", "Ask it to build something concrete", level=2)
    add_para(doc, "At the prompt, paste:")
    add_code_block(doc,
        "> add a simple `is_palindrome(s)` helper in utils.py with a unit test. "
        "Strip non-alphanumerics, ignore case.")
    add_para(doc, "Claude Code will plan the file, write it, show you the diff, and ask before saving. "
                  "Press `y` to accept, `e` to edit inline, or describe what you want changed.")
    add_section_heading(doc, "③", "Iterate", level=2)
    add_para(doc, "Ask it to run the test:")
    add_code_block(doc, "> run the test")
    add_para(doc, "It'll create the test file, run it, and show output. If a test fails, ask it to fix and re-run.")
    add_section_heading(doc, "④", "Useful slash commands", level=2)
    add_table(doc,
        ["Slash command", "What it does"],
        [
            ["/status",  "Show the resolved env vars, model, and working directory."],
            ["/clear",   "Reset the conversation in this session."],
            ["/compact", "Compress earlier turns to free up context."],
            ["/help",    "Full command reference."],
        ],
        col_widths_cm=[5.5, 10.5],
    )
    add_section_heading(doc, "💡", "Tips for getting good results", level=2)
    add_bullets(doc, [
        "Be specific. \"Refactor this function for clarity, keep the public signature\" beats \"make it better.\"",
        "Smaller scopes win. One file at a time. Long multi-file refactors are where local models struggle.",
        "Show, don't tell. Pasting a small example of the desired output usually beats describing it.",
        "Verify the diff. The model occasionally hallucinates an import or path — read before accepting.",
        "If it gets stuck in a loop, `/clear` and rephrase. Don't keep nudging a confused conversation.",
    ])
    _hr(doc)

    add_section_heading(doc, "🛡️", "The permission model")
    add_para(doc, "Claude Code never silently edits files or runs commands. Every side-effect goes through a prompt. It will ask before:")
    add_bullets(doc, [
        "Writing to or deleting a file",
        "Running shell commands (especially anything destructive)",
        "Installing packages or making network calls",
    ])
    add_callout(doc, "tip", "💡 SHIFT-TAB",
        "Press Shift+Tab to cycle between three modes: ask (default), auto-accept edits, plan-only. "
        "The current mode is shown in the bottom-of-terminal status line.")
    _hr(doc)

    add_section_heading(doc, "⌨️", "Keyboard shortcuts cheat-sheet")
    add_table(doc,
        ["Shortcut", "What it does"],
        [
            ["Esc",       "Cancel the current generation."],
            ["Shift+Tab", "Cycle permission modes (ask / auto-accept / plan)."],
            ["Ctrl+R",    "Toggle thinking mode (show/hide reasoning)."],
            ["Ctrl+C",    "Quit Claude Code."],
            ["@filename", "Anchor a specific file as context."],
            ["!cmd",      "Drop into a shell, run cmd, pipe output back."],
        ],
        col_widths_cm=[4.0, 12.0],
    )
    _hr(doc)

    add_section_heading(doc, "📖", "Recipe gallery")
    add_para(doc, "\"Write a script\" is the obvious one. Higher-value patterns most newcomers don't think to try:")
    add_bullets(doc, [
        "Onboard yourself to an unfamiliar repo: ask for a tour of `src/`, then have it explain the data model.",
        "Diagnose a failing test: paste the failure, ask for a hypothesis and a minimal repro.",
        "Write tests for code you didn't write: \"add 5 unit tests covering edge cases of `parse_xyz`.\"",
        "Pre-PR self-review: \"review the diff against main and flag anything that'd embarrass me.\"",
        "Refactor with intent: \"extract the validation logic into a pure function, keep behavior identical.\"",
        "Explain something opaque: paste a regex, a SQL plan, a stack trace — ask for plain-English.",
        "Add docstrings to legacy code: \"add Google-style docstrings to all public methods in this file.\"",
        "Translate between languages: \"port this Python function to Go, idiomatic.\"",
    ])
    _hr(doc)

    add_section_heading(doc, "📝", "The CLAUDE.md trick — make it know your project")
    add_para(doc,
        "Drop a `CLAUDE.md` at the root of your repo and Claude Code reads it on every run. "
        "Teach it your conventions once instead of repeating yourself in every prompt.")
    add_para(doc, "Generate a starter automatically:")
    add_code_block(doc, "claude /init")
    add_para(doc,
        "AGENTS.md is the cross-tool version of the same idea — read by Claude Code, Cursor, Aider, "
        "Codex CLI, Gemini CLI. Same format, neutral name. Either works; teams that mix tools usually use AGENTS.md.")
    _hr(doc)

    add_section_heading(doc, "🆚", "Prefer an IDE? Use the VS Code extension")
    add_para(doc,
        "Install the \"Claude Code\" extension from the VS Code marketplace. Same env-var config, same "
        "slash commands, same permission model — just rendered as a side panel in VS Code.")
    add_bullets(doc, [
        "Diffs render inline in the editor",
        "Selected code is auto-attached as context",
        "Same env vars work — no extra config",
    ])
    _hr(doc)

    add_section_heading(doc, "🔄", "Switching back to real Claude")
    add_para(doc, "When you want real Opus / Sonnet for the heavy lifting, unset the five env vars:")
    add_section_heading(doc, "🍎🐧", "macOS / Linux", level=2)
    add_code_block(doc,
        "unset ANTHROPIC_BASE_URL ANTHROPIC_AUTH_TOKEN \\\n"
        "      ANTHROPIC_DEFAULT_OPUS_MODEL ANTHROPIC_DEFAULT_SONNET_MODEL ANTHROPIC_DEFAULT_HAIKU_MODEL")
    add_section_heading(doc, "🪟", "Windows PowerShell", level=2)
    add_code_block(doc,
        "[System.Environment]::SetEnvironmentVariable('ANTHROPIC_BASE_URL',            $null, 'User')\n"
        "[System.Environment]::SetEnvironmentVariable('ANTHROPIC_AUTH_TOKEN',             $null, 'User')\n"
        "[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_OPUS_MODEL',   $null, 'User')\n"
        "[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_SONNET_MODEL', $null, 'User')\n"
        "[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_HAIKU_MODEL',  $null, 'User')")
    add_para(doc, "Then `claude login` to authenticate with your Anthropic account or API key.")
    _hr(doc)

    add_section_heading(doc, "⚠️", "Caveats — read this once")
    add_callout(doc, "important", "❗ NOT CLAUDE",
        "These are smaller open-weight models running on Mac hardware. Expect a quality drop from "
        "Opus / Sonnet, especially on long multi-file refactors and complex tool chains. Use them for "
        "what they're good at; reach for real Claude when the task warrants it.")
    add_bullets(doc, [
        "Smaller context window than Claude. Use /compact often, /clear when conversations drift.",
        "Tool-call reliability varies by model. If it loops or emits malformed JSON, simplify the prompt.",
        "Every turn re-prefills the conversation (no local prompt cache today) — see the Speed & latency section above.",
        "If something's genuinely broken (not just \"lower quality than Claude\"), file it.",
    ])
    _hr(doc)

    add_section_heading(doc, "🌐", "Network gotchas")
    add_bullets(doc, [
        "Reachable from corp WiFi/wired LAN, or from home over the corp VPN.",
        "Not reachable from the corporate proxy-only or fully off-VPN networks.",
        "If `claude` hangs at startup, check that you can reach the gateway (below).",
    ])
    add_section_heading(doc, "🩺", "Quick reachability check", level=2)
    add_code_block(doc,
        "curl -H \"Authorization: Bearer $ANTHROPIC_AUTH_TOKEN\" "
        "$ANTHROPIC_BASE_URL/v1/models")
    add_para(doc, "Expected: a JSON list with `glm-4.7-flash` and `laguna`. "
                  "If you get 401, the API key is wrong. If the connection times out, you're off-network.")
    _hr(doc)

    add_section_heading(doc, "❓", "FAQ — things people ping me about")
    faq = [
        ("Q: 🪟 What's the easiest way to install on a corp-managed Windows laptop?",
         "A: Try Software Center first — search for 'Claude'; if it's listed, click Install and you're done (no admin prompt, no PATH fiddling, jump to Step 3). The SC rollout is still in progress though, so most laptops don't have it yet — if Claude isn't there for you, fall back to the winget + npm steps in Step 1 and Step 2."),
        ("Q: 🪟 winget install fails with 'server certificate did not match' (0x8a15005e).",
         "A: If Claude is in your Software Center, that path skips winget entirely (search 'Claude' → Install). Otherwise: corp SSL inspection breaks the Microsoft Store source. Pin winget explicitly with --source winget. The error message lists 'winget' as a working source — that's the one to use."),
        ("Q: 🪟 I don't have admin rights on my Windows laptop — can I still install?",
         "A: Yes. Easiest path is Software Center if Claude is listed for you (no admin needed). If it isn't, use winget with --scope user: `winget install OpenJS.NodeJS.LTS --source winget --scope user`. Node installs into your profile and only your user PATH is modified. Last-resort fallback is the portable zip from https://nodejs.org/dist/: extract to %USERPROFILE%\\nodejs, then add that folder to your *user* PATH (System Properties → Environment Variables → User variables → Path → New)."),
        ("Q: It told me it can't access the internet. Is something broken?",
         "A: No — that's expected. The local models run fully offline by design. To fetch a webpage, run `!curl ...` and pipe the output back in."),
        ("Q: My output got cut off mid-sentence.",
         "A: You hit the context window. `/compact` to compress earlier turns, or `/clear` to start fresh. Local context is smaller than Claude's — keep conversations focused."),
        ("Q: It hallucinated a function / import / file path.",
         "A: Known weakness of smaller models. Always read the diff before accepting. If it keeps inventing things, scope down — ask about one file at a time and use @filename to anchor it."),
        ("Q: It refused to do something benign.",
         "A: Rephrase. Adding context like \"this is my own project, the file is mine to edit\" usually unblocks it."),
        ("Q: Tool calls are failing or producing malformed JSON.",
         "A: Simplify the request — break it into smaller steps."),
        ("Q: Every turn feels slow. Why isn't the second one faster?",
         "A: Anthropic's cloud caches your system prompt + tools server-side, so repeat turns skip prefill. We don't have that locally yet — every turn re-prefills the conversation. The first turn pays for ~9K tokens of system prompt + tools; each subsequent turn pays for that plus everything since. Use `/compact` and `/clear` to keep context lean. See the Speed & latency section near the top of this doc for measured numbers."),
        ("Q: Can I use this for confidential / customer data?",
         "A: Prompts and responses stay on our hardware — nothing is sent to Anthropic or any third party. Follow normal data-handling policy, but you don't have the \"sending to a vendor\" worry."),
        ("Q: Can I run two `claude` sessions at once?",
         "A: Yes. Open a second terminal in a different repo. Independent contexts."),
        ("Q: How do I see what env vars Claude Code is actually using?",
         "A: Type `/status` inside the prompt — shows resolved BASE_URL, model, permission mode, working directory."),
    ]
    for q, a in faq:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, q, bold=True, color=SLATE_DEEP, size=10.5)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _run(p, a, color=SLATE, size=10.5)
    _hr(doc)

    add_section_heading(doc, "🆘", "Where to get help")
    add_bullets(doc, [
        "Inside Claude Code: `/help` for the full command list.",
        "Setup or routing issues (\"my requests aren't reaching the internal endpoint\"): ping the team lead.",
        "Bugs in the model itself (refused a reasonable task, kept hallucinating): file with the team lead so we can patch the bridge or tweak the config.",
        "Questions about Claude Code itself: https://docs.anthropic.com/en/docs/claude-code — public docs apply, just substitute our endpoint for the Anthropic API.",
    ])
    _hr(doc)

    add_section_heading(doc, "📚", "References")
    add_bullets(doc, [
        "Claude Code official docs: https://docs.anthropic.com/en/docs/claude-code",
        "VS Code extension: search \"Claude Code\" in the marketplace",
        "Internal admin guide: docs/CLAUDE_CODE_ADMIN_SETUP.docx",
    ])
    return doc


# ---------------------------------------------------------------------------
# Admin doc
# ---------------------------------------------------------------------------

ADMIN_TOC = [
    ("🧱", "Architecture"),
    ("📋", "Components & ports"),
    ("⚙️", "Service operations"),
    ("🔧", "Configuration files"),
    ("➕", "Adding a new model"),
    ("🖥️", "Mac backends"),
    ("🔐", "Security & secrets"),
    ("🚨", "Troubleshooting"),
    ("💾", "Backup & recovery"),
    ("📚", "References"),
]


def build_admin_doc() -> Document:
    doc = Document()
    set_default_styles(doc)

    add_hero(
        doc,
        "🛠️  Claude Code Local Stack — Admin Guide",
        ["scope: lab-vm1 + Mac fleet", "consumers: any Claude Code client on corp net", "status: production"],
        "Operating manual for the router that lets Claude Code clients talk to our self-hosted "
        "vllm-mlx and Ollama backends. Covers architecture, day-2 ops, troubleshooting, and how to "
        "add new models.",
    )

    add_kpi_grid(doc, [
        ("🧠", "2 backends",
         "studio1 GLM-Flash (vllm-mlx), studio1 Laguna (Ollama)."),
        ("🚪", "1 endpoint",
         "lab-vm1:8051 — single URL clients point at; bearer-auth gated."),
        ("⚙️", "2 services",
         "ir-claude-router (8050) + ir-claude-router-shim (8051) on lab-vm1."),
    ])

    add_section_heading(doc, "📑", "Table of contents")
    for emoji, title in ADMIN_TOC:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        _run(p, f"→  {emoji}  {title}", color=BLUE_LIGHT, size=10.5)
    _hr(doc)

    add_section_heading(doc, "🧱", "Architecture")
    add_para(doc,
        "Two services on lab-vm1, two Mac backends — both on studio1. The shim is the public face; "
        "ccr is the internal translator that ships the requests onward.")
    add_code_block(doc,
        "[claude client]                                                  \n"
        "      │  ANTHROPIC_BASE_URL=http://lab-vm1:8051                   \n"
        "      ▼                                                          \n"
        "  lab-vm1:8051   ir-claude-router-shim   (FastAPI)               \n"
        "      │   • exposes GET /v1/models for SDK / curl discovery     \n"
        "      │   • rewrites friendly id → provider,model                \n"
        "      │   • bearer-auth gate                                     \n"
        "      ▼                                                          \n"
        "  127.0.0.1:8050   ir-claude-router      (claude-code-router)    \n"
        "      │   • Anthropic /v1/messages → OpenAI /v1/chat/completions \n"
        "      │   • routes by `provider,model` to one of two upstreams   \n"
        "      ▼                                                          \n"
        "  ┌───────────────┬──────────────────────┐                        \n"
        "  │ 127.0.0.1:8024│ 127.0.0.1:8022       │                        \n"
        "  │ studio1 GLM   │ studio1 Laguna       │                        \n"
        "  │ vllm-mlx      │ Ollama               │                        \n"
        "  └───────────────┴──────────────────────┘                        \n"
        "  (each is a reverse SSH tunnel from studio1 into lab-vm1)       ")

    add_callout(doc, "note", "📌 WHY TWO LAYERS",
        "ccr (claude-code-router, an npm package) handles the Anthropic↔OpenAI translation and "
        "multi-provider routing — but it expects requests in `provider,model` form and doesn't "
        "expose /v1/models for discovery. The shim adds /v1/models (for SDK / curl / IDE-plugin "
        "enumeration), translates friendly model ids to ccr's `provider,model` form on incoming "
        "/v1/messages, and gates everything behind a bearer token. Note: Claude Code's /model "
        "picker is hardcoded to Opus / Sonnet / Haiku and does NOT read /v1/models — users wire "
        "each tier to one of our ids via ANTHROPIC_DEFAULT_{OPUS,SONNET,HAIKU}_MODEL. ~120 lines "
        "of FastAPI; no logic of its own beyond the rewrite.")
    _hr(doc)

    add_section_heading(doc, "📋", "Components & ports")
    add_table(doc,
        ["Service", "Port", "Purpose", "Source"],
        [
            ["ir-claude-router",      "8050",
             "claude-code-router (npm). Anthropic↔OpenAI + provider routing.",
             "~/.claude-code-router/config.json"],
            ["ir-claude-router-shim", "8051",
             "FastAPI front door. /v1/models (discovery), id-rewrite, bearer auth.",
             "deployment/claude_router_shim.py"],
        ],
        col_widths_cm=[4.0, 1.5, 6.5, 4.0],
    )

    add_para(doc, "Backends (each lives behind a reverse SSH tunnel from studio1):")
    add_table(doc,
        ["Tunnel port", "Mac", "Engine", "Model"],
        [
            ["8024", "studio1",  "vllm-mlx", "mlx-community/GLM-4.7-Flash-8bit"],
            ["8022", "studio1",  "Ollama",   "laguna-xs.2:q8_0"],
        ],
        col_widths_cm=[2.5, 3.0, 3.0, 7.5],
    )
    _hr(doc)

    add_section_heading(doc, "⚙️", "Service operations")
    add_section_heading(doc, "🩺", "Status", level=2)
    add_code_block(doc,
        "systemctl --user status ir-claude-router ir-claude-router-shim\n"
        "systemctl --user is-active ir-claude-router ir-claude-router-shim")

    add_section_heading(doc, "🔄", "Restart", level=2)
    add_code_block(doc,
        "# config.json edit → restart router\n"
        "systemctl --user restart ir-claude-router\n\n"
        "# claude_router_shim.py edit → restart shim\n"
        "systemctl --user restart ir-claude-router-shim")

    add_section_heading(doc, "📜", "Logs", level=2)
    add_table(doc,
        ["Log", "Path"],
        [
            ["Shim (Python uvicorn)",     "data/transient/logs/claude_router_shim.log"],
            ["Router (ccr — systemd capture)", "data/transient/logs/claude_router.log"],
            ["Router (ccr — pino server)", "~/.claude-code-router/logs/ccr-*.log"],
        ],
        col_widths_cm=[6.5, 9.5],
    )
    _hr(doc)

    add_section_heading(doc, "🔧", "Configuration files")

    add_section_heading(doc, "①", "~/.claude-code-router/config.json", level=2)
    add_para(doc, "ccr config — providers, model lists, routing rules. Not under git (contains live URLs).")
    add_code_block(doc,
        "{\n"
        "  \"HOST\": \"0.0.0.0\",\n"
        "  \"PORT\": 8050,\n"
        "  \"APIKEY\": \"$CCR_APIKEY\",\n"
        "  \"Providers\": [\n"
        "    { \"name\": \"glm\",    \"api_base_url\": \"http://127.0.0.1:8024/v1/chat/completions\",\n"
        "      \"api_key\": \"sk-no-key\",\n"
        "      \"models\": [\"glm-4.7-flash\"] },\n"
        "    { \"name\": \"laguna\", \"api_base_url\": \"http://127.0.0.1:8022/v1/chat/completions\",\n"
        "      \"api_key\": \"sk-no-key\",\n"
        "      \"models\": [\"laguna-xs.2:q8_0\"] }\n"
        "  ],\n"
        "  \"Router\": {\n"
        "    \"default\":     \"glm,glm-4.7-flash\",\n"
        "    \"background\":  \"glm,glm-4.7-flash\",\n"
        "    \"think\":       \"glm,glm-4.7-flash\",\n"
        "    \"longContext\": \"glm,glm-4.7-flash\"\n"
        "  }\n"
        "}")

    add_section_heading(doc, "②", "deployment/claude_router_shim.py", level=2)
    add_para(doc,
        "Two dicts at the top of the file decide what shows up in /v1/models and how aliases map to "
        "ccr providers. Edit, save, restart `ir-claude-router-shim`.")
    add_code_block(doc,
        "MODEL_MAP = {\n"
        "    \"glm-4.7-flash\":  \"glm,glm-4.7-flash\",\n"
        "    \"laguna\":         \"laguna,laguna-xs.2:q8_0\",\n"
        "}\n"
        "DISPLAY_NAMES = {\n"
        "    \"glm-4.7-flash\": \"GLM 4.7 Flash\",\n"
        "    \"laguna\":        \"Laguna xs.2\",\n"
        "}")

    add_section_heading(doc, "③", "data/transient/.env  →  CCR_APIKEY", level=2)
    add_para(doc,
        "Single bearer token validated by both the shim (own check) and ccr (forwarded). Clients send "
        "the same value as `Authorization: Bearer <CCR_APIKEY>` and as `ANTHROPIC_AUTH_TOKEN`.")
    add_callout(doc, "warning", "⚠️ KEEP THE KEY OUT OF GIT",
        "`.env` is gitignored. Don't paste the token into commit messages, PRs, or code. Distribute "
        "out-of-band (1Password, encrypted message) and rotate by regenerating then restarting both services.")
    _hr(doc)

    add_section_heading(doc, "➕", "Adding a new model")
    add_para(doc, "End-to-end walkthrough — example: gemma3 on studio2.")

    add_section_heading(doc, "①", "Stand up the model on the Mac", level=2)
    add_bullets(doc, [
        "Install vllm-mlx (or use the existing service plist as a template — see studio1 reference).",
        "Run with `--served-model-name gemma3` so the model id is human-readable; otherwise the full path is used.",
        "Bind to 127.0.0.1:<port> on the Mac. Don't expose externally.",
    ])

    add_section_heading(doc, "②", "Add a reverse SSH tunnel", level=2)
    add_para(doc,
        "On the Mac, edit `~/Library/LaunchAgents/com.ir.tunnel-to-labvm.plist` to add a new "
        "`-R <lab-vm-port>:127.0.0.1:<mac-port>` mapping. Pick a free lab-vm1 port (e.g. 8024). "
        "Reload with `launchctl kickstart -k gui/$(id -u)/com.ir.tunnel-to-labvm`.")
    add_para(doc, "Verify on lab-vm1:")
    add_code_block(doc, "curl -s http://127.0.0.1:8024/v1/models")

    add_section_heading(doc, "③", "Add provider to ccr config", level=2)
    add_para(doc, "Edit `~/.claude-code-router/config.json` — add a Providers entry:")
    add_code_block(doc,
        "{ \"name\": \"gemma\",\n"
        "  \"api_base_url\": \"http://127.0.0.1:8024/v1/chat/completions\",\n"
        "  \"api_key\": \"sk-no-key\",\n"
        "  \"models\": [\"gemma3\"] }")

    add_section_heading(doc, "④", "Add alias to the shim", level=2)
    add_para(doc, "Edit `deployment/claude_router_shim.py`:")
    add_code_block(doc,
        "MODEL_MAP[\"claude-gemma3\"]      = \"gemma,gemma3\"\n"
        "DISPLAY_NAMES[\"claude-gemma3\"]  = \"Gemma 3\"")

    add_section_heading(doc, "⑤", "Restart and verify", level=2)
    add_code_block(doc,
        "systemctl --user restart ir-claude-router ir-claude-router-shim\n\n"
        "CCR_KEY=$(grep '^CCR_APIKEY=' data/transient/.env | cut -d= -f2)\n"
        "curl -s -H \"Authorization: Bearer $CCR_KEY\" http://127.0.0.1:8051/v1/models | jq")
    add_para(doc, "Expected: the new `claude-gemma3` entry appears alongside the others. Send a test /v1/messages call to confirm routing.")
    _hr(doc)

    add_section_heading(doc, "🖥️", "Mac backends")
    add_para(doc, "All three Claude Code backends now live on studio1, behind a single reverse-tunnel session to lab-vm1. (mac-m1 still runs GLM-4.7-Flash for Pokedex + Win.AI on its own tunnel, but is no longer in the Claude Code path as of 2026-05-06.)")

    add_section_heading(doc, "🎙️", "studio1 (GLM + Laguna, two stacks)", level=2)
    add_bullets(doc, [
        "vllm-mlx GLM: `mlx-community/GLM-4.7-Flash-8bit`, parser `glm47`, reasoning `deepseek_r1`, tunnel lab-vm1:8024 → studio1:8002 (~30 GB on disk)",
        "Ollama: laguna-xs.2:q8_0, tunnel lab-vm1:8022 → studio1:11434 (~40 GB cold, KEEP_ALIVE=30s so unloads when idle)",
        "SSH backchannel: `ssh -p 2224 vvobbilichetty@127.0.0.1` from lab-vm1",
        "GLM reload: `launchctl bootout gui/$(id -u)/com.ir.vllm-mlx-glm && launchctl bootstrap gui/$(id -u) ~/Library/LaunchAgents/com.ir.vllm-mlx-glm.plist`",
        "Qwen3-32B vllm-mlx is downloaded (~33 GB) but the launchctl agent is **disabled** (2026-05-06) to avoid memory contention with GLM. Re-enable with `launchctl enable gui/$(id -u)/com.ir.vllm-mlx-qwen && launchctl bootstrap gui/$(id -u) ~/Library/LaunchAgents/com.ir.vllm-mlx-qwen.plist`.",
    ])

    add_callout(doc, "warning", "⚠️ LAUNCHCTL DOMAIN",
        "studio1's vllm-mlx runs in the `gui/$UID` domain, but the tunnel agent is in `user/$UID`. "
        "`launchctl print user/501` shows the vllm services as \"enabled\" but kickstart/bootout there "
        "fails with \"Could not find service in domain.\" Use `gui/$(id -u)/...` for vllm; "
        "`user/$(id -u)/...` for the tunnel.")
    add_callout(doc, "important", "⚡ SYSTEM-PROMPT KV CACHE PATCH",
        "vllm-mlx's `--continuous-batching` flag (which gates the engine-level prefix cache) crashes "
        "mlx-lm on first cache-hit decode with `RuntimeError: There is no Stream(gpu, X) in current thread`. "
        "We work around it with a local patch to `vllm_mlx/engine/simple.py` that adds single-slot "
        "system-prompt KV caching to the pure-LLM `stream_chat()` path. Result: ~17x speedup on cache hits "
        "(measured 9.8s cold → 0.58s hit on Qwen2.5-Coder with a 2.5K-token system prompt). "
        "Patch + idempotent apply script: `deployment/vllm_mlx_patches/`. "
        "Re-run `apply.sh` after every `pip install --upgrade vllm-mlx` and bounce the launchctl agent.")
    _hr(doc)

    add_section_heading(doc, "🔐", "Security & secrets")
    add_bullets(doc, [
        "Single bearer (`CCR_APIKEY`) validates both the shim and ccr.",
        "Shim binds 0.0.0.0 — accessible from any host on the corp network. No external exposure.",
        "ccr is bound 0.0.0.0 too but should only be hit via the shim. Future hardening: bind ccr to 127.0.0.1.",
        "All traffic between lab-vm1 and the Macs is over reverse SSH tunnels (encrypted + key-auth).",
        "vllm-mlx and Ollama on the Macs bind 127.0.0.1 only — not reachable except through the tunnel.",
    ])
    add_section_heading(doc, "🔁", "Rotating the bearer token", level=2)
    add_code_block(doc,
        "# 1. regenerate\n"
        "openssl rand -hex 32\n\n"
        "# 2. update data/transient/.env\n"
        "sed -i \"s|^CCR_APIKEY=.*|CCR_APIKEY=<new-value>|\" data/transient/.env\n\n"
        "# 3. restart\n"
        "systemctl --user restart ir-claude-router ir-claude-router-shim\n\n"
        "# 4. distribute new value to users")
    _hr(doc)

    add_section_heading(doc, "🚨", "Troubleshooting")
    add_table(doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["401 Unauthorized on /v1/models",
             "Wrong or missing bearer token",
             "Check `ANTHROPIC_AUTH_TOKEN` matches `CCR_APIKEY` in `data/transient/.env`."],
            ["404 Not Found on /v1/models",
             "Client pointing at ccr (8050) instead of shim (8051)",
             "Set `ANTHROPIC_BASE_URL=http://lab-vm1:8051`."],
            ["/v1/messages returns \"fetch failed\"",
             "Upstream Mac unreachable",
             "Check tunnel: `ss -tlnp | grep 80<port>`. SSH the Mac, verify the engine is up."],
            ["Connection reset on GLM path",
             "studio1 vllm-mlx-glm not running",
             "ssh studio1, `launchctl bootstrap gui/$(id -u) ~/Library/LaunchAgents/com.ir.vllm-mlx-glm.plist`."],
            ["Picker doesn't show models",
             "Stale gateway-models cache on client",
             "Delete `~/.claude/cache/gateway-models.json` on client, restart `claude`."],
            ["Tool calls flaky",
             "Model-specific (smaller models drop tool args)",
             "Switch to glm-4.7-flash — most reliable for tool use."],
            ["Shim restart with no effect",
             "Edited `config.json` but didn't restart ccr",
             "Both services restart together for any provider change."],
        ],
        col_widths_cm=[4.5, 4.5, 7.0],
    )
    _hr(doc)

    add_section_heading(doc, "💾", "Backup & recovery")
    add_para(doc, "All four artifacts are tiny and rsync'd to lab-vm2 by the weekly backup cron:")
    add_bullets(doc, [
        "`~/.claude-code-router/config.json` — provider routing",
        "`deployment/claude_router_shim.py` — alias map (under git)",
        "`deployment/systemd/ir-claude-router*.service` — units (under git)",
        "`data/transient/.env` — `CCR_APIKEY` line",
    ])
    add_para(doc, "Recovery: restore those four files, `systemctl --user daemon-reload`, then start both services. No DB, no schema migrations, no data state to rehydrate.")
    _hr(doc)

    add_section_heading(doc, "📚", "References")
    add_bullets(doc, [
        "User-facing setup guide: `docs/CLAUDE_CODE_USER_SETUP.docx`",
        "claude-code-router repo: https://github.com/musistudio/claude-code-router",
        "Memory: `~/.claude/projects/-home-vinay-IR/memory/project_claude_code_router.md`",
        "Memory: `~/.claude/projects/-home-vinay-IR/memory/project_studio1_qwen3_vllm.md`",
        "Generator script: `misc_scripts/build_claude_code_docs.py`",
    ])

    return doc


# ---------------------------------------------------------------------------
# Markdown counterparts
# ---------------------------------------------------------------------------

USER_MD = """# 🤖 Claude Code → Our Internal LLM

> **status: live** · **backend: internal** · **corp network only** · **cost: $0**

Run Anthropic's Claude Code CLI against our self-hosted models. No API key. No usage caps. No per-token billing. Just point and go. ⚡

| 🔒 100% Local | 💰 $0 Cost | 🏢 Our Infra |
|---|---|---|
| Every prompt and response stays on hardware we own. | No API key, no per-token billing, no usage caps. | Runs on Macs we own, fronted by lab-vm1 on the corp network. |

## 📑 Table of contents
- [What is Claude Code?](#-what-is-claude-code)
- [The headline](#-the-headline)
- [Speed & latency — what to expect](#️-speed--latency--what-to-expect)
- [1️⃣ Install Node.js](#1️⃣-install-nodejs)
- [2️⃣ Install Claude Code](#2️⃣-install-claude-code)
- [3️⃣ Configure five env vars](#3️⃣-configure-five-env-vars)
- [4️⃣ Take it for a spin](#4️⃣-take-it-for-a-spin)
- [Switch models on the fly](#️-switch-models-on-the-fly)
- [5️⃣ Your first real task](#5️⃣-your-first-real-task--a-2-minute-tutorial)
- [Permission model](#️-the-permission-model)
- [Keyboard shortcuts](#️-keyboard-shortcuts-cheat-sheet)
- [Recipe gallery](#-recipe-gallery)
- [The CLAUDE.md trick](#-the-claudemd-trick--make-it-know-your-project)
- [VS Code extension](#-prefer-an-ide-use-the-vs-code-extension)
- [Switching back to real Claude](#-switching-back-to-real-claude)
- [Caveats](#️-caveats--read-this-once)
- [Network gotchas](#-network-gotchas)
- [FAQ](#-faq--things-people-ping-me-about)
- [Where to get help](#-where-to-get-help)
- [References](#-references)

---

## 🧠 What is Claude Code?
Claude Code is a terminal AI pair programmer. You run `claude` in a project directory, type what you want in plain English, and it reads files, writes code, runs commands, and shows you a diff before saving. Think "a teammate who pair-programs in your repo," not "a chatbot in a browser tab."

With this setup, all of that runs against our self-hosted models instead of Anthropic's cloud. Same CLI, same workflow — just our hardware doing the thinking.

## ✨ The headline
Three local models exposed through one endpoint.

| Model id | Best for | Notes |
|---|---|---|
| `glm-4.7-flash` | Default — coding, tool use, chat | Quickest first token. The Opus and Sonnet picker tiers both point here. |
| `qwen2.5-coder-32b` | Coding-heavy sessions with lots of tool calls | Code-tuned 32B; alternative if GLM-Flash misbehaves on your task. |
| `laguna` | Long-form prose, summaries | Runs via Ollama; slower first hit (cold load). Wired to Haiku tier. |

---

## ⏱️ Speed & latency — what to expect

> **❗ Read this before you judge.** Every turn re-prefills the entire conversation. There is **no prompt cache** locally today. Plan for **1–3 minutes per turn**, not the sub-second feel of anthropic.com. Your first "Hi" can take 2+ minutes — that's the floor, not a bug.

Measured on studio1 (Apple Silicon, GLM-4.7-Flash 8-bit) with Claude Code's stock system prompt + ~60 tools:

| Turn | Typical latency | Why |
|---|---|---|
| First "Hi" in a fresh session | ≈ 90–150 s | Prefill of system prompt + tool definitions (~9K tokens) at ~90 tok/s. |
| Tool-using turn (read a file, suggest an edit) | ≈ 90–180 s | Same prefill plus the file you just attached + the prior turns. |
| Long session (10+ turns, large files in context) | Grows turn-over-turn | Conversation keeps re-prefilling. Use `/compact` and `/clear` proactively. |

### 🤔 Why so much slower than anthropic.com?
- **Prompt caching** — Anthropic caches your system prompt and tools server-side, so a repeat turn skips prefill entirely (sub-second first token). Our stack doesn't have this yet (the underlying flag is broken in our current mlx-lm version).
- **Hardware** — a Mac Studio is not a datacenter GPU. Cloud Claude runs on accelerator clusters with orders-of-magnitude more memory bandwidth.
- **Model size** — GLM-4.7-Flash is ~30 GB on disk; Opus / Sonnet are far larger and run on far bigger machines. Smaller model partly compensates for the slower hardware, but only partly.

### ✅ Good fit for
- Learning Claude Code's workflow without burning Anthropic credits.
- Single-file edits, code review, explanations of opaque code.
- Pre-PR self-review on small diffs.
- Anything where data shouldn't leave the LAN.

### 🚫 Less good fit for
- Tight iterative loops on large files (each turn pays full prefill).
- Multi-file refactors that need long agentic chains.
- Anything where you'd notice a 60-second wait every turn.

For those, switch back to real Claude — see "Switching back to real Claude" below.

---

## 1️⃣ Install Node.js
You need Node 18+ (LTS recommended).

### 🪟 Windows

> **🏢 Corp-managed laptop? Try Software Center first.** Open Software Center (Start menu → "Software Center"), search for **Claude**, click **Install**. If it's there, you're done — no admin prompt, no PATH wrangling. **Skip ahead to Step 3 (env vars)**. The rollout is still in progress though, so most laptops don't have it yet — if Claude isn't listed for you, use the winget steps below to install Node, then continue to Step 2.

If Software Center doesn't list Claude — install Node manually. The flags below work **without admin rights** and **on corp Wi-Fi**:
```powershell
winget install OpenJS.NodeJS.LTS --source winget --scope user
```

> **📘 Why those two flags** — `--source winget` pins the Microsoft *winget* source instead of the default *msstore*, which fails with `0x8a15005e` (server certificate did not match) on corp Wi-Fi because SSL inspection breaks the Microsoft Store source. `--scope user` installs Node into your profile only and modifies your user PATH; no admin elevation needed. Close and reopen your terminal afterwards so the new PATH takes effect.

### 🍎 macOS
```bash
brew install node
```

### 🐧 Linux (Ubuntu / Debian)
```bash
sudo apt install -y nodejs npm
```

Sanity check:
```bash
node -v && npm -v
```

---

## 2️⃣ Install Claude Code

> **🏢 Installed via Software Center on Windows?** You're done — SC installs Claude Code alongside Node. Skip to Step 3 (env vars).

Otherwise (Software Center doesn't list Claude yet, or you're on Mac/Linux), same one-liner everywhere:
```bash
npm install -g @anthropic-ai/claude-code
```
Verify:
```bash
claude --version
```

---

## 3️⃣ Configure five env vars
These point Claude Code at our internal endpoint and pick which local model each tier resolves to. Get the API key from the team lead.

### 🍎🐧 macOS / Linux — make it permanent
Open `~/.zshrc` (or `~/.bashrc` on bash) in your editor — pick whichever you have:
```bash
subl ~/.zshrc      # Sublime Text
code ~/.zshrc      # VS Code
nano ~/.zshrc      # nano (no install needed)
```
Append these five lines and save:
```bash
export ANTHROPIC_BASE_URL=__LLM_URL__
export ANTHROPIC_AUTH_TOKEN=<your-bearer-token>
export ANTHROPIC_DEFAULT_OPUS_MODEL=glm-4.7-flash
export ANTHROPIC_DEFAULT_SONNET_MODEL=glm-4.7-flash
export ANTHROPIC_DEFAULT_HAIKU_MODEL=laguna
```
Reload:
```bash
source ~/.zshrc
```

### 🪟 Windows — PowerShell, persistent (user-level)
```powershell
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_BASE_URL',            '__LLM_URL__', 'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_AUTH_TOKEN',           '<your-bearer-token>', 'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_OPUS_MODEL',   'glm-4.7-flash',          'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_SONNET_MODEL', 'glm-4.7-flash',          'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_HAIKU_MODEL',  'laguna',             'User')
```
Then close and reopen the terminal.

### 🩺 Verify — confirm the values are set

Open a **fresh** terminal (so it picks up the new vars), then run the line for your shell. All five values should print non-empty — if any are blank, the vars didn't persist and `claude` will fall back to api.anthropic.com.

🪟 **PowerShell**
```powershell
$env:ANTHROPIC_BASE_URL; $env:ANTHROPIC_AUTH_TOKEN.Substring(0,8) + "..."; $env:ANTHROPIC_DEFAULT_OPUS_MODEL; $env:ANTHROPIC_DEFAULT_SONNET_MODEL; $env:ANTHROPIC_DEFAULT_HAIKU_MODEL
```

🪟 **Windows CMD**
```cmd
echo %ANTHROPIC_BASE_URL% & echo %ANTHROPIC_AUTH_TOKEN:~0,8%... & echo %ANTHROPIC_DEFAULT_OPUS_MODEL% & echo %ANTHROPIC_DEFAULT_SONNET_MODEL% & echo %ANTHROPIC_DEFAULT_HAIKU_MODEL%
```

🍎🐧 **macOS / Linux**
```bash
echo "$ANTHROPIC_BASE_URL"; echo "${ANTHROPIC_AUTH_TOKEN:0:8}..."; echo "$ANTHROPIC_DEFAULT_OPUS_MODEL"; echo "$ANTHROPIC_DEFAULT_SONNET_MODEL"; echo "$ANTHROPIC_DEFAULT_HAIKU_MODEL"
```

The token is truncated to its first 8 characters so you can sanity-check it's set without echoing the full secret to your terminal scrollback.

---

## 4️⃣ Take it for a spin
```bash
cd ~/some/repo
claude
```
Inside the prompt, type `/status` — confirm `ANTHROPIC_BASE_URL` shows __LLM_URL__. Then say hi:
```
> hi, what model are you?
```

---

## 🎛️ Switch models on the fly

The five env vars in step 3 set your defaults. To try a different model for one session without editing your shell config, override at launch.

### ① Env-var prefix (one session)
```bash
ANTHROPIC_MODEL=qwen2.5-coder-32b claude
```
`ANTHROPIC_MODEL` takes precedence over the per-tier vars. Banner will read `qwen2.5-coder-32b[1m]` instead of the default. Closes the override when the session ends.

### ② CLI flag (one session)
```bash
claude --model qwen2.5-coder-32b
```
Same effect as the env-var prefix; pick whichever feels natural.

### ③ `/model` picker (mid-session)
Inside Claude Code, `/model` switches between the Opus / Sonnet / Haiku tiers. Each tier resolves to whichever id you set in `ANTHROPIC_DEFAULT_{OPUS,SONNET,HAIKU}_MODEL`. Useful if you wired Sonnet to `qwen2.5-coder-32b` — `/model` then becomes a GLM ↔ Qwen toggle without restarting.

> **💡 When to reach for each model.** Default to `glm-4.7-flash`. If a coding turn comes back empty, drops a tool call, or the model talks about code instead of writing it, retry with `qwen2.5-coder-32b` — it's code-tuned and doesn't have a thinking-mode prefix that can swallow short answers.

---

## 5️⃣ Your first real task — a 2-minute tutorial

### ① Open a project (or start fresh)
`cd` into an existing repo and run `claude` there — it'll see your code and edit files in place.

### ② Ask it to build something concrete
At the prompt, paste:
```
> add a simple is_palindrome(s) helper in utils.py with a unit test. Strip non-alphanumerics, ignore case.
```
Claude Code will plan the file, write it, show you the diff, and ask before saving. Press `y` to accept, `e` to edit inline, or describe what you want changed.

### ③ Iterate
```
> run the test
```
It'll create the test file, run it, and show output. If a test fails, ask it to fix and re-run.

### ④ Useful slash commands

| Command | What it does |
|---|---|
| `/status` | Show the resolved env vars, model, and working directory. |
| `/clear` | Reset the conversation in this session. |
| `/compact` | Compress earlier turns to free up context. |
| `/help` | Full command reference. |

### 💡 Tips for getting good results
- Be specific. "Refactor this function for clarity, keep the public signature" beats "make it better."
- Smaller scopes win. One file at a time. Long multi-file refactors are where local models struggle.
- Show, don't tell. Pasting a small example of the desired output usually beats describing it.
- Verify the diff. The model occasionally hallucinates an import or path — read before accepting.
- If it gets stuck, `/clear` and rephrase. Don't keep nudging a confused conversation.

---

## 🛡️ The permission model
Claude Code never silently edits files or runs commands. Every side-effect goes through a prompt. It will ask before:
- Writing to or deleting a file
- Running shell commands (especially anything destructive)
- Installing packages or making network calls

> **💡 Shift+Tab** cycles between three modes: ask (default), auto-accept edits, plan-only. Current mode shows in the bottom-of-terminal status line.

---

## ⌨️ Keyboard shortcuts cheat-sheet

| Shortcut | What it does |
|---|---|
| `Esc` | Cancel the current generation. |
| `Shift+Tab` | Cycle permission modes (ask / auto-accept / plan). |
| `Ctrl+R` | Toggle thinking mode. |
| `Ctrl+C` | Quit Claude Code. |
| `@filename` | Anchor a specific file as context. |
| `!cmd` | Drop into a shell, run cmd, pipe output back. |

---

## 📖 Recipe gallery
"Write a script" is the obvious one. Higher-value patterns most newcomers don't think to try:
- **Onboard yourself to an unfamiliar repo** — ask for a tour of `src/`, then have it explain the data model.
- **Diagnose a failing test** — paste the failure, ask for a hypothesis and a minimal repro.
- **Write tests for code you didn't write** — "add 5 unit tests covering edge cases of `parse_xyz`."
- **Pre-PR self-review** — "review the diff against main and flag anything that'd embarrass me."
- **Refactor with intent** — "extract the validation logic into a pure function, keep behavior identical."
- **Explain something opaque** — paste a regex, a SQL plan, a stack trace — ask for plain English.
- **Add docstrings to legacy code** — "add Google-style docstrings to all public methods in this file."
- **Translate between languages** — "port this Python function to Go, idiomatic."

---

## 📝 The CLAUDE.md trick — make it know your project
Drop a `CLAUDE.md` at the root of your repo and Claude Code reads it on every run. Teach it your conventions once instead of repeating yourself in every prompt.

Generate a starter automatically:
```
claude /init
```

`AGENTS.md` is the cross-tool version of the same idea — read by Claude Code, Cursor, Aider, Codex CLI, Gemini CLI. Same format, neutral name. Either works; teams that mix tools usually use AGENTS.md.

---

## 🆚 Prefer an IDE? Use the VS Code extension
Install the "Claude Code" extension from the VS Code marketplace. Same env-var config, same slash commands, same permission model — just rendered as a side panel in VS Code.
- Diffs render inline in the editor
- Selected code is auto-attached as context
- Same env vars work — no extra config

---

## 🔄 Switching back to real Claude
When you want real Opus / Sonnet for the heavy lifting, unset the five env vars:

### 🍎🐧 macOS / Linux
```bash
unset ANTHROPIC_BASE_URL ANTHROPIC_AUTH_TOKEN \
      ANTHROPIC_DEFAULT_OPUS_MODEL ANTHROPIC_DEFAULT_SONNET_MODEL ANTHROPIC_DEFAULT_HAIKU_MODEL
```

### 🪟 Windows PowerShell
```powershell
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_BASE_URL',            $null, 'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_AUTH_TOKEN',             $null, 'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_OPUS_MODEL',   $null, 'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_SONNET_MODEL', $null, 'User')
[System.Environment]::SetEnvironmentVariable('ANTHROPIC_DEFAULT_HAIKU_MODEL',  $null, 'User')
```
Then `claude login` to authenticate.

---

## ⚠️ Caveats — read this once

> **❗ Not Claude.** These are smaller open-weight models running on Mac hardware. Expect a quality drop from Opus / Sonnet, especially on long multi-file refactors and complex tool chains. Use them for what they're good at; reach for real Claude when the task warrants it.

- Smaller context window than Claude. Use `/compact` often, `/clear` when conversations drift.
- Tool-call reliability varies by model. If it loops or emits malformed JSON, simplify the prompt.
- Every turn re-prefills the conversation (no local prompt cache today) — see the Speed & latency section above.
- If something's genuinely broken (not just "lower quality than Claude"), file it.

---

## 🌐 Network gotchas
- Reachable from corp WiFi/wired LAN, or from home over the corp VPN.
- Not reachable from the corporate proxy-only or fully off-VPN networks.
- If `claude` hangs at startup, check that you can reach the gateway.

### 🩺 Quick reachability check
```bash
curl -H "Authorization: Bearer $ANTHROPIC_AUTH_TOKEN" $ANTHROPIC_BASE_URL/v1/models
```
Expected: a JSON list with `glm-4.7-flash` and `laguna`.

---

## ❓ FAQ — things people ping me about

**Q: 🪟 What's the easiest way to install on a corp-managed Windows laptop?**
A: Try **Software Center** first — search for **Claude**; if it's listed, click **Install** and you're done (no admin prompt, no PATH fiddling, jump to Step 3). The SC rollout is still in progress though, so most laptops don't have it yet — if Claude isn't there for you, fall back to the winget + npm steps in Step 1 and Step 2.

**Q: 🪟 winget install fails with "server certificate did not match" (`0x8a15005e`).**
A: If Claude is in your Software Center, that path skips winget entirely (search "Claude" → Install). Otherwise: corp SSL inspection breaks the Microsoft Store source. Pin winget explicitly: `winget install OpenJS.NodeJS.LTS --source winget`. The error message lists `winget` as a working source — that's the one to use.

**Q: 🪟 I don't have admin rights on my Windows laptop — can I still install?**
A: Yes. Easiest path is Software Center if Claude is listed for you (no admin needed). If it isn't, use winget with `--scope user`: `winget install OpenJS.NodeJS.LTS --source winget --scope user`. Node installs into your profile and only your user PATH is modified. Last-resort fallback is the portable zip from https://nodejs.org/dist/: extract to `%USERPROFILE%\nodejs`, then add that folder to your **user** PATH (System Properties → Environment Variables → User variables → Path → New).

**Q: It told me it can't access the internet. Is something broken?**
A: No — that's expected. The local models run fully offline by design. To fetch a webpage, run `!curl ...` and pipe the output back in.

**Q: My output got cut off mid-sentence.**
A: You hit the context window. `/compact` to compress earlier turns, or `/clear` to start fresh.

**Q: It hallucinated a function / import / file path.**
A: Known weakness of smaller models. Always read the diff before accepting. If it keeps inventing things, scope down — ask about one file at a time and use `@filename` to anchor it.

**Q: It refused to do something benign.**
A: Rephrase. Adding context like "this is my own project, the file is mine to edit" usually unblocks it.

**Q: Tool calls are failing or producing malformed JSON.**
A: Simplify the request — break it into smaller steps.

**Q: Every turn feels slow. Why isn't the second one faster?**
A: Anthropic's cloud caches your system prompt + tools server-side, so repeat turns skip prefill. We don't have that locally yet — every turn re-prefills the conversation. The first turn pays for ~9K tokens of system prompt + tools; each subsequent turn pays for that plus everything since. Use `/compact` and `/clear` to keep context lean. See the Speed & latency section near the top of this doc for measured numbers.

**Q: Can I use this for confidential / customer data?**
A: Prompts and responses stay on our hardware — nothing is sent to Anthropic or any third party. Follow normal data-handling policy.

**Q: Can I run two `claude` sessions at once?**
A: Yes. Open a second terminal in a different repo. Independent contexts.

**Q: How do I see what env vars Claude Code is using?**
A: Type `/status` inside the prompt — shows resolved BASE_URL, model, permission mode, working directory.

---

## 🆘 Where to get help
- Inside Claude Code: `/help` for the full command list.
- Setup or routing issues: ping the team lead.
- Bugs in the model itself: file with the team lead so we can patch the bridge or tweak the config.
- Questions about Claude Code: https://docs.anthropic.com/en/docs/claude-code — public docs apply, just substitute our endpoint.

---

## 📚 References
- Claude Code official docs: https://docs.anthropic.com/en/docs/claude-code
- VS Code extension: search "Claude Code" in the marketplace
- Internal admin guide: `docs/CLAUDE_CODE_ADMIN_SETUP.md`
"""

ADMIN_MD = """# 🛠️ Claude Code Local Stack — Admin Guide

> **scope: lab-vm1 + Mac fleet** · **consumers: any Claude Code client on corp net** · **status: production**

Operating manual for the router that lets Claude Code clients talk to our self-hosted vllm-mlx and Ollama backends.

| 🧠 2 backends | 🚪 1 endpoint | ⚙️ 2 services |
|---|---|---|
| studio1 GLM-Flash (vllm-mlx), studio1 Laguna (Ollama). | `lab-vm1:8051` — single URL, bearer-auth gated. | `ir-claude-router` (8050) + `ir-claude-router-shim` (8051). |

## 📑 Table of contents
- [Architecture](#-architecture)
- [Components & ports](#-components--ports)
- [Service operations](#️-service-operations)
- [Configuration files](#-configuration-files)
- [Adding a new model](#-adding-a-new-model)
- [Mac backends](#️-mac-backends)
- [Security & secrets](#-security--secrets)
- [Troubleshooting](#-troubleshooting)
- [Backup & recovery](#-backup--recovery)
- [References](#-references)

---

## 🧱 Architecture

Two services on lab-vm1, two Mac backends — both on studio1. The shim is the public face; ccr is the internal translator.

```
[claude client]
      │  ANTHROPIC_BASE_URL=http://lab-vm1:8051
      ▼
  lab-vm1:8051   ir-claude-router-shim   (FastAPI)
      │   • exposes GET /v1/models for SDK / curl discovery
      │   • rewrites friendly id → provider,model
      │   • bearer-auth gate
      ▼
  127.0.0.1:8050   ir-claude-router      (claude-code-router)
      │   • Anthropic /v1/messages → OpenAI /v1/chat/completions
      │   • routes by `provider,model` to one of two upstreams
      ▼
  ┌──────────────┬─────────────────────┐
  │ 8024         │ 8022                │
  │ studio1 GLM  │ studio1 Laguna      │
  │ vllm-mlx     │ Ollama              │
  └──────────────┴─────────────────────┘
  (each is a reverse SSH tunnel from studio1 into lab-vm1)
```

> **📌 Why two layers** — ccr handles the Anthropic↔OpenAI translation and multi-provider routing, but expects requests in `provider,model` form and doesn't expose `/v1/models` for discovery. The shim adds `/v1/models` (for SDK / curl / IDE-plugin enumeration), translates friendly model ids to ccr's `provider,model` form on incoming `/v1/messages`, and gates everything behind a bearer token. Note: Claude Code's `/model` picker is hardcoded to Opus / Sonnet / Haiku and does NOT read `/v1/models` — users wire each tier to one of our ids via `ANTHROPIC_DEFAULT_{OPUS,SONNET,HAIKU}_MODEL`. ~120 lines of FastAPI; no logic of its own beyond the rewrite.

---

## 📋 Components & ports

| Service | Port | Purpose | Source |
|---|---|---|---|
| `ir-claude-router` | 8050 | claude-code-router (npm). Anthropic↔OpenAI + provider routing. | `~/.claude-code-router/config.json` |
| `ir-claude-router-shim` | 8051 | FastAPI front door. `/v1/models` (discovery), id-rewrite, bearer auth. | `deployment/claude_router_shim.py` |

Backends (each lives behind a reverse SSH tunnel from its Mac):

| Tunnel port | Mac | Engine | Model |
|---|---|---|---|
| 8024 | studio1 | vllm-mlx | `mlx-community/GLM-4.7-Flash-8bit` |
| 8022 | studio1 | Ollama | `laguna-xs.2:q8_0` |

---

## ⚙️ Service operations

### 🩺 Status
```bash
systemctl --user status ir-claude-router ir-claude-router-shim
systemctl --user is-active ir-claude-router ir-claude-router-shim
```

### 🔄 Restart
```bash
# config.json edit → restart router
systemctl --user restart ir-claude-router

# claude_router_shim.py edit → restart shim
systemctl --user restart ir-claude-router-shim
```

### 📜 Logs

| Log | Path |
|---|---|
| Shim (Python uvicorn) | `data/transient/logs/claude_router_shim.log` |
| Router (ccr — systemd capture) | `data/transient/logs/claude_router.log` |
| Router (ccr — pino server) | `~/.claude-code-router/logs/ccr-*.log` |

---

## 🔧 Configuration files

### ① `~/.claude-code-router/config.json`
ccr config — providers, model lists, routing rules. Not under git (contains live URLs).

```json
{
  "HOST": "0.0.0.0",
  "PORT": 8050,
  "APIKEY": "$CCR_APIKEY",
  "Providers": [
    { "name": "glm",    "api_base_url": "http://127.0.0.1:8024/v1/chat/completions",
      "api_key": "sk-no-key",
      "models": ["glm-4.7-flash"] },
    { "name": "laguna", "api_base_url": "http://127.0.0.1:8022/v1/chat/completions",
      "api_key": "sk-no-key",
      "models": ["laguna-xs.2:q8_0"] }
  ],
  "Router": {
    "default":     "glm,glm-4.7-flash",
    "background":  "glm,glm-4.7-flash",
    "think":       "glm,glm-4.7-flash",
    "longContext": "glm,glm-4.7-flash"
  }
}
```

### ② `deployment/claude_router_shim.py`
Two dicts at the top decide what shows up in `/v1/models` and how aliases map to ccr providers. Edit, save, restart `ir-claude-router-shim`.

```python
MODEL_MAP = {
    "glm-4.7-flash":  "glm,glm-4.7-flash",
    "laguna":         "laguna,laguna-xs.2:q8_0",
}
DISPLAY_NAMES = {
    "glm-4.7-flash": "GLM 4.7 Flash",
    "laguna":        "Laguna xs.2",
}
```

### ③ `data/transient/.env` → `CCR_APIKEY`
Single bearer token validated by both the shim (own check) and ccr (forwarded). Clients send the same value as `Authorization: Bearer <CCR_APIKEY>` and as `ANTHROPIC_AUTH_TOKEN`.

> **⚠️ Keep the key out of git.** `.env` is gitignored. Distribute out-of-band (1Password, encrypted message). Rotate by regenerating, then restarting both services.

---

## ➕ Adding a new model
End-to-end walkthrough — example: `gemma3` on studio2.

### ① Stand up the model on the Mac
- Install vllm-mlx (or use the existing service plist as a template — see studio1 reference).
- Run with `--served-model-name gemma3` so the model id is human-readable; otherwise the full path is used.
- Bind to `127.0.0.1:<port>` on the Mac. Don't expose externally.

### ② Add a reverse SSH tunnel
On the Mac, edit `~/Library/LaunchAgents/com.ir.tunnel-to-labvm.plist` to add a new `-R <lab-vm-port>:127.0.0.1:<mac-port>` mapping. Pick a free lab-vm1 port (e.g. 8024). Reload with `launchctl kickstart -k gui/$(id -u)/com.ir.tunnel-to-labvm`.

Verify on lab-vm1:
```bash
curl -s http://127.0.0.1:8024/v1/models
```

### ③ Add provider to ccr config
Edit `~/.claude-code-router/config.json` — add a `Providers` entry:
```json
{ "name": "gemma",
  "api_base_url": "http://127.0.0.1:8024/v1/chat/completions",
  "api_key": "sk-no-key",
  "models": ["gemma3"] }
```

### ④ Add alias to the shim
Edit `deployment/claude_router_shim.py`:
```python
MODEL_MAP["claude-gemma3"]      = "gemma,gemma3"
DISPLAY_NAMES["claude-gemma3"]  = "Gemma 3"
```

### ⑤ Restart and verify
```bash
systemctl --user restart ir-claude-router ir-claude-router-shim

CCR_KEY=$(grep '^CCR_APIKEY=' data/transient/.env | cut -d= -f2)
curl -s -H "Authorization: Bearer $CCR_KEY" http://127.0.0.1:8051/v1/models | jq
```
Expected: the new `claude-gemma3` entry appears alongside the others.

---

## 🖥️ Mac backends

All three Claude Code backends now live on **studio1**, behind a single reverse-tunnel session to lab-vm1. (mac-m1 still runs GLM-4.7-Flash for Pokedex + Win.AI on its own tunnel, but is no longer in the Claude Code path as of 2026-05-06.)

### 🎙️ studio1 (GLM + Laguna, two stacks)
- vllm-mlx GLM: `mlx-community/GLM-4.7-Flash-8bit`, parser `glm47`, reasoning `deepseek_r1`, tunnel `lab-vm1:8024 → studio1:8002` (~30 GB on disk)
- Ollama: `laguna-xs.2:q8_0`, tunnel `lab-vm1:8022 → studio1:11434` (~40 GB cold, `KEEP_ALIVE=30s` so it unloads when idle)
- SSH backchannel: `ssh -p 2224 vvobbilichetty@127.0.0.1` from lab-vm1
- GLM reload: `launchctl bootout gui/$(id -u)/com.ir.vllm-mlx-glm && launchctl bootstrap gui/$(id -u) ~/Library/LaunchAgents/com.ir.vllm-mlx-glm.plist`
- Qwen3-32B vllm-mlx is downloaded (~33 GB) but the launchctl agent is **disabled** (2026-05-06) to avoid memory contention with GLM. Re-enable with `launchctl enable gui/$(id -u)/com.ir.vllm-mlx-qwen && launchctl bootstrap gui/$(id -u) ~/Library/LaunchAgents/com.ir.vllm-mlx-qwen.plist`.

> **⚠️ launchctl domain** — studio1's vllm-mlx services run in the `gui/$UID` domain, but the tunnel agent is in `user/$UID`. `launchctl print user/501` shows the vllm services as "enabled" but kickstart/bootout there fails with "Could not find service in domain." Use `gui/$(id -u)/...` for vllm; `user/$(id -u)/...` for the tunnel.

> **⚡ System-prompt KV cache patch** — vllm-mlx's `--continuous-batching` flag (which gates the engine-level prefix cache) crashes mlx-lm on first cache-hit decode with `RuntimeError: There is no Stream(gpu, X) in current thread`. We work around it with a local patch to `vllm_mlx/engine/simple.py` that adds single-slot system-prompt KV caching to the pure-LLM `stream_chat()` path. Result: ~17x speedup on cache hits (9.8s cold → 0.58s hit on Qwen2.5-Coder with a 2.5K-token system prompt). Patch + idempotent apply script: `deployment/vllm_mlx_patches/`. Re-run `apply.sh` after every `pip install --upgrade vllm-mlx` and bounce the launchctl agent.

---

## 🔐 Security & secrets
- Single bearer (`CCR_APIKEY`) validates both the shim and ccr.
- Shim binds `0.0.0.0` — accessible from any host on the corp network. No external exposure.
- ccr is bound `0.0.0.0` too but should only be hit via the shim. Future hardening: bind ccr to `127.0.0.1`.
- All traffic between lab-vm1 and the Macs is over reverse SSH tunnels (encrypted + key-auth).
- vllm-mlx and Ollama on the Macs bind `127.0.0.1` only — not reachable except through the tunnel.

### 🔁 Rotating the bearer token
```bash
# 1. regenerate
openssl rand -hex 32

# 2. update data/transient/.env
sed -i "s|^CCR_APIKEY=.*|CCR_APIKEY=<new-value>|" data/transient/.env

# 3. restart
systemctl --user restart ir-claude-router ir-claude-router-shim

# 4. distribute new value to users
```

---

## 🚨 Troubleshooting

| Symptom | Likely cause | Fix |
|---|---|---|
| 401 on `/v1/models` | Wrong/missing bearer token | Check `ANTHROPIC_AUTH_TOKEN` matches `CCR_APIKEY` in `data/transient/.env`. |
| 404 on `/v1/models` | Client pointing at ccr (8050) instead of shim (8051) | Set `ANTHROPIC_BASE_URL=http://lab-vm1:8051`. |
| `/v1/messages` "fetch failed" | Upstream Mac unreachable | Check tunnel: `ss -tlnp \\| grep 80<port>`. SSH the Mac, verify the engine is up. |
| Connection reset on GLM path | studio1 `vllm-mlx-glm` not running | ssh studio1, `launchctl bootstrap gui/$(id -u) ~/Library/LaunchAgents/com.ir.vllm-mlx-glm.plist`. |
| Picker doesn't show models | Stale gateway-models cache on client | Delete `~/.claude/cache/gateway-models.json` on client, restart `claude`. |
| Tool calls flaky | Model-specific (smaller models drop tool args) | Switch to `glm-4.7-flash` — most reliable for tool use. |
| Shim restart with no effect | Edited `config.json` but didn't restart ccr | Restart both for any provider change. |

---

## 💾 Backup & recovery
All four artifacts are tiny and rsync'd to lab-vm2 by the weekly backup cron:
- `~/.claude-code-router/config.json` — provider routing
- `deployment/claude_router_shim.py` — alias map (under git)
- `deployment/systemd/ir-claude-router*.service` — units (under git)
- `data/transient/.env` — `CCR_APIKEY` line

Recovery: restore those four files, `systemctl --user daemon-reload`, then start both services. No DB, no schema migrations.

---

## 📚 References
- User-facing setup guide: `docs/CLAUDE_CODE_USER_SETUP.docx`
- claude-code-router repo: https://github.com/musistudio/claude-code-router
- Memory: `~/.claude/projects/-home-vinay-IR/memory/project_claude_code_router.md`
- Memory: `~/.claude/projects/-home-vinay-IR/memory/project_studio1_qwen3_vllm.md`
- Generator script: `misc_scripts/build_claude_code_docs.py`
"""


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    DOCS.mkdir(exist_ok=True)

    user_doc = build_user_doc()
    user_doc.save(DOCS / "CLAUDE_CODE_USER_SETUP.docx")
    (DOCS / "CLAUDE_CODE_USER_SETUP.md").write_text(USER_MD.replace("__LLM_URL__", LLM_PUBLIC_URL))

    admin_doc = build_admin_doc()
    admin_doc.save(DOCS / "CLAUDE_CODE_ADMIN_SETUP.docx")
    (DOCS / "CLAUDE_CODE_ADMIN_SETUP.md").write_text(ADMIN_MD)

    print(f"Wrote {DOCS / 'CLAUDE_CODE_USER_SETUP.docx'}")
    print(f"Wrote {DOCS / 'CLAUDE_CODE_USER_SETUP.md'}")
    print(f"Wrote {DOCS / 'CLAUDE_CODE_ADMIN_SETUP.docx'}")
    print(f"Wrote {DOCS / 'CLAUDE_CODE_ADMIN_SETUP.md'}")


if __name__ == "__main__":
    main()
