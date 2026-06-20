"""Chat routes: Sleuth AI chat, Aide chat, page-context chat widget."""

import datetime as _dt
import json
import logging
import time

import requests
from flask import Blueprint, Response, jsonify, render_template, request, session, current_app

from src.utils.logging_utils import log_web_activity, get_client_ip
from src.components.web import sleuth_handler, aide_handler, approved_testing_handler
from src.components.web import page_chat_handler
from src.utils import bot_logs_db
from web.auth import helpers as auth_helpers
from web.config import CONFIG, EASTERN, prod_list_handler, prod_ticket_handler
from web.extensions import limiter

# Public-playground backend registry: short key → (base-url config attr,
# display label, api-key config attr or None). The served model id is
# fetched live from `<base_url>/models` so renaming a model on a server
# doesn't need a code change. Endpoints accept a 'model' key from clients;
# resolution happens server-side so a client can never point us at an
# arbitrary base_url. The dropdown options on the playground page are
# also built from this registry via /api/sleuth-models.
#
# studio1 vllm-mlx instances (8003, 8004) share a single --api-key; we
# reuse `embeds_api_key` for the Qwen entry rather than carrying a
# duplicate setting.
_PUBLIC_BACKENDS = {
    "glm":  ("m1_analysis_base_url",  "GLM-4.7 Flash (M1)",        None),
    "qwen": ("studio1_qwen_base_url", "Qwen3-Coder 30B (Studio1)", "embeds_api_key"),
}
_DEFAULT_PUBLIC_MODEL = "glm"
_PUBLIC_TEMP_MIN = 0.0
_PUBLIC_TEMP_MAX = 1.5
_PUBLIC_TEMP_DEFAULT = 0.3

_MODEL_ID_CACHE: dict[str, tuple[float, str]] = {}
_MODEL_ID_TTL_S = 300  # 5 min — survives a model swap quickly enough


def _fetch_served_model_id(base_url: str, api_key: str | None = None) -> str | None:
    cache_key = (base_url, bool(api_key))
    cached = _MODEL_ID_CACHE.get(cache_key)
    if cached and (time.time() - cached[0]) < _MODEL_ID_TTL_S:
        return cached[1]
    try:
        headers = {"Authorization": f"Bearer {api_key}"} if api_key else None
        r = requests.get(f"{base_url.rstrip('/')}/models", timeout=10, headers=headers)
        r.raise_for_status()
        data = r.json().get("data") or []
        model_id = (data[0] or {}).get("id") if data else None
        if model_id:
            _MODEL_ID_CACHE[cache_key] = (time.time(), model_id)
            return model_id
    except Exception as exc:
        logging.getLogger(__name__).warning("served-model lookup failed for %s: %s", base_url, exc)
    return None


def _resolve_public_model(model_key: str) -> tuple[str, str, str, str | None] | None:
    """Map a client-supplied model key to (key, base_url, model_id, api_key).
    Returns None if the key is unknown, its base_url isn't configured, or
    the backend's /v1/models doesn't respond."""
    entry = _PUBLIC_BACKENDS.get(model_key)
    if not entry:
        return None
    cfg_attr, _label, api_key_attr = entry
    base_url = getattr(CONFIG, cfg_attr, None)
    if not base_url:
        return None
    api_key = getattr(CONFIG, api_key_attr, None) if api_key_attr else None
    model_id = _fetch_served_model_id(base_url, api_key)
    if not model_id:
        return None
    return model_key, base_url, model_id, api_key


def _clamp_temperature(raw) -> float:
    try:
        t = float(raw)
    except (TypeError, ValueError):
        return _PUBLIC_TEMP_DEFAULT
    if t < _PUBLIC_TEMP_MIN:
        return _PUBLIC_TEMP_MIN
    if t > _PUBLIC_TEMP_MAX:
        return _PUBLIC_TEMP_MAX
    return t

logger = logging.getLogger(__name__)
chat_bp = Blueprint('chat', __name__)

# Lazy-init a lightweight LLM for the Defense Pulse chat widget
_dp_llm = None


def _get_dp_llm():
    global _dp_llm
    if _dp_llm is None:
        from my_bot.utils.llm_factory import create_llm
        _dp_llm = create_llm(temperature=0.1)
    return _dp_llm

# Lazy imports for Sleuth components
SLEUTH_AVAILABLE = True

try:
    from my_bot.core.my_model import ask, ask_stream
    from my_bot.core.state_manager import get_state_manager
except Exception as e:
    logger.warning(f"Sleuth components unavailable: {e}")
    SLEUTH_AVAILABLE = False

    def ask(*_args, **_kwargs):
        return "Model not available in this environment"

    def ask_stream(*_args, **_kwargs):
        yield "Model not available in this environment"

    def get_state_manager():
        return None


# --- Sleuth Chat ---

@chat_bp.route('/sleuth')
@auth_helpers.admin_required
@log_web_activity
def sleuth_chat():
    """Sleuth AI chat interface"""
    return render_template('sleuth_chat.html')


@chat_bp.route('/sleuth-demo')
@auth_helpers.admin_required
@log_web_activity
def sleuth_demo():
    """Guided demo page showcasing Sleuth capabilities."""
    return render_template('sleuth_demo.html')


# Cache fixtures for 5 min so repeat clicks don't hammer the APIs.
_PKX_DEMO_FIXTURES_CACHE: dict = {'ts': 0.0, 'payload': None}
_PKX_DEMO_FIXTURES_TTL_S = 300


def _refresh_sleuth_demo_fixtures() -> dict:
    """Pull a recent XSOAR ticket (with hostname) and a recent ServiceNow INC.

    Returns a dict with `fixtures` (the values), `sources` (where each came from),
    and per-source error notes when something fails. Always returns a usable
    payload — falls back to the static placeholders if APIs are unreachable.
    """
    import time as _t
    import sqlite3 as _sql
    from pathlib import Path as _Path

    fixtures = {'xsoar_ticket': '932447', 'host': 'LTPX-47291', 'sn_inc': 'INC0123456'}
    sources: dict = {'xsoar_ticket': 'fallback (placeholder)', 'host': 'fallback (placeholder)', 'sn_inc': 'fallback (placeholder)'}
    errors: dict = {}

    # XSOAR ticket + hostname from the local timeline DB (closed tickets, has hostname)
    try:
        # Resolve relative to this worktree's root (web/routes/chat.py -> repo
        # root) so the dev instance reads its own data/, not prod's.
        db = _Path(__file__).resolve().parents[2] / 'data' / 'xsoar_timeline' / 'xsoar_timeline.db'
        with _sql.connect(str(db)) as conn:
            row = conn.execute(
                """
                SELECT id, hostname, closed_date, security_category
                FROM xsoar_tickets
                WHERE status = 2
                  AND hostname IS NOT NULL AND hostname != ''
                  AND length(hostname) > 5
                  AND security_category IS NOT NULL
                ORDER BY closed_date DESC
                LIMIT 1
                """
            ).fetchone()
        if row:
            fixtures['xsoar_ticket'] = str(row[0])
            fixtures['host'] = str(row[1])
            sources['xsoar_ticket'] = f"xsoar_timeline.db · closed {row[2]} · {row[3]}"
            sources['host'] = f"hostname from XSOAR #{row[0]}"
    except Exception as e:
        errors['xsoar'] = str(e)

    # ServiceNow — recent MIM group incident (last 7d), pick the latest
    try:
        from services.service_now import ServiceNowClient
        client = ServiceNowClient()
        latest = None
        for group in ('GTO-Major Incident management-US', 'GTO-Major Incident management-EMEA'):
            try:
                rows = client.get_recent_incidents_by_group_name(group, minutes=10080)
                if isinstance(rows, list):
                    for r in rows:
                        num = (r.get('number') or '').strip()
                        if num.startswith('INC'):
                            opened = r.get('opened_at') or ''
                            if latest is None or opened > latest[1]:
                                latest = (num, opened, group)
            except Exception as inner:
                errors.setdefault('servicenow_groups', []).append(f'{group}: {inner}')
        if latest:
            fixtures['sn_inc'] = latest[0]
            sources['sn_inc'] = f"ServiceNow · {latest[2]} · opened {latest[1]}"
    except Exception as e:
        errors['servicenow'] = str(e)

    return {
        'ok': True,
        'fixtures': fixtures,
        'sources': sources,
        'errors': errors,
        'refreshed_at': _t.strftime('%Y-%m-%dT%H:%M:%S'),
    }


@chat_bp.route('/api/sleuth-demo/fixtures')
@auth_helpers.admin_required
@log_web_activity
def api_sleuth_demo_fixtures():
    """Return fresh placeholder values for the demo script (host, XSOAR ticket, SN INC).

    Cached server-side for ~5 minutes; pass ?force=1 to bypass.
    """
    import time as _t
    force = request.args.get('force') in ('1', 'true', 'yes')
    cached = _PKX_DEMO_FIXTURES_CACHE.get('payload')
    ts = _PKX_DEMO_FIXTURES_CACHE.get('ts', 0.0)
    if cached and not force and (_t.time() - ts) < _PKX_DEMO_FIXTURES_TTL_S:
        return jsonify({**cached, 'cached': True, 'age_s': round(_t.time() - ts, 1)})
    payload = _refresh_sleuth_demo_fixtures()
    _PKX_DEMO_FIXTURES_CACHE['ts'] = _t.time()
    _PKX_DEMO_FIXTURES_CACHE['payload'] = payload
    return jsonify({**payload, 'cached': False, 'age_s': 0.0})


@chat_bp.route('/api/sleuth-demo/infra-check')
@auth_helpers.admin_required
@log_web_activity
def api_sleuth_demo_infra_check():
    """Pre-demo readiness probe. Verifies Sleuth bot + critical LLM endpoints + demo fixtures.

    Calls into the local bot_status_api (port 8040) for service/LLM state and runs the demo
    fixture refresh inline so the presenter knows whether XSOAR + ServiceNow IDs are live.
    """
    import requests as _rq
    import os as _os
    import time as _time
    from concurrent.futures import ThreadPoolExecutor as _Pool
    BOT_STATUS_BASE = 'http://localhost:8040'
    checks: list[dict] = []

    # 1+2. Sleuth bot + Mac host health in parallel (both are bot_status_api calls).
    # Mac host status is the prerequisite for the m1/studio1 LLM probes — if the
    # host is down, the TCP tunnel binding on lab-vm is stale and a probe will
    # hang for the full timeout. Short-circuit instead.
    def _get_bot():
        try:
            return _rq.get(f'{BOT_STATUS_BASE}/api/status/sleuth', timeout=3)
        except Exception as e:
            return e

    def _get_macs():
        try:
            return _rq.get(f'{BOT_STATUS_BASE}/api/mac-health', timeout=3)
        except Exception as e:
            return e

    with _Pool(max_workers=2) as pool:
        bot_f = pool.submit(_get_bot)
        macs_f = pool.submit(_get_macs)
        bot_resp = bot_f.result()
        macs_resp = macs_f.result()

    # Sleuth bot row
    if isinstance(bot_resp, Exception):
        checks.append({'key': 'sleuth-bot', 'label': 'Sleuth bot (Webex)', 'ok': False,
                       'detail': f'bot-status-api unreachable: {str(bot_resp)[:80]}'})
    elif bot_resp.ok:
        d = bot_resp.json().get('bot', {}) or {}
        status = d.get('status', 'unknown')
        ok = (status == 'running')
        extras = []
        if d.get('cpu_percent') is not None:
            extras.append(f"CPU {d['cpu_percent']}%")
        if d.get('memory_mb') is not None:
            extras.append(f"Mem {int(d['memory_mb'])}MB")
        detail = status + ((' · ' + ' · '.join(extras)) if (ok and extras) else '')
        checks.append({'key': 'sleuth-bot', 'label': 'Sleuth bot (Webex)', 'ok': ok, 'detail': detail})
    else:
        checks.append({'key': 'sleuth-bot', 'label': 'Sleuth bot (Webex)', 'ok': False,
                       'detail': f'bot-status-api HTTP {bot_resp.status_code}'})

    # Parse mac host status
    mac_hosts: dict = {}
    if not isinstance(macs_resp, Exception) and macs_resp.ok:
        mac_hosts = (macs_resp.json() or {}).get('hosts', {}) or {}

    def _host_state(host_key: str) -> tuple[bool, str]:
        """Return (is_up, error_string_if_down)."""
        h = mac_hosts.get(host_key) or {}
        up = h.get('status') == 'up'
        err = (h.get('error') or 'host status unknown')
        return up, err

    # Hosts summary row — surfaces the underlying infra at-a-glance
    if mac_hosts:
        host_bits = []
        all_hosts_ok = True
        for k in ('mac-m1', 'studio1'):  # only the hosts demo prompts depend on
            up, _ = _host_state(k)
            host_bits.append(f"{k} {'✓' if up else '✗'}")
            if not up:
                all_hosts_ok = False
        checks.append({
            'key': 'mac-hosts', 'label': 'Mac hosts (SSH reachability)',
            'ok': all_hosts_ok, 'detail': ' · '.join(host_bits),
        })
    else:
        err_str = str(macs_resp) if isinstance(macs_resp, Exception) else f'HTTP {macs_resp.status_code}'
        checks.append({'key': 'mac-hosts', 'label': 'Mac hosts (SSH reachability)',
                       'ok': False, 'detail': f'mac-health unavailable: {err_str[:80]}'})

    # 3-5. Critical LLM endpoints. Skip the actual HTTP probe when the host is
    # known-down (stale TCP listener would hang for the full timeout).
    embeds_token = _os.environ.get('EMBEDS_API_KEY', '')
    m1_up, m1_err = _host_state('mac-m1')
    studio1_up, studio1_err = _host_state('studio1')
    llm_targets = [
        {'key': 'm1-analysis', 'label': 'Agent LLM — M1 Analysis (tool-calling)',
         'url': 'http://localhost:8015/v1/models', 'host_key': 'mac-m1',
         'host_up': m1_up, 'host_err': m1_err},
        {'key': 'm1-router', 'label': 'Router LLM — M1 (tool selection)',
         'url': 'http://localhost:8016/v1/models', 'host_key': 'mac-m1',
         'host_up': m1_up, 'host_err': m1_err},
        {'key': 'embed', 'label': 'Embeddings — S1 (RAG / Act 03)',
         'url': 'http://studio1.lab:8004/v1/models', 'host_key': 'studio1',
         'host_up': studio1_up, 'host_err': studio1_err,
         'headers': {'Authorization': f'Bearer {embeds_token}'} if embeds_token else {}},
    ]

    def _probe_llm(t):
        if not t['host_up']:
            return {'key': t['key'], 'label': t['label'], 'ok': False,
                    'detail': f"host {t['host_key']} down — {t['host_err'][:60]}"}
        start = _time.time()
        try:
            resp = _rq.get(t['url'], headers=t.get('headers', {}), timeout=5)
            lat = int((_time.time() - start) * 1000)
            if resp.ok:
                try:
                    models = [m.get('id', '').split('/')[-1] for m in (resp.json() or {}).get('data', [])]
                    model_str = models[0] if models else '—'
                except Exception:
                    model_str = '—'
                return {'key': t['key'], 'label': t['label'], 'ok': True,
                        'detail': f"{model_str} · {lat}ms"}
            return {'key': t['key'], 'label': t['label'], 'ok': False,
                    'detail': f"HTTP {resp.status_code} · {lat}ms"}
        except Exception as e:
            return {'key': t['key'], 'label': t['label'], 'ok': False,
                    'detail': f"unreachable: {str(e)[:80]}"}

    with _Pool(max_workers=3) as pool:
        for result in pool.map(_probe_llm, llm_targets):
            checks.append(result)

    # 5. Demo fixtures (XSOAR + Host + ServiceNow) — checks that placeholder swap will work live.
    # Reuse the fixtures cache (5 min TTL) when fresh; only refresh if stale or empty.
    try:
        import time as __t
        cached = _PKX_DEMO_FIXTURES_CACHE.get('payload')
        ts = _PKX_DEMO_FIXTURES_CACHE.get('ts', 0.0)
        if cached and (__t.time() - ts) < _PKX_DEMO_FIXTURES_TTL_S:
            payload = cached
        else:
            payload = _refresh_sleuth_demo_fixtures()
            _PKX_DEMO_FIXTURES_CACHE['ts'] = __t.time()
            _PKX_DEMO_FIXTURES_CACHE['payload'] = payload
        fixtures = payload.get('fixtures', {}) or {}
        sources = payload.get('sources', {}) or {}
        xsoar_live = 'fallback' not in (sources.get('xsoar_ticket') or '')
        host_live = 'fallback' not in (sources.get('host') or '')
        sn_live = 'fallback' not in (sources.get('sn_inc') or '')
        ok = xsoar_live and host_live and sn_live
        bits = [
            f"XSOAR {'✓' if xsoar_live else '✗'} #{fixtures.get('xsoar_ticket', '?')}",
            f"Host {'✓' if host_live else '✗'} {fixtures.get('host', '?')}",
            f"SN {'✓' if sn_live else '✗'} {fixtures.get('sn_inc', '?')}",
        ]
        checks.append({'key': 'fixtures', 'label': 'Demo fixtures (XSOAR · Host · SN)',
                       'ok': ok, 'detail': ' · '.join(bits)})
    except Exception as e:
        checks.append({'key': 'fixtures', 'label': 'Demo fixtures', 'ok': False,
                       'detail': f'refresh failed: {str(e)[:80]}'})

    ok_count = sum(1 for c in checks if c['ok'])
    fail_count = len(checks) - ok_count
    return jsonify({
        'checks': checks,
        'summary': {'ok': ok_count, 'fail': fail_count, 'total': len(checks)},
        'all_ok': fail_count == 0,
    })


@chat_bp.route('/api/sleuth-status')
@auth_helpers.admin_required
@limiter.limit("30 per minute")
@log_web_activity
def api_sleuth_status():
    """Health check endpoint for Pokédex chat availability"""
    status = sleuth_handler.check_sleuth_status(get_state_manager)
    return jsonify(status)


@chat_bp.route('/api/sleuth-models')
@auth_helpers.admin_required
@limiter.limit("30 per minute")
def api_sleuth_models():
    """Return the configured backends for the playground model dropdown.
    Backends without a base_url in config are filtered out."""
    models = [
        {"key": key, "label": label}
        for key, (cfg_attr, label, _api_key_attr) in _PUBLIC_BACKENDS.items()
        if getattr(CONFIG, cfg_attr, None)
    ]
    return jsonify({"models": models, "default": _DEFAULT_PUBLIC_MODEL})


@chat_bp.route('/api/sleuth-chat', methods=['POST'])
@auth_helpers.admin_required
@limiter.limit("10 per minute")
@log_web_activity
def api_sleuth_chat():
    """API endpoint for Sleuth chat messages"""
    try:
        data = request.get_json()
        user_message = data.get('message', '').strip()
        session_id = data.get('session_id', '')

        if not user_message:
            return jsonify({'success': False, 'error': 'Message is required'}), 400

        if len(user_message) > 4000:
            return jsonify({'success': False, 'error': 'Message too long (max 4,000 characters)'}), 400

        if not session_id:
            return jsonify({'success': False, 'error': 'Session ID is required'}), 400

        response_text = sleuth_handler.handle_sleuth_chat(
            user_message,
            session_id,
            get_client_ip(),
            ask
        )

        return jsonify({'success': True, 'response': response_text})

    except Exception as exc:
        logger.error(f"Error in Sleuth chat API: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'Failed to get response from AI. Please try again.'}), 500


def _sleuth_rate_limit_exempt() -> bool:
    """Skip the 5/min rate limit for the recording orchestrator.

    Gated on two conditions that together can only be met by something
    running on the same host: an authenticated admin session AND a
    request that originated from 127.0.0.1. External clients reach
    Flask through nginx + ProxyFix, which rewrites remote_addr to the
    real client IP, so they never appear as 127.0.0.1 here.
    """
    if request.remote_addr not in ('127.0.0.1', '::1'):
        return False
    return bool(auth_helpers.is_admin())


@chat_bp.route('/api/sleuth-chat-stream', methods=['POST'])
@auth_helpers.admin_required
@limiter.limit("5 per minute; 100 per day", exempt_when=_sleuth_rate_limit_exempt)
@log_web_activity
def api_sleuth_chat_stream():
    """Public-playground streaming endpoint.

    Reads `message`, `session_id`, optional `model` ('glm'|'qwen'),
    optional `temperature` ([0.0, 1.5]) from JSON body. Routes the turn
    through state_manager.execute_query_stream_public so only allowlisted
    read tools can run, and writes one row to public_chat_log on completion.
    """
    try:
        data = request.get_json() or {}
        user_message = (data.get('message') or '').strip()
        session_id = (data.get('session_id') or '').strip()
        model_key_raw = (data.get('model') or _DEFAULT_PUBLIC_MODEL).strip()
        temperature = _clamp_temperature(data.get('temperature'))

        if not user_message:
            return jsonify({'success': False, 'error': 'Message is required'}), 400
        if len(user_message) > 4000:
            return jsonify({'success': False, 'error': 'Message too long (max 4,000 characters)'}), 400
        if not session_id:
            return jsonify({'success': False, 'error': 'Session ID is required'}), 400

        resolved = _resolve_public_model(model_key_raw)
        if not resolved:
            return jsonify({
                'success': False,
                'error': f"Unknown or unconfigured model: {model_key_raw!r}. "
                         f"Choose one of: {sorted(_PUBLIC_BACKENDS.keys())}"
            }), 400
        model_key, base_url, model_id, api_key = resolved

        user_ip = get_client_ip()
        state_manager = get_state_manager()
        if state_manager is None:
            return jsonify({'success': False, 'error': 'Sleuth backend unavailable'}), 503

        llm_overrides = {
            "base_url": base_url,
            "model": model_id,
            "temperature": temperature,
        }
        if api_key:
            llm_overrides["api_key"] = api_key
        allowlist = state_manager.PUBLIC_TOOL_ALLOWLIST

        logger.info(
            f"Sleuth public stream: ip={user_ip} model={model_key} "
            f"temp={temperature} chars={len(user_message)}"
        )

        def generate():
            start_time = time.time()
            first_token_time = None
            stream_metrics = None
            response_chunks: list[str] = []
            try:
                for token in state_manager.execute_query_stream_public(
                    user_message,
                    llm_overrides=llm_overrides,
                    name_allowlist=allowlist,
                ):
                    if isinstance(token, dict) and token.get('_metrics'):
                        stream_metrics = token
                        continue

                    if first_token_time is None:
                        first_token_time = time.time()

                    response_chunks.append(token)
                    yield f"data: {json.dumps({'token': token})}\n\n"

                elapsed = round(time.time() - start_time, 1)
                ttft = round(first_token_time - start_time, 1) if first_token_time else None

                done_payload = {'done': True}
                if stream_metrics:
                    done_payload['metrics'] = {
                        'time': elapsed,
                        'eval_time': stream_metrics.get('eval_time'),
                        'gen_time': stream_metrics.get('gen_time'),
                        'input_tokens': stream_metrics.get('input_tokens'),
                        'output_tokens': stream_metrics.get('output_tokens'),
                        'speed': stream_metrics.get('speed'),
                        'iterations': stream_metrics.get('iterations'),
                        'route': stream_metrics.get('route'),
                        'ttft': ttft,
                        'model': model_key,
                        'temperature': temperature,
                    }

                yield f"data: {json.dumps(done_payload)}\n\n"

            except Exception as stream_err:
                logger.error(f"Error in public stream: {stream_err}", exc_info=True)
                yield f"data: {json.dumps({'error': 'Streaming error occurred'})}\n\n"
            finally:
                # Always log the turn — success or error path
                try:
                    bot_logs_db.log_public_chat(
                        timestamp=_dt.datetime.now(EASTERN).strftime('%Y-%m-%d %H:%M:%S'),
                        ip=user_ip,
                        model=model_key,
                        temperature=temperature,
                        prompt=user_message,
                        response=''.join(response_chunks),
                        tool_calls_json=json.dumps(
                            (stream_metrics or {}).get('tools_used') or []
                        ),
                        elapsed_s=round(time.time() - start_time, 2),
                    )
                except Exception as log_err:
                    logger.warning(f"public_chat_log insert failed: {log_err}")

        return current_app.response_class(
            generate(),
            mimetype='text/event-stream',
            headers={
                'Cache-Control': 'no-cache',
                'X-Accel-Buffering': 'no'
            }
        )

    except Exception as exc:
        logger.error(f"Error in Sleuth streaming chat API: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'An unexpected error occurred. Please try again.'}), 500


# --- Aide Chat ---

@chat_bp.route('/aide')
@auth_helpers.admin_required
@log_web_activity
def aide_chat():
    """Aide chat interface — admin-gated."""
    return render_template('aide_chat.html')


@chat_bp.route('/api/aide/login', methods=['POST'])
@auth_helpers.admin_required
@limiter.limit("5 per minute")
@log_web_activity
def api_aide_login():
    """API endpoint for Aide authentication"""
    try:
        data = request.get_json()
        password = data.get('password', '').strip()
        email = data.get('email', '').strip()

        success, error = aide_handler.authenticate_aide(password, CONFIG.aide_password)

        if success:
            session['aide_authenticated'] = True
            session['aide_user_email'] = email
            session.permanent = True
            return jsonify({'success': True, 'message': 'Authentication successful'})
        else:
            return jsonify({'success': False, 'error': error}), 401

    except Exception as exc:
        logger.error(f"Error in Aide login: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred'}), 500


@chat_bp.route('/api/aide/logout', methods=['POST'])
@auth_helpers.admin_required
@log_web_activity
def api_aide_logout():
    """API endpoint to logout from Aide"""
    session.pop('aide_authenticated', None)
    return jsonify({'success': True, 'message': 'Logged out successfully'})


@chat_bp.route('/api/aide/create-x-ticket', methods=['POST'])
@auth_helpers.admin_required
@log_web_activity
def api_create_x_ticket():
    """API endpoint to create X ticket"""
    try:
        data = request.get_json()
        title = data.get('title', '').strip()
        details = data.get('details', '').strip()
        detection_source = data.get('detection_source', '').strip()
        user_email = data.get('user_email', '').strip()

        if not title or not details or not detection_source:
            return jsonify({'success': False, 'error': 'All fields are required'}), 400

        message = aide_handler.create_x_ticket(
            title,
            details,
            detection_source,
            user_email,
            get_client_ip(),
            prod_ticket_handler,
            CONFIG.xsoar_prod_ui_base_url
        )

        return jsonify({'success': True, 'message': message})

    except Exception as exc:
        logger.error(f"Error creating X ticket: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred'}), 500


@chat_bp.route('/api/aide/approved-testing', methods=['POST'])
@auth_helpers.admin_required
@log_web_activity
def api_approved_testing():
    """API endpoint to add approved testing entry"""
    try:
        data = request.get_json()

        try:
            message = approved_testing_handler.submit_aide_approved_testing(
                data,
                prod_list_handler,
                CONFIG.team_name,
                EASTERN,
                get_client_ip()
            )
            return jsonify({'success': True, 'message': message})

        except ValueError as val_err:
            return jsonify({'success': False, 'error': str(val_err)}), 400

    except Exception as exc:
        logger.error(f"Error adding approved testing: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred'}), 500


@chat_bp.route('/api/aide/ioc-hunt', methods=['POST'])
@auth_helpers.admin_required
@log_web_activity
def api_ioc_hunt():
    """API endpoint to create IOC hunt"""
    try:
        data = request.get_json()
        ioc_title = data.get('ioc_title', '').strip()
        iocs = data.get('iocs', '').strip()
        user_email = data.get('user_email', '').strip()

        if not ioc_title or not iocs:
            return jsonify({'success': False, 'error': 'All fields are required'}), 400

        message = aide_handler.create_ioc_hunt(
            ioc_title,
            iocs,
            user_email,
            get_client_ip(),
            prod_ticket_handler,
            CONFIG.xsoar_prod_ui_base_url
        )

        return jsonify({'success': True, 'message': message})

    except Exception as exc:
        logger.error(f"Error creating IOC hunt: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred'}), 500


@chat_bp.route('/api/aide/threat-hunt', methods=['POST'])
@auth_helpers.admin_required
@log_web_activity
def api_threat_hunt():
    """API endpoint to create threat hunt"""
    try:
        data = request.get_json()
        threat_title = data.get('threat_title', '').strip()
        threat_description = data.get('threat_description', '').strip()
        user_email = data.get('user_email', '').strip()

        if not threat_title or not threat_description:
            return jsonify({'success': False, 'error': 'All fields are required'}), 400

        message = aide_handler.create_threat_hunt(
            threat_title,
            threat_description,
            user_email,
            get_client_ip(),
            prod_ticket_handler,
            CONFIG.xsoar_prod_ui_base_url
        )

        return jsonify({'success': True, 'message': message})

    except Exception as exc:
        logger.error(f"Error creating threat hunt: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred'}), 500


@chat_bp.route('/api/aide/oncall', methods=['GET'])
@auth_helpers.admin_required
@log_web_activity
def api_oncall():
    """API endpoint to get on-call information"""
    try:
        on_call_person = aide_handler.get_oncall_info()
        return jsonify({'success': True, 'data': on_call_person})

    except Exception as exc:
        logger.error(f"Error getting on-call info: {exc}", exc_info=True)
        return jsonify({'success': False, 'error': 'An internal error occurred'}), 500


# --- Page-Context Chat Widget (shared across all dashboard pages) ---

@chat_bp.route('/api/page-chat/stream', methods=['POST'])
# Public: read-only LLM query over client-supplied page context (no mutation);
# rate-limited. Lets the chat widget work on public read-only dashboards.
@limiter.limit("10 per minute")
@log_web_activity
def api_page_chat_stream():
    """Streaming chat widget. The page sends its own context as report_md."""
    try:
        data = request.get_json()
        user_message = (data.get('message') or '').strip()
        report_md = (data.get('report_md') or '').strip()
        session_id = (data.get('session_id') or '').strip()
        language = (data.get('language') or 'English').strip()

        if not user_message:
            return jsonify({'success': False, 'error': 'Message is required'}), 400
        if len(user_message) > 2000:
            return jsonify({'success': False, 'error': 'Message too long (max 2 000 chars)'}), 400
        if not report_md:
            return jsonify({'success': False, 'error': 'No report context provided'}), 400
        if not session_id:
            return jsonify({'success': False, 'error': 'Session ID is required'}), 400

        llm = _get_dp_llm()
        if llm is None:
            return jsonify({'success': False, 'error': 'LLM unavailable'}), 503

        def generate():
            try:
                for payload in page_chat_handler.handle_chat_stream(
                    user_message, report_md, session_id, llm, language=language
                ):
                    yield f"data: {json.dumps(payload)}\n\n"
            except Exception as err:
                logger.error("Page chat stream error: %s", err, exc_info=True)
                yield f"data: {json.dumps({'error': 'Streaming error'})}\n\n"

        return current_app.response_class(
            generate(),
            mimetype='text/event-stream',
            headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
        )

    except Exception as exc:
        logger.error("Page chat error: %s", exc, exc_info=True)
        return jsonify({'success': False, 'error': 'Chat error'}), 500


@chat_bp.route('/api/page-chat/translate', methods=['POST'])
# Public: read-only (see /api/page-chat/stream).
@limiter.limit("20 per minute")
@log_web_activity
def api_page_chat_translate():
    """Translate a list of short UI strings into the requested language. Used to localize chat widget greeting + suggestion chips."""
    try:
        data = request.get_json() or {}
        language = (data.get('language') or '').strip()
        strings = data.get('strings') or []
        if not language or language.lower() == 'english':
            return jsonify({'success': True, 'translations': strings})
        if not isinstance(strings, list) or not strings:
            return jsonify({'success': False, 'error': 'strings (non-empty list) is required'}), 400
        if len(strings) > 30:
            return jsonify({'success': False, 'error': 'too many strings (max 30)'}), 400
        if any(not isinstance(s, str) or len(s) > 500 for s in strings):
            return jsonify({'success': False, 'error': 'each string must be <=500 chars'}), 400

        llm = _get_dp_llm()
        if llm is None:
            return jsonify({'success': False, 'error': 'LLM unavailable'}), 503

        translations = page_chat_handler.translate_strings(strings, language, llm)
        return jsonify({'success': True, 'translations': translations})

    except Exception as exc:
        logger.error("Translate error: %s", exc, exc_info=True)
        return jsonify({'success': False, 'error': 'Translate error'}), 500


@chat_bp.route('/api/page-chat/download', methods=['POST'])
# Public: read-only (see /api/page-chat/stream).
@limiter.limit("10 per minute")
@log_web_activity
def api_page_chat_download():
    """Build a .docx from the supplied chat history and return it as a download."""
    try:
        import datetime as _dt
        from io import BytesIO
        from flask import send_file
        from docx import Document
        from docx.shared import Pt

        data = request.get_json() or {}
        title = (data.get('title') or 'Chat').strip()[:200]
        language = (data.get('language') or 'English').strip()[:50]
        entries = data.get('entries') or []
        if not isinstance(entries, list) or not entries:
            return jsonify({'success': False, 'error': 'entries (non-empty list) required'}), 400
        if len(entries) > 200:
            return jsonify({'success': False, 'error': 'too many entries (max 200)'}), 400

        doc = Document()
        doc.add_heading(title, level=1)
        meta = doc.add_paragraph()
        meta_run = meta.add_run(
            f"Exported {_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')} — Language: {language}"
        )
        meta_run.italic = True
        meta_run.font.size = Pt(9)

        for e in entries:
            role = (e.get('role') or '').strip()
            text = str(e.get('text') or '')[:20000]
            ts = e.get('ts')
            who = 'You' if role == 'user' else 'Assistant'
            ts_str = ''
            if isinstance(ts, (int, float)):
                try:
                    ts_str = ' — ' + _dt.datetime.fromtimestamp(ts / 1000).strftime('%Y-%m-%d %H:%M:%S')
                except Exception:
                    pass
            doc.add_heading(who + ts_str, level=2)
            # Preserve paragraph breaks (blank-line separated); keep line breaks within paragraphs as soft breaks
            for block in text.split('\n\n'):
                p = doc.add_paragraph()
                lines = block.split('\n')
                for i, line in enumerate(lines):
                    p.add_run(line)
                    if i < len(lines) - 1:
                        p.add_run().add_break()

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        stamp = _dt.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'chat-{stamp}.docx',
        )
    except Exception as exc:
        logger.error("Page chat download error: %s", exc, exc_info=True)
        return jsonify({'success': False, 'error': 'Download error'}), 500


@chat_bp.route('/api/page-chat/clear', methods=['POST'])
# Public: read-only (see /api/page-chat/stream).
@limiter.limit("10 per minute")
@log_web_activity
def api_page_chat_clear():
    """Clear conversation history for the caller's session."""
    data = request.get_json(silent=True) or {}
    session_id = (data.get('session_id') or '').strip()
    if not session_id:
        return jsonify({'success': False, 'error': 'Session ID is required'}), 400
    page_chat_handler.clear_history(session_id)
    return jsonify({'success': True})


# --- Docs Library RAG Chat ---

@chat_bp.route('/api/docs-library/chat/stream', methods=['POST'])
@auth_helpers.login_required
@limiter.limit("10 per minute")
@log_web_activity
def api_docs_library_chat_stream():
    """RAG chat over the local document store. Retrieves relevant chunks from ChromaDB then streams an LLM response."""
    from src.components.web import docs_library_chat_handler as dl_chat
    try:
        data = request.get_json()
        user_message = (data.get('message') or '').strip()
        session_id = (data.get('session_id') or '').strip()

        if not user_message:
            return jsonify({'success': False, 'error': 'Message is required'}), 400
        if len(user_message) > 2000:
            return jsonify({'success': False, 'error': 'Message too long (max 2 000 chars)'}), 400
        if not session_id:
            return jsonify({'success': False, 'error': 'Session ID is required'}), 400

        llm = _get_dp_llm()
        if llm is None:
            return jsonify({'success': False, 'error': 'LLM unavailable'}), 503

        def generate():
            try:
                for payload in dl_chat.handle_chat_stream(user_message, session_id, llm):
                    yield f"data: {json.dumps(payload)}\n\n"
            except Exception as err:
                logger.error("Docs library chat stream error: %s", err, exc_info=True)
                yield f"data: {json.dumps({'error': 'Streaming error'})}\n\n"

        return current_app.response_class(
            generate(),
            mimetype='text/event-stream',
            headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'},
        )

    except Exception as exc:
        logger.error("Docs library chat error: %s", exc, exc_info=True)
        return jsonify({'success': False, 'error': 'Chat error'}), 500


@chat_bp.route('/api/docs-library/chat/clear', methods=['POST'])
@auth_helpers.login_required
@limiter.limit("10 per minute")
@log_web_activity
def api_docs_library_chat_clear():
    """Clear docs library chat session history."""
    from src.components.web import docs_library_chat_handler as dl_chat
    data = request.get_json(silent=True) or {}
    session_id = (data.get('session_id') or '').strip()
    if not session_id:
        return jsonify({'success': False, 'error': 'Session ID is required'}), 400
    dl_chat.clear_history(session_id)
    return jsonify({'success': True})
