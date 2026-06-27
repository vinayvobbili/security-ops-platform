"""Third-Party Cyber Risk Assessment (TPCRA) — DB, evidence retrieval, evaluation.

Backs the `/vendor-risk-assessment` web page. An analyst opens an assessment for
a vendor, uploads that vendor's cyber documentation (SOC 2 reports, policies,
pen-test summaries) plus the Aravo questionnaire export, and the system:

  1. indexes the vendor's evidence into a per-assessment vector store,
  2. evaluates each baseline IT-security control against ONLY that vendor's
     evidence (Met / Not Met / Not Applicable + evidence summary, validation
     source, notes/gaps),
  3. lets the analyst review/edit every determination,
  4. exports a fully-populated Due-Diligence (DD) Form as a Word document.

This is the worked twin of `customer_assurance_handler` (RAG + local-LLM drafting),
with one key difference: in Customer Assurance the knowledge base is ONE shared
internal-policy corpus; here the evidence is PER VENDOR, so every assessment gets
its own slice of the evidence collection (filtered by `assessment_id`). That
removes the legacy 20-document ceiling — we control chunking + embedding, so
there is no hard cap on how many vendor documents an assessment can ingest.

v1 SKELETON: the baseline control set below (`BASELINE_CONTROLS`) and the DD Form
export layout are placeholders modelled on a typical due-diligence form. Swap in
the real baseline control document + DD Form template when they land — the
pipeline around them does not change.
"""

import hashlib
import logging
import math
import os
import re
import sqlite3
import threading
from contextlib import contextmanager
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

# Paths — data-isolated under data/tpcra/ (dev twin gets its own copy)
_REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent
DATA_DIR = _REPO_ROOT / "data" / "tpcra"
DB_PATH = DATA_DIR / "tpcra.db"
UPLOADS_DIR = DATA_DIR / "uploads"            # per-assessment vendor docs + DD-form exports
CHROMA_PATH = DATA_DIR / "chroma_evidence"    # dedicated Chroma store (all assessments, filtered)

for _p in (DATA_DIR, UPLOADS_DIR, CHROMA_PATH):
    _p.mkdir(parents=True, exist_ok=True)

EVIDENCE_COLLECTION_NAME = "tpcra_evidence"

# Retrieval confidence gate. If the best evidence chunk for a control falls below
# this similarity, we DO NOT let the model assert "Met" from thin air — the
# control defaults to "Not Met" with the standardized "not provided" note, which
# is the governing rule: absence of evidence is valid assessment data.
EVIDENCE_MIN_SCORE = float(os.environ.get("TPCRA_EVIDENCE_MIN_SCORE", "0.45"))

MAX_FILE_SIZE = 40 * 1024 * 1024  # 40 MB — SOC 2 reports run large
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".txt", ".csv", ".md"}

# Document roles within one assessment.
DOC_KINDS = ["vendor_evidence", "baseline_control", "aravo_export"]

# Enum-ish constants (free-text in DB; canonical values for the UI).
VENDOR_TIERS = ["Tier 1 — Critical", "Tier 2 — High", "Tier 3 — Moderate", "Tier 4 — Low"]
ASSESSMENT_TYPES = ["Initial Due Diligence", "Reassessment", "Targeted Review", "Pre-Contract"]
ASSESSMENT_STATUSES = ["new", "evaluating", "in_review", "ready", "delivered", "archived"]
CONTROL_STATUSES = ["pending", "evaluated", "confirmed", "eval_failed"]
DETERMINATIONS = ["Met", "Not Met", "Not Applicable"]
RISK_RATINGS = ["Low", "Moderate", "High", "Critical"]

# Standardized statements (the locked assessment prompt, section 3). Used whenever the
# evidence does not support a control so the field is never left blank.
STD_NOT_PROVIDED = "Not provided by vendor documentation reviewed as part of this assessment."

# Stamped when an evaluation could not actually run (evidence store / model
# unreachable) so an analyst never mistakes a degraded result for a real one.
EVAL_FAILED_NOTICE = (
    "This control could NOT be evaluated — the evidence index or the analysis "
    "model was unreachable. The determination below is a placeholder, not grounded "
    "in vendor evidence. Re-run the evaluation once connectivity is restored."
)


# ------------------------------------------------------------------ Baseline controls
#
# PLACEHOLDER baseline IT-security control set, modelled on a typical third-party
# cyber due-diligence form (domains drawn from SIG / NIST CSF / ISO 27001 themes).
# Replace with the real baseline control document — `seed_baseline_controls`
# is the only thing that reads this, so swapping it in is a one-line change.

BASELINE_CONTROLS: List[Dict[str, str]] = [
    # Governance & Risk
    {"domain": "Governance & Risk", "control_id": "GRC-1",
     "control_text": "The vendor maintains a documented information security program with executive ownership and annual review."},
    {"domain": "Governance & Risk", "control_id": "GRC-2",
     "control_text": "The vendor performs periodic risk assessments and tracks remediation of identified risks."},
    {"domain": "Governance & Risk", "control_id": "GRC-3",
     "control_text": "Security policies are reviewed and approved at least annually and communicated to personnel."},
    # Access Control & Identity
    {"domain": "Access Control & Identity", "control_id": "IAM-1",
     "control_text": "Multi-factor authentication is enforced for remote access and privileged/administrative accounts."},
    {"domain": "Access Control & Identity", "control_id": "IAM-2",
     "control_text": "Access is granted on least-privilege and reviewed periodically; terminations revoke access promptly."},
    {"domain": "Access Control & Identity", "control_id": "IAM-3",
     "control_text": "Privileged access is logged and monitored, with unique, non-shared administrative accounts."},
    # Data Protection
    {"domain": "Data Protection", "control_id": "DATA-1",
     "control_text": "Data is encrypted in transit using TLS 1.2+ and at rest using AES-256 or equivalent."},
    {"domain": "Data Protection", "control_id": "DATA-2",
     "control_text": "Encryption keys are managed under a documented key-management process with restricted access."},
    {"domain": "Data Protection", "control_id": "DATA-3",
     "control_text": "Data retention and secure disposal procedures are documented and enforced."},
    # Network & Infrastructure Security
    {"domain": "Network & Infrastructure", "control_id": "NET-1",
     "control_text": "Network segmentation, firewalls, and intrusion detection/prevention are in place."},
    {"domain": "Network & Infrastructure", "control_id": "NET-2",
     "control_text": "Endpoint protection (EDR/AV) is deployed and centrally managed across the fleet."},
    {"domain": "Network & Infrastructure", "control_id": "NET-3",
     "control_text": "A documented vulnerability management program patches critical findings within defined SLAs."},
    # Vulnerability & Testing
    {"domain": "Vulnerability & Testing", "control_id": "VULN-1",
     "control_text": "Independent penetration testing is performed at least annually with findings remediated."},
    {"domain": "Vulnerability & Testing", "control_id": "VULN-2",
     "control_text": "Secure software development lifecycle (SAST/DAST, code review) controls are in place where applicable."},
    # Incident Response & Resilience
    {"domain": "Incident Response & Resilience", "control_id": "IR-1",
     "control_text": "A documented incident response plan is tested periodically and defines customer breach-notification timelines."},
    {"domain": "Incident Response & Resilience", "control_id": "IR-2",
     "control_text": "Business continuity and disaster recovery plans are documented and tested, with defined RTO/RPO."},
    {"domain": "Incident Response & Resilience", "control_id": "IR-3",
     "control_text": "System and security event logging is centralized and retained for an appropriate period."},
    # Third-Party & Compliance
    {"domain": "Third-Party & Compliance", "control_id": "TPC-1",
     "control_text": "The vendor manages security of its own subservice organizations / fourth parties."},
    {"domain": "Third-Party & Compliance", "control_id": "TPC-2",
     "control_text": "A current independent audit (SOC 2 Type II, ISO 27001) covers the services in scope."},
    {"domain": "Third-Party & Compliance", "control_id": "TPC-3",
     "control_text": "Security awareness training is delivered to personnel at hire and at least annually."},
]


# ------------------------------------------------------------------ DB setup

@contextmanager
def _get_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


def init_db():
    """Create all tables if they don't exist. Safe to call repeatedly."""
    with _get_connection() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS assessments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                vendor_name      TEXT NOT NULL,
                vendor_tier      TEXT,
                assessment_type  TEXT,
                title            TEXT NOT NULL,
                aravo_ref        TEXT,
                scope_notes      TEXT,
                status           TEXT NOT NULL DEFAULT 'new',
                inherent_risk    TEXT,
                residual_risk    TEXT,
                overall_statement TEXT,
                owner            TEXT,
                assigned_to      TEXT,
                created_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at       DATETIME DEFAULT CURRENT_TIMESTAMP,
                delivered_at     DATETIME
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS controls (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                assessment_id   INTEGER NOT NULL REFERENCES assessments(id) ON DELETE CASCADE,
                seq             INTEGER NOT NULL,
                domain          TEXT,
                control_id      TEXT,
                control_text    TEXT NOT NULL,
                determination   TEXT,
                evidence_summary TEXT,
                validation_source TEXT,
                notes_gaps      TEXT,
                confidence      REAL,
                status          TEXT DEFAULT 'pending',
                updated_at      DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS citations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                control_id  INTEGER NOT NULL REFERENCES controls(id) ON DELETE CASCADE,
                source_path TEXT NOT NULL,
                chunk_text  TEXT,
                score       REAL
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                assessment_id INTEGER NOT NULL REFERENCES assessments(id) ON DELETE CASCADE,
                filename      TEXT NOT NULL,
                path          TEXT NOT NULL,
                kind          TEXT,
                indexed       INTEGER DEFAULT 0,
                chunk_count   INTEGER DEFAULT 0,
                uploaded_at   DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS audit_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                assessment_id INTEGER REFERENCES assessments(id) ON DELETE CASCADE,
                actor   TEXT,
                action  TEXT,
                detail  TEXT,
                at      DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_controls_assessment ON controls(assessment_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_citations_control ON citations(control_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_documents_assessment ON documents(assessment_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_audit_assessment ON audit_log(assessment_id)")
    logger.info(f"TPCRA database initialized at {DB_PATH}")


init_db()


# ------------------------------------------------------------------ Assessments CRUD

def create_assessment(
    vendor_name: str,
    title: str,
    vendor_tier: Optional[str] = None,
    assessment_type: Optional[str] = None,
    aravo_ref: Optional[str] = None,
    scope_notes: Optional[str] = None,
    assigned_to: Optional[str] = None,
    owner: Optional[str] = None,
) -> int:
    with _get_connection() as conn:
        cur = conn.execute(
            """INSERT INTO assessments
               (vendor_name, vendor_tier, assessment_type, title, aravo_ref,
                scope_notes, assigned_to, owner, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'new')""",
            (vendor_name, vendor_tier, assessment_type, title, aravo_ref,
             scope_notes, assigned_to, owner),
        )
        rid = cur.lastrowid
        conn.execute(
            "INSERT INTO audit_log (assessment_id, actor, action, detail) VALUES (?, ?, ?, ?)",
            (rid, owner or "system", "created", f"Assessment opened for {vendor_name}"),
        )
    return rid


def get_assessment(assessment_id: int) -> Optional[Dict[str, Any]]:
    with _get_connection() as conn:
        row = conn.execute("SELECT * FROM assessments WHERE id = ?", (assessment_id,)).fetchone()
    return dict(row) if row else None


def list_assessments(status: Optional[str] = None) -> List[Dict[str, Any]]:
    q = "SELECT * FROM assessments"
    params: List[Any] = []
    if status:
        q += " WHERE status = ?"
        params.append(status)
    q += " ORDER BY updated_at DESC, id DESC"
    with _get_connection() as conn:
        rows = conn.execute(q, params).fetchall()
    return [dict(r) for r in rows]


def update_assessment_status(assessment_id: int, status: str, actor: str = "system") -> bool:
    if status not in ASSESSMENT_STATUSES:
        raise ValueError(f"Invalid status: {status}")
    with _get_connection() as conn:
        delivered = ", delivered_at = CURRENT_TIMESTAMP" if status == "delivered" else ""
        conn.execute(
            f"UPDATE assessments SET status = ?, updated_at = CURRENT_TIMESTAMP{delivered} WHERE id = ?",
            (status, assessment_id),
        )
        conn.execute(
            "INSERT INTO audit_log (assessment_id, actor, action, detail) VALUES (?, ?, ?, ?)",
            (assessment_id, actor, "status", f"→ {status}"),
        )
    return True


def set_risk_ratings(
    assessment_id: int,
    inherent_risk: Optional[str],
    residual_risk: Optional[str],
    overall_statement: Optional[str],
    actor: str = "analyst",
) -> bool:
    with _get_connection() as conn:
        conn.execute(
            """UPDATE assessments
               SET inherent_risk = ?, residual_risk = ?, overall_statement = ?,
                   updated_at = CURRENT_TIMESTAMP
               WHERE id = ?""",
            (inherent_risk, residual_risk, overall_statement, assessment_id),
        )
        conn.execute(
            "INSERT INTO audit_log (assessment_id, actor, action, detail) VALUES (?, ?, ?, ?)",
            (assessment_id, actor, "risk_rated",
             f"inherent={inherent_risk} residual={residual_risk}"),
        )
    return True


def delete_assessment(assessment_id: int) -> bool:
    """Delete an assessment, its rows, and its evidence chunks."""
    with _get_connection() as conn:
        cur = conn.execute("DELETE FROM assessments WHERE id = ?", (assessment_id,))
        deleted = cur.rowcount > 0
    if deleted:
        _delete_assessment_evidence(assessment_id)
    return deleted


# ------------------------------------------------------------------ Controls

def seed_baseline_controls(assessment_id: int) -> int:
    """Populate the assessment with the baseline control set. Idempotent —
    skips if controls already exist for the assessment."""
    with _get_connection() as conn:
        existing = conn.execute(
            "SELECT COUNT(*) AS n FROM controls WHERE assessment_id = ?", (assessment_id,)
        ).fetchone()["n"]
        if existing:
            return 0
        for seq, c in enumerate(BASELINE_CONTROLS, start=1):
            conn.execute(
                """INSERT INTO controls (assessment_id, seq, domain, control_id, control_text)
                   VALUES (?, ?, ?, ?, ?)""",
                (assessment_id, seq, c["domain"], c["control_id"], c["control_text"]),
            )
        conn.execute(
            "INSERT INTO audit_log (assessment_id, actor, action, detail) VALUES (?, ?, ?, ?)",
            (assessment_id, "system", "seeded_controls", f"{len(BASELINE_CONTROLS)} baseline controls"),
        )
    return len(BASELINE_CONTROLS)


def get_control(control_id: int) -> Optional[Dict[str, Any]]:
    with _get_connection() as conn:
        row = conn.execute("SELECT * FROM controls WHERE id = ?", (control_id,)).fetchone()
    return dict(row) if row else None


def list_controls(assessment_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM controls WHERE assessment_id = ? ORDER BY seq", (assessment_id,)
        ).fetchall()
    return [dict(r) for r in rows]


def save_control_determination(
    control_id: int,
    determination: Optional[str],
    evidence_summary: Optional[str],
    validation_source: Optional[str],
    notes_gaps: Optional[str],
    status: str = "evaluated",
    confidence: Optional[float] = None,
    citations: Optional[List[Dict[str, Any]]] = None,
) -> None:
    with _get_connection() as conn:
        conn.execute(
            """UPDATE controls
               SET determination = ?, evidence_summary = ?, validation_source = ?,
                   notes_gaps = ?, status = ?, confidence = ?, updated_at = CURRENT_TIMESTAMP
               WHERE id = ?""",
            (determination, evidence_summary, validation_source, notes_gaps,
             status, confidence, control_id),
        )
        if citations is not None:
            conn.execute("DELETE FROM citations WHERE control_id = ?", (control_id,))
            for c in citations:
                conn.execute(
                    """INSERT INTO citations (control_id, source_path, chunk_text, score)
                       VALUES (?, ?, ?, ?)""",
                    (control_id, c.get("source_path", ""), c.get("chunk_text"), c.get("score")),
                )


def get_citations(control_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT source_path, chunk_text, score FROM citations WHERE control_id = ? ORDER BY score DESC",
            (control_id,),
        ).fetchall()
    return [dict(r) for r in rows]


# ------------------------------------------------------------------ Documents

def save_upload(assessment_id: int, file: FileStorage, kind: str = "vendor_evidence") -> Optional[Dict[str, Any]]:
    if not file or not file.filename:
        return None
    name = secure_filename(file.filename)
    if not name or Path(name).suffix.lower() not in ALLOWED_EXTENSIONS:
        return None
    dest_dir = UPLOADS_DIR / str(assessment_id)
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / name
    file.save(str(dest))
    with _get_connection() as conn:
        cur = conn.execute(
            """INSERT INTO documents (assessment_id, filename, path, kind)
               VALUES (?, ?, ?, ?)""",
            (assessment_id, name, str(dest), kind),
        )
        doc_id = cur.lastrowid
    return {"id": doc_id, "filename": name, "path": str(dest), "kind": kind}


def list_documents(assessment_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM documents WHERE assessment_id = ? ORDER BY id", (assessment_id,)
        ).fetchall()
    return [dict(r) for r in rows]


def get_document(doc_id: int) -> Optional[Dict[str, Any]]:
    with _get_connection() as conn:
        row = conn.execute("SELECT * FROM documents WHERE id = ?", (doc_id,)).fetchone()
    return dict(row) if row else None


# ------------------------------------------------------------------ Audit

def log_audit(assessment_id: int, action: str, detail: str = "", actor: str = "system") -> None:
    with _get_connection() as conn:
        conn.execute(
            "INSERT INTO audit_log (assessment_id, actor, action, detail) VALUES (?, ?, ?, ?)",
            (assessment_id, actor, action, detail),
        )


def get_audit_log(assessment_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT actor, action, detail, at FROM audit_log WHERE assessment_id = ? ORDER BY id DESC",
            (assessment_id,),
        ).fetchall()
    return [dict(r) for r in rows]


# ------------------------------------------------------------------ Evidence vector store
#
# One Chroma collection holds chunks for ALL assessments; every chunk carries an
# `assessment_id` in its metadata and retrieval filters on it, so a control is
# only ever evaluated against its own vendor's evidence.

_evi_state: Dict[str, Any] = {"client": None, "collection": None, "embedding_fn": None}


def _get_evidence_collection():
    if _evi_state["collection"] is not None:
        return _evi_state["collection"]
    import chromadb
    client = chromadb.PersistentClient(path=str(CHROMA_PATH))
    collection = client.get_or_create_collection(
        name=EVIDENCE_COLLECTION_NAME,
        metadata={"description": "TPCRA per-vendor evidence (filtered by assessment_id)"},
    )
    _evi_state["client"] = client
    _evi_state["collection"] = collection
    return collection


def _get_embedding_fn():
    if _evi_state["embedding_fn"] is None:
        from my_bot.document.document_processor import OllamaEmbeddingFunction
        _evi_state["embedding_fn"] = OllamaEmbeddingFunction()
    return _evi_state["embedding_fn"]


def _chunk_id(content: str, source: str, assessment_id: int) -> str:
    h = hashlib.sha256(f"{assessment_id}:{source}:{content}".encode("utf-8")).hexdigest()
    return h[:32]


def _load_document_chunks(path: str, fname: str) -> List[Dict[str, Any]]:
    """Load one file into a list of {content, metadata} chunks. Dispatch by ext,
    mirroring the customer_assurance ingest (table rows stay atomic, narrative is
    split). Returns [] on failure (logged by caller)."""
    from langchain_core.documents import Document as LCDocument
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

    ext = Path(fname).suffix.lower()
    atomic: List[Dict[str, Any]] = []
    splittable: List[LCDocument] = []

    if ext == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader
        for d in PyPDFLoader(path).load():
            d.metadata = {**(d.metadata or {}), "source": fname, "element_type": "pdf_page"}
            splittable.append(d)
    elif ext == ".docx":
        from my_bot.document.document_processor import load_word_structured
        for item in load_word_structured(path):
            if item["metadata"].get("element_type") == "table_row":
                atomic.append(item)
            else:
                splittable.append(LCDocument(page_content=item["content"], metadata=item["metadata"]))
    elif ext == ".xlsx":
        from my_bot.document.document_processor import load_excel_structured
        atomic.extend(load_excel_structured(path))
    elif ext == ".doc":
        from langchain_community.document_loaders import UnstructuredWordDocumentLoader
        for d in UnstructuredWordDocumentLoader(path).load():
            d.metadata = {**(d.metadata or {}), "source": fname, "element_type": "legacy_doc"}
            splittable.append(d)
    elif ext == ".xls":
        from langchain_community.document_loaders import UnstructuredExcelLoader
        for d in UnstructuredExcelLoader(path).load():
            d.metadata = {**(d.metadata or {}), "source": fname, "element_type": "legacy_xls"}
            splittable.append(d)
    else:  # .txt / .csv / .md
        text = Path(path).read_text(errors="ignore")
        splittable.append(LCDocument(page_content=text, metadata={"source": fname, "element_type": "text"}))

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500, chunk_overlap=300,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
    )
    chunks: List[Dict[str, Any]] = []
    for c in (splitter.split_documents(splittable) if splittable else []):
        meta = dict(c.metadata or {})
        meta.setdefault("source", fname)
        meta.setdefault("element_type", "chunk")
        chunks.append({"content": c.page_content, "metadata": meta})
    chunks.extend(atomic)
    return chunks


def ingest_assessment_documents(assessment_id: int, reset: bool = False) -> Dict[str, Any]:
    """Index every not-yet-indexed document for an assessment into the evidence
    store (tagged with assessment_id). Returns a summary dict."""
    summary = {"files": 0, "chunks_added": 0, "chunks_skipped": 0, "errors": []}

    if reset:
        _delete_assessment_evidence(assessment_id)
        with _get_connection() as conn:
            conn.execute(
                "UPDATE documents SET indexed = 0, chunk_count = 0 WHERE assessment_id = ?",
                (assessment_id,),
            )

    docs = [d for d in list_documents(assessment_id) if not d.get("indexed")]
    if not docs:
        return summary
    summary["files"] = len(docs)

    collection = _get_evidence_collection()
    try:
        existing_ids = set(collection.get(where={"assessment_id": assessment_id}).get("ids") or [])
    except Exception as e:
        logger.warning(f"[tpcra] could not read existing ids: {e}")
        existing_ids = set()

    embedding_fn = _get_embedding_fn()
    for d in docs:
        try:
            chunks = _load_document_chunks(d["path"], d["filename"])
        except Exception as e:
            summary["errors"].append(f"{d['filename']}: {e}")
            logger.error(f"[tpcra] load failed {d['filename']}: {e}", exc_info=True)
            continue

        ids, contents, metadatas = [], [], []
        for ch in chunks:
            cid = _chunk_id(ch["content"], d["filename"], assessment_id)
            if cid in existing_ids or cid in ids:
                summary["chunks_skipped"] += 1
                continue
            meta = {
                k: (v if isinstance(v, (str, int, float, bool)) else str(v))
                for k, v in (ch["metadata"] or {}).items() if v is not None
            }
            meta["assessment_id"] = assessment_id
            meta["doc_kind"] = d.get("kind") or "vendor_evidence"
            ids.append(cid)
            contents.append(ch["content"])
            metadatas.append(meta)

        added = 0
        if ids:
            try:
                embeddings = embedding_fn(contents)
                collection.upsert(ids=ids, documents=contents, metadatas=metadatas, embeddings=embeddings)
                added = len(ids)
                summary["chunks_added"] += added
                existing_ids.update(ids)
            except Exception as e:
                summary["errors"].append(f"{d['filename']} upsert: {e}")
                logger.error(f"[tpcra] upsert failed {d['filename']}: {e}", exc_info=True)
                continue

        with _get_connection() as conn:
            conn.execute(
                "UPDATE documents SET indexed = 1, chunk_count = ? WHERE id = ?",
                (added, d["id"]),
            )

    log_audit(assessment_id, "indexed",
              f"{summary['chunks_added']} chunk(s) from {summary['files']} file(s)")
    logger.info(f"[tpcra] ingest a#{assessment_id}: {summary}")
    return summary


def _delete_assessment_evidence(assessment_id: int) -> None:
    try:
        collection = _get_evidence_collection()
        collection.delete(where={"assessment_id": assessment_id})
    except Exception as e:
        logger.warning(f"[tpcra] evidence delete failed a#{assessment_id}: {e}")


def evidence_stats(assessment_id: int) -> Dict[str, Any]:
    docs = list_documents(assessment_id)
    indexed = [d for d in docs if d.get("indexed")]
    chunk_count = sum(d.get("chunk_count") or 0 for d in indexed)
    return {
        "doc_count": len(docs),
        "indexed_count": len(indexed),
        "pending_count": len(docs) - len(indexed),
        "chunk_count": chunk_count,
    }


def _cosine(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(x * x for x in b))
    return dot / (na * nb) if na and nb else 0.0


# ------------------------------------------------------------------ Evaluation

def _build_eval_prompt(control: Dict[str, Any], context: str) -> str:
    """Build the per-control evaluation prompt. Encodes the locked TPCRA rules:
    evidence-only, no inference, absence-of-evidence is valid data, fixed output
    fields. The model returns a labelled block we parse below."""
    return f"""You are a Third-Party Cyber Risk Assessment Analyst. Evaluate ONE baseline IT-security control against ONLY the vendor evidence provided below. Your evaluation is reviewed by an analyst before release.

RULES:
- Use ONLY the vendor evidence provided. Do NOT infer, assume, or speculate about controls, tooling, or processes.
- The absence of evidence is valid assessment data. If the evidence does not support the control, the determination is "Not Met". "Not Met" means lack of evidence, NOT a control failure.
- Use "Not Applicable" ONLY when the control is clearly out of scope for this vendor's services.
- Never invent specific dates, names, versions, or counts that the evidence does not contain.
- Be neutral, factual, and executive-ready. No audit, QA, or second-line language.

CONTROL [{control.get('control_id')}] — {control.get('domain')}:
{control.get('control_text')}

VENDOR EVIDENCE:
{context if context.strip() else "(No relevant evidence was retrieved for this control.)"}

Respond in EXACTLY this format, each field on its own line:
DETERMINATION: <Met | Not Met | Not Applicable>
EVIDENCE SUMMARY: <plain-language explanation grounded in the evidence; if none, write "{STD_NOT_PROVIDED}">
VALIDATION SOURCE: <the specific document name(s) the evidence came from, or "{STD_NOT_PROVIDED}">
NOTES / GAPS: <scope limitations, missing artifacts, or follow-ups; "None" if fully supported>"""


_FIELD_RE = {
    "determination": re.compile(r"DETERMINATION:\s*(.+)", re.IGNORECASE),
    "evidence_summary": re.compile(r"EVIDENCE SUMMARY:\s*(.+?)(?=\n[A-Z/ ]+:|\Z)", re.IGNORECASE | re.DOTALL),
    "validation_source": re.compile(r"VALIDATION SOURCE:\s*(.+?)(?=\n[A-Z/ ]+:|\Z)", re.IGNORECASE | re.DOTALL),
    "notes_gaps": re.compile(r"NOTES\s*/?\s*GAPS:\s*(.+?)(?=\n[A-Z/ ]+:|\Z)", re.IGNORECASE | re.DOTALL),
}


def _normalize_determination(raw: str) -> str:
    t = (raw or "").strip().lower()
    if t.startswith("not applic") or t in ("na", "n/a"):
        return "Not Applicable"
    if t.startswith("not met") or t.startswith("n "):
        return "Not Met"
    if t.startswith("met") or t.startswith("y "):
        return "Met"
    # Defensive default: ambiguous answers are treated as lack of evidence.
    return "Not Met"


def _parse_eval(raw: str) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for field, rx in _FIELD_RE.items():
        m = rx.search(raw or "")
        out[field] = m.group(1).strip() if m else ""
    out["determination"] = _normalize_determination(out.get("determination", ""))
    return out


def _degraded_eval(control: Dict[str, Any], detail: str) -> Dict[str, Any]:
    """Evaluation could not run (evidence/model unreachable). Save a clearly-
    flagged placeholder so the analyst re-runs it rather than trusting it."""
    save_control_determination(
        control["id"],
        determination="Not Met",
        evidence_summary=EVAL_FAILED_NOTICE,
        validation_source=STD_NOT_PROVIDED,
        notes_gaps=f"Evaluation degraded — {detail}.",
        status="eval_failed",
        confidence=None,
        citations=[],
    )
    log_audit(control["assessment_id"], "evaluated", f"{control.get('control_id')} (degraded — {detail})")
    return {"failed": True, "message": EVAL_FAILED_NOTICE, "determination": "Not Met"}


def evaluate_control(control_id: int) -> Dict[str, Any]:
    """Evaluate one control against its assessment's evidence.

    Orchestration (retrieve -> rerank -> confidence gate -> draft) runs through
    the shared `attestq` kernel via attestq_bridge; this function owns only the
    TPCRA-specific bits: the evidence-only prompt, the labelled-field parser, the
    "Not Met = absence of evidence" gate outcome, and persistence. The kernel and
    its adapters are imported lazily so the page still loads if attestq is absent.
    """
    control = get_control(control_id)
    if not control:
        raise ValueError("control not found")

    try:
        from attestq import Citation, Question

        from . import attestq_bridge as bridge
    except Exception as e:
        logger.error(f"[tpcra] attestq unavailable: {e}", exc_info=True)
        return _degraded_eval(control, "kernel unavailable")

    # The LLM (local). Not configured/unavailable -> degraded.
    try:
        chat = bridge.LocalLLMChat(timeout=90)
        if not chat.is_configured():
            return _degraded_eval(control, "model not configured")
    except Exception as e:
        logger.warning(f"[tpcra] the LLM unavailable: {e}")
        return _degraded_eval(control, "model unavailable")

    store = bridge.ChromaCollectionStore(
        _get_evidence_collection(), namespace_field="assessment_id", namespace_cast=int,
    )

    def _prompt_builder(question, hits):
        context = "\n\n".join(f"[{i+1}] Source: {h.source}\n{h.text}" for i, h in enumerate(hits))
        ctl = {"control_id": question.id, "domain": question.domain, "control_text": question.prompt}
        return _build_eval_prompt(ctl, context)

    def _parser(raw, hits):
        parsed = _parse_eval(raw)
        cites = [Citation(source=h.source, snippet=(h.text or "")[:500],
                          score=h.score or 0.0, chunk_id=h.id) for h in hits[:8]]
        return parsed["determination"], (parsed.get("evidence_summary") or STD_NOT_PROVIDED), cites

    engine = bridge.build_engine(
        chat=chat,
        embed=_get_embedding_fn(),
        store=store,
        prompt_builder=_prompt_builder,
        response_parser=_parser,
        min_confidence=EVIDENCE_MIN_SCORE,
        insufficient_determination="Not Met",
        insufficient_summary=STD_NOT_PROVIDED,
        k=12,
        rerank_top_k=8,
    )

    question = Question(
        id=control.get("control_id") or str(control_id),
        prompt=control["control_text"],
        choices=["Met", "Not Met", "Not Applicable"],
        domain=control.get("domain"),
    )

    try:
        ans = engine.evaluate(question, namespace=str(control["assessment_id"]))
    except Exception as e:
        logger.error(f"[tpcra] evaluation failed: {e}", exc_info=True)
        return _degraded_eval(control, "model error")

    # Confidence gate fired: thin/empty evidence -> Not Met, no model call.
    if ans.insufficient_evidence:
        save_control_determination(
            control_id,
            determination="Not Met",
            evidence_summary=STD_NOT_PROVIDED,
            validation_source=STD_NOT_PROVIDED,
            notes_gaps="No vendor evidence above the relevance threshold was found for this control.",
            status="evaluated",
            confidence=ans.confidence,
            citations=[],
        )
        log_audit(control["assessment_id"], "evaluated",
                  f"{control.get('control_id')} (no evidence, score={ans.confidence})")
        return {"determination": "Not Met", "source": "gate", "confidence": ans.confidence}

    # Model answered: re-parse the raw block for the full TPCRA field set.
    parsed = _parse_eval(ans.raw)
    save_control_determination(
        control_id,
        determination=parsed["determination"],
        evidence_summary=parsed.get("evidence_summary") or STD_NOT_PROVIDED,
        validation_source=parsed.get("validation_source") or STD_NOT_PROVIDED,
        notes_gaps=parsed.get("notes_gaps") or "None",
        status="evaluated",
        confidence=ans.confidence,
        citations=bridge.citations_to_rows(ans.citations),
    )
    src = chat.last_source
    log_audit(control["assessment_id"], "evaluated",
              f"{control.get('control_id')} = {parsed['determination']} ({src})")
    return {"determination": parsed["determination"], "source": src, "confidence": ans.confidence, **parsed}


# ------------------------------------------------------------ Background evaluation
#
# A full baseline run is many LLM calls; firing them in one request would blow
# the timeout and the gateway rate limit. "Evaluate All" kicks off a paced
# background worker that does ONE control at a time, persisting as it goes. The
# controls table is the source of truth, so progress survives a page refresh.

_EVALUATABLE = ("pending", "eval_failed")
EVAL_PACE_SECONDS = 1.0
AVG_SECONDS_PER_CONTROL = 7.0

_WORKERS: Dict[int, Dict[str, Any]] = {}
_WORKER_LOCK = threading.Lock()


def _pending_controls(assessment_id: int) -> List[Dict[str, Any]]:
    return [c for c in list_controls(assessment_id) if c.get("status") in _EVALUATABLE]


def eval_all_status(assessment_id: int) -> Dict[str, Any]:
    controls = list_controls(assessment_id)
    total = len(controls)
    done = sum(1 for c in controls if c.get("status") in ("evaluated", "confirmed"))
    pending = sum(1 for c in controls if c.get("status") in _EVALUATABLE)
    with _WORKER_LOCK:
        running = assessment_id in _WORKERS
    eta = int(pending * AVG_SECONDS_PER_CONTROL) if running else 0
    return {"total": total, "done": done, "pending": pending, "running": running, "eta_seconds": eta}


def start_eval_all(assessment_id: int) -> Dict[str, Any]:
    with _WORKER_LOCK:
        if assessment_id in _WORKERS:
            return {"started": False, **eval_all_status(assessment_id)}
        pending = _pending_controls(assessment_id)
        if not pending:
            return {"started": False, **eval_all_status(assessment_id)}
        _WORKERS[assessment_id] = {"total": len(pending)}
    update_assessment_status(assessment_id, "evaluating", actor="system")
    t = threading.Thread(target=_eval_all_worker, args=(assessment_id,), daemon=True)
    t.start()
    return {"started": True, **eval_all_status(assessment_id)}


def _eval_all_worker(assessment_id: int) -> None:
    import time
    try:
        for control in _pending_controls(assessment_id):
            try:
                evaluate_control(control["id"])
            except Exception as e:
                logger.error(f"[tpcra] eval failed control#{control['id']}: {e}", exc_info=True)
            time.sleep(EVAL_PACE_SECONDS)
    finally:
        with _WORKER_LOCK:
            _WORKERS.pop(assessment_id, None)
        # Move to review once the batch is done (unless analyst already advanced it).
        a = get_assessment(assessment_id)
        if a and a.get("status") == "evaluating":
            update_assessment_status(assessment_id, "in_review", actor="system")


# ------------------------------------------------------------------ Risk synthesis

def synthesize_risk(assessment_id: int) -> Dict[str, Any]:
    """Ask the model to propose inherent/residual ratings + an overall statement
    from the control determinations. Analyst-editable afterward. Falls back to a
    deterministic heuristic if the model is unavailable."""
    controls = list_controls(assessment_id)
    a = get_assessment(assessment_id)
    if not a:
        raise ValueError("assessment not found")

    met = sum(1 for c in controls if c.get("determination") == "Met")
    not_met = sum(1 for c in controls if c.get("determination") == "Not Met")
    na = sum(1 for c in controls if c.get("determination") == "Not Applicable")
    scored = met + not_met
    coverage = (met / scored) if scored else 0.0

    def _heuristic() -> Tuple[str, str, str]:
        if coverage >= 0.9:
            residual = "Low"
        elif coverage >= 0.75:
            residual = "Moderate"
        elif coverage >= 0.5:
            residual = "High"
        else:
            residual = "Critical"
        inherent = a.get("vendor_tier", "").split("—")[-1].strip() or "Moderate"
        inherent = {"Critical": "Critical", "High": "High", "Moderate": "Moderate", "Low": "Low"}.get(inherent, "Moderate")
        stmt = (f"Of {scored} in-scope baseline controls, {met} were Met and {not_met} were Not Met "
                f"(evidence not provided), with {na} Not Applicable. Residual risk is assessed as "
                f"{residual} based on evidence coverage of {coverage:.0%}.")
        return inherent, residual, stmt

    summary = "\n".join(
        f"- [{c.get('control_id')}] {c.get('domain')}: {c.get('determination') or 'pending'}"
        for c in controls
    )
    prompt = f"""You are a Third-Party Cyber Risk Assessment Analyst writing the risk summary for a vendor due-diligence form. Base your assessment ONLY on the control determinations below. Be neutral, evidence-based, and executive-ready. "Not Met" reflects lack of provided evidence, not a confirmed failure.

Vendor: {a.get('vendor_name')} | Tier: {a.get('vendor_tier') or 'unspecified'}
Control determinations ({met} Met, {not_met} Not Met, {na} N/A):
{summary}

Respond in EXACTLY this format:
INHERENT RISK: <Low | Moderate | High | Critical>
RESIDUAL RISK: <Low | Moderate | High | Critical>
OVERALL STATEMENT: <2-4 sentence neutral, executive-ready risk statement>"""

    inherent = residual = statement = None
    try:
        from my_bot.utils.llm_factory import create_llm
        from langchain_core.messages import HumanMessage
        llm = create_llm(timeout=90)
        resp = llm.invoke([HumanMessage(content=prompt)])
        raw = (resp.content or "")
        if raw:
            mi = re.search(r"INHERENT RISK:\s*(\w+)", raw, re.IGNORECASE)
            mr = re.search(r"RESIDUAL RISK:\s*(\w+)", raw, re.IGNORECASE)
            ms = re.search(r"OVERALL STATEMENT:\s*(.+)", raw, re.IGNORECASE | re.DOTALL)
            if mi:
                inherent = mi.group(1).strip().title()
            if mr:
                residual = mr.group(1).strip().title()
            if ms:
                statement = ms.group(1).strip()
    except Exception as e:
        logger.warning(f"[tpcra] risk synthesis model failed, using heuristic: {e}")

    if not (inherent and residual and statement):
        h_inherent, h_residual, h_stmt = _heuristic()
        inherent = inherent or h_inherent
        residual = residual or h_residual
        statement = statement or h_stmt

    set_risk_ratings(assessment_id, inherent, residual, statement, actor="system")
    return {"inherent_risk": inherent, "residual_risk": residual, "overall_statement": statement}


# ------------------------------------------------------------------ DD Form export

def export_dd_form_docx(assessment_id: int) -> Optional[Path]:
    """Export the populated Due-Diligence Form as a .docx.

    PLACEHOLDER layout — replace with the real DD Form template. Every field
    is populated (the mandatory rule): a control with no determination still
    emits the standardized "not provided" statement rather than a blank cell."""
    try:
        from docx import Document
    except ImportError:
        logger.error("[tpcra] python-docx not available")
        return None

    a = get_assessment(assessment_id)
    if not a:
        return None
    controls = list_controls(assessment_id)

    doc = Document()
    doc.add_heading(f"Third-Party Cybersecurity Due-Diligence Form — {a['vendor_name']}", level=0)
    doc.add_paragraph(f"Assessment: {a['title']}")
    doc.add_paragraph(f"Vendor Tier: {a.get('vendor_tier') or 'Not specified'} | "
                      f"Type: {a.get('assessment_type') or 'Not specified'}")
    if a.get("aravo_ref"):
        doc.add_paragraph(f"Aravo Reference: {a['aravo_ref']}")
    doc.add_paragraph(f"Scope: {a.get('scope_notes') or STD_NOT_PROVIDED}")
    doc.add_paragraph("")

    # Risk summary
    doc.add_heading("Risk Assessment Summary", level=1)
    doc.add_paragraph(f"Inherent Risk Rating: {a.get('inherent_risk') or 'Pending'}")
    doc.add_paragraph(f"Residual Risk Rating: {a.get('residual_risk') or 'Pending'}")
    doc.add_paragraph(f"Overall Risk Statement: {a.get('overall_statement') or STD_NOT_PROVIDED}")
    doc.add_paragraph("")

    # Baseline control evaluation
    doc.add_heading("Baseline IT Security Control Evaluation", level=1)
    current_domain = None
    for c in controls:
        if c.get("domain") and c["domain"] != current_domain:
            current_domain = c["domain"]
            doc.add_heading(current_domain, level=2)
        doc.add_heading(f"[{c.get('control_id')}] {c.get('control_text')}", level=3)
        doc.add_paragraph(f"Determination: {c.get('determination') or 'Not Met'}")
        doc.add_paragraph(f"Evidence Summary: {c.get('evidence_summary') or STD_NOT_PROVIDED}")
        doc.add_paragraph(f"Validation Source: {c.get('validation_source') or STD_NOT_PROVIDED}")
        doc.add_paragraph(f"Notes / Gaps: {c.get('notes_gaps') or 'None'}")
        if c.get("status") == "eval_failed":
            p = doc.add_paragraph()
            r = p.add_run("⚠ Automated evaluation failed — re-run before relying on this determination")
            r.italic = True
        doc.add_paragraph("")

    # Sign-off (section 8 defaults)
    doc.add_heading("Reviewer & Sign-Off", level=1)
    doc.add_paragraph("Assessment Completed By: Third-Party Risk Assessment Analyst")
    doc.add_paragraph("Peer Reviewer: Pending internal review")
    doc.add_paragraph("Final Approval Date: To be completed in Aravo")

    exports_dir = UPLOADS_DIR / str(assessment_id)
    exports_dir.mkdir(parents=True, exist_ok=True)
    from datetime import datetime
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = exports_dir / f"DD_Form_{a['vendor_name'].replace(' ', '_')}_{stamp}.docx"
    doc.save(str(path))

    with _get_connection() as conn:
        conn.execute(
            "INSERT INTO documents (assessment_id, filename, path, kind, indexed) VALUES (?, ?, ?, ?, 1)",
            (assessment_id, path.name, str(path), "dd_form_export"),
        )
    log_audit(assessment_id, "exported", path.name)
    return path


# ------------------------------------------------------------------ Sample seeding

def seed_sample_assessment(run_eval: bool = True, wipe_existing: bool = True) -> Dict[str, Any]:
    """Create one fully-worked sample assessment for a fictional vendor.

    Writes the sample evidence docs to disk, indexes them, and (by default)
    evaluates every control against them and synthesizes the risk summary so a
    reviewer opens a finished example. Idempotent: with ``wipe_existing`` it
    removes any prior sample (matched on vendor_name + title) first.

    Returns a summary dict. Evaluation uses the live model; if the gateway/index
    is unreachable, affected controls land as ``eval_failed`` (visibly flagged)
    rather than failing the whole seed.
    """
    from src.components.web.tpcra_sample import SAMPLE_VENDOR, SAMPLE_DOCS

    if wipe_existing:
        for a in list_assessments():
            if a["vendor_name"] == SAMPLE_VENDOR["vendor_name"] and a["title"] == SAMPLE_VENDOR["title"]:
                delete_assessment(a["id"])

    aid = create_assessment(
        vendor_name=SAMPLE_VENDOR["vendor_name"],
        title=SAMPLE_VENDOR["title"],
        vendor_tier=SAMPLE_VENDOR.get("vendor_tier"),
        assessment_type=SAMPLE_VENDOR.get("assessment_type"),
        aravo_ref=SAMPLE_VENDOR.get("aravo_ref"),
        scope_notes=SAMPLE_VENDOR.get("scope_notes"),
        owner="sample",
    )
    seed_baseline_controls(aid)

    # Write the sample evidence files to the assessment's upload dir + register them.
    dest_dir = UPLOADS_DIR / str(aid)
    dest_dir.mkdir(parents=True, exist_ok=True)
    for doc in SAMPLE_DOCS:
        path = dest_dir / doc["filename"]
        path.write_text(doc["content"])
        with _get_connection() as conn:
            conn.execute(
                "INSERT INTO documents (assessment_id, filename, path, kind) VALUES (?, ?, ?, ?)",
                (aid, doc["filename"], str(path), "vendor_evidence"),
            )
    log_audit(aid, "uploaded", f"{len(SAMPLE_DOCS)} sample evidence document(s)", actor="sample")

    summary: Dict[str, Any] = {"assessment_id": aid, "documents": len(SAMPLE_DOCS)}
    ingest = ingest_assessment_documents(aid)
    summary["chunks_indexed"] = ingest.get("chunks_added", 0)
    summary["index_errors"] = ingest.get("errors", [])

    if run_eval:
        evaluated = 0
        for control in list_controls(aid):
            try:
                evaluate_control(control["id"])
                evaluated += 1
            except Exception as e:
                logger.warning(f"[tpcra] sample eval failed control#{control['id']}: {e}")
        summary["controls_evaluated"] = evaluated
        try:
            summary["risk"] = synthesize_risk(aid)
        except Exception as e:
            logger.warning(f"[tpcra] sample risk synthesis failed: {e}")
        update_assessment_status(aid, "in_review", actor="sample")

    logger.info(f"[tpcra] sample assessment seeded: {summary}")
    return summary
