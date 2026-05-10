"""Customer Assurance drafting assistant — DB, KB retrieval, and CRUD helpers.

This handler backs the `/customer-assurance` web page. Analysts intake customer
security questionnaires (RFPs, SOC2 questionnaires, vendor security reviews),
the system retrieves relevant policy/evidence chunks from a dedicated ChromaDB
collection, an LLM drafts answers, and analysts review/edit/export.

Schema follows the ai_intake_handler pattern (inline CREATE TABLE, auto-init on
import). KB retrieval reuses `ChromaRetriever` from `document_processor` but
points at a separate collection (`customer_assurance_kb`) so the docs_library
`local_documents` collection stays untouched.
"""

import hashlib
import logging
import math
import os
import sqlite3
from contextlib import contextmanager
from pathlib import Path
from typing import Any, Dict, List, Optional

from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename

logger = logging.getLogger(__name__)

# Paths
_REPO_ROOT = Path(__file__).resolve().parent.parent.parent.parent
DATA_DIR = _REPO_ROOT / "data" / "customer_assurance"
DB_PATH = DATA_DIR / "customer_assurance.db"
UPLOADS_DIR = DATA_DIR / "uploads"           # per-request inbound questionnaires + exports
KB_SOURCE_DIR = DATA_DIR / "kb_source"       # raw policy/SOC2/standard-response docs
CHROMA_KB_PATH = DATA_DIR / "chroma_kb"      # dedicated Chroma store

for _p in (DATA_DIR, UPLOADS_DIR, KB_SOURCE_DIR, CHROMA_KB_PATH):
    _p.mkdir(parents=True, exist_ok=True)

KB_COLLECTION_NAME = "customer_assurance_kb"

# Past-answer reuse thresholds (cosine similarity over question embeddings).
# At/above HIGH: reuse the past final_answer verbatim, skip the LLM.
# Between MED and HIGH: include the past answer as a priority context block.
# Below MED: ignore.
PAST_ANSWER_HIGH_THRESHOLD = float(os.environ.get("CA_PAST_ANSWER_HIGH_THRESHOLD", "0.88"))
PAST_ANSWER_MED_THRESHOLD = float(os.environ.get("CA_PAST_ANSWER_MED_THRESHOLD", "0.72"))

# KB retrieval confidence gate. If the best KB chunk falls below this similarity
# and no past approved answer is available, we auto-route the question to SME
# review instead of letting the LLM draft from thin evidence. Protects external
# customers from confidently-wrong answers.
KB_RETRIEVAL_MIN_SCORE = float(os.environ.get("CA_KB_MIN_SCORE", "0.55"))

MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 MB — questionnaires can be large
ALLOWED_EXTENSIONS = {
    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.txt', '.csv', '.md',
}

# Enum-ish constants (free-text in DB; these are just canonical values)
SEGMENTS = ["National", "Regional", "Small Market", "Pensions", "Public Sector", "MIM", "Other"]
REQUEST_TYPES = ["Questionnaire", "General Questions", "RFP", "Meeting Request",
                 "Vulnerability Request", "On-Site Assessment"]
SOURCE_FORMATS = ["Online", "Email", "Excel", "Word", "PowerPoint", "PDF"]
PRIORITIES = ["Low", "Medium", "High", "Urgent"]
REQUEST_STATUSES = ["new", "drafting", "needs_legal", "ready", "delivered", "archived"]
QUESTION_STATUSES = ["pending", "drafted", "approved", "needs_sme"]


# ------------------------------------------------------------------ DB setup

@contextmanager
def _get_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def init_db():
    """Create all tables if they don't exist. Safe to call repeatedly."""
    with _get_connection() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS requests (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                customer_name        TEXT NOT NULL,
                customer_segment     TEXT NOT NULL,
                account_team_contact TEXT,
                request_type         TEXT NOT NULL,
                source_format        TEXT,
                due_date             DATE,
                priority             TEXT,
                title                TEXT NOT NULL,
                raw_text             TEXT,
                notes                TEXT,
                status               TEXT NOT NULL DEFAULT 'new',
                needs_legal_review   INTEGER DEFAULT 0,
                legal_note           TEXT,
                assigned_to          TEXT,
                archer_ref           TEXT,
                salesforce_ref       TEXT,
                created_at           DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at           DATETIME DEFAULT CURRENT_TIMESTAMP,
                delivered_at         DATETIME
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS questions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                request_id   INTEGER NOT NULL REFERENCES requests(id) ON DELETE CASCADE,
                seq          INTEGER NOT NULL,
                section      TEXT,
                question     TEXT NOT NULL,
                draft_answer TEXT,
                final_answer TEXT,
                confidence   REAL,
                status       TEXT DEFAULT 'pending',
                sme_owner    TEXT,
                updated_at   DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS citations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                question_id  INTEGER NOT NULL REFERENCES questions(id) ON DELETE CASCADE,
                source_path  TEXT NOT NULL,
                chunk_text   TEXT,
                score        REAL
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS uploads (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                request_id   INTEGER NOT NULL REFERENCES requests(id) ON DELETE CASCADE,
                filename     TEXT NOT NULL,
                path         TEXT NOT NULL,
                kind         TEXT,
                uploaded_at  DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS audit_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                request_id INTEGER REFERENCES requests(id) ON DELETE CASCADE,
                actor      TEXT,
                action     TEXT,
                detail     TEXT,
                at         DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_questions_request ON questions(request_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_citations_question ON citations(question_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_uploads_request ON uploads(request_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_audit_request ON audit_log(request_id)")

        # Migration: source coordinates for round-trip xlsx export. Questions
        # extracted from an uploaded .xlsx remember the (sheet, row,
        # response_col) they came from so the export can write final answers
        # back into the customer's original spreadsheet instead of producing
        # a fresh .docx. Older rows / pasted-text questions leave these NULL
        # and fall back to the .docx export path.
        existing_cols = {r["name"] for r in conn.execute("PRAGMA table_info(questions)").fetchall()}
        if "source_sheet" not in existing_cols:
            conn.execute("ALTER TABLE questions ADD COLUMN source_sheet TEXT")
        if "source_row" not in existing_cols:
            conn.execute("ALTER TABLE questions ADD COLUMN source_row INTEGER")
        if "source_response_col" not in existing_cols:
            conn.execute("ALTER TABLE questions ADD COLUMN source_response_col INTEGER")
    logger.info(f"Customer Assurance database initialized at {DB_PATH}")


init_db()


# ------------------------------------------------------------------ Requests CRUD

def create_request(
    customer_name: str,
    customer_segment: str,
    request_type: str,
    title: str,
    *,
    account_team_contact: Optional[str] = None,
    source_format: Optional[str] = None,
    due_date: Optional[str] = None,
    priority: Optional[str] = None,
    raw_text: Optional[str] = None,
    notes: Optional[str] = None,
    assigned_to: Optional[str] = None,
    archer_ref: Optional[str] = None,
    salesforce_ref: Optional[str] = None,
) -> int:
    with _get_connection() as conn:
        cur = conn.execute(
            """INSERT INTO requests
               (customer_name, customer_segment, account_team_contact, request_type,
                source_format, due_date, priority, title, raw_text, notes,
                assigned_to, archer_ref, salesforce_ref)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (customer_name, customer_segment, account_team_contact, request_type,
             source_format, due_date, priority, title, raw_text, notes,
             assigned_to, archer_ref, salesforce_ref),
        )
        request_id = cur.lastrowid
        conn.execute(
            "INSERT INTO audit_log (request_id, actor, action, detail) VALUES (?, ?, ?, ?)",
            (request_id, assigned_to or "system", "created", f"Request '{title}' created"),
        )
    return request_id


def get_request(request_id: int) -> Optional[Dict[str, Any]]:
    with _get_connection() as conn:
        row = conn.execute("SELECT * FROM requests WHERE id = ?", (request_id,)).fetchone()
        return dict(row) if row else None


def list_requests(
    status: Optional[str] = None,
    segment: Optional[str] = None,
    exclude_demo: bool = False,
) -> List[Dict[str, Any]]:
    sql = "SELECT * FROM requests"
    params: List[Any] = []
    clauses = []
    if status:
        clauses.append("status = ?")
        params.append(status)
    if segment:
        clauses.append("customer_segment = ?")
        params.append(segment)
    if exclude_demo:
        clauses.append("(assigned_to IS NULL OR assigned_to != 'demo')")
    if clauses:
        sql += " WHERE " + " AND ".join(clauses)
    sql += " ORDER BY COALESCE(due_date, '9999-12-31') ASC, created_at DESC"
    with _get_connection() as conn:
        rows = conn.execute(sql, params).fetchall()
        return [dict(r) for r in rows]


def update_request_status(request_id: int, status: str, actor: str = "system") -> bool:
    if status not in REQUEST_STATUSES:
        raise ValueError(f"Invalid status: {status}")
    with _get_connection() as conn:
        extra = ""
        params: List[Any] = [status]
        if status == "delivered":
            extra = ", delivered_at = CURRENT_TIMESTAMP"
        params.append(request_id)
        cur = conn.execute(
            f"UPDATE requests SET status = ?, updated_at = CURRENT_TIMESTAMP{extra} WHERE id = ?",
            params,
        )
        if cur.rowcount:
            conn.execute(
                "INSERT INTO audit_log (request_id, actor, action, detail) VALUES (?, ?, ?, ?)",
                (request_id, actor, "status_change", f"-> {status}"),
            )
        return cur.rowcount > 0


def flag_legal_review(request_id: int, note: str, actor: str = "system") -> bool:
    with _get_connection() as conn:
        cur = conn.execute(
            """UPDATE requests
                  SET needs_legal_review = 1, legal_note = ?, status = 'needs_legal',
                      updated_at = CURRENT_TIMESTAMP
                WHERE id = ?""",
            (note, request_id),
        )
        if cur.rowcount:
            conn.execute(
                "INSERT INTO audit_log (request_id, actor, action, detail) VALUES (?, ?, ?, ?)",
                (request_id, actor, "legal_flagged", note[:500] if note else ""),
            )
        return cur.rowcount > 0


def delete_request(request_id: int) -> bool:
    with _get_connection() as conn:
        cur = conn.execute("DELETE FROM requests WHERE id = ?", (request_id,))
    if cur.rowcount:
        _invalidate_past_answers_cache()
    return cur.rowcount > 0


# ------------------------------------------------------------------ Questions CRUD

def add_questions(request_id: int, items: List[Dict[str, Any]]) -> List[int]:
    """Bulk insert questions. `items` is a list of {question, section?, seq?} dicts.

    Returns the list of new question IDs in insertion order.
    """
    ids: List[int] = []
    with _get_connection() as conn:
        # Figure out next seq
        row = conn.execute(
            "SELECT COALESCE(MAX(seq), 0) AS max_seq FROM questions WHERE request_id = ?",
            (request_id,),
        ).fetchone()
        next_seq = (row["max_seq"] or 0) + 1
        for item in items:
            seq = item.get("seq") or next_seq
            next_seq = seq + 1
            cur = conn.execute(
                """INSERT INTO questions
                   (request_id, seq, section, question,
                    source_sheet, source_row, source_response_col)
                   VALUES (?, ?, ?, ?, ?, ?, ?)""",
                (request_id, seq, item.get("section"), item["question"],
                 item.get("source_sheet"), item.get("source_row"),
                 item.get("source_response_col")),
            )
            ids.append(cur.lastrowid)
    return ids


def get_question(question_id: int) -> Optional[Dict[str, Any]]:
    with _get_connection() as conn:
        row = conn.execute("SELECT * FROM questions WHERE id = ?", (question_id,)).fetchone()
        return dict(row) if row else None


def list_questions(request_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM questions WHERE request_id = ? ORDER BY seq ASC",
            (request_id,),
        ).fetchall()
        return [dict(r) for r in rows]


def save_draft(
    question_id: int,
    draft_answer: str,
    confidence: Optional[float] = None,
    citations: Optional[List[Dict[str, Any]]] = None,
) -> None:
    """Persist LLM draft + replace citations for this question."""
    with _get_connection() as conn:
        conn.execute(
            """UPDATE questions
                  SET draft_answer = ?, confidence = ?, status = 'drafted',
                      updated_at = CURRENT_TIMESTAMP
                WHERE id = ?""",
            (draft_answer, confidence, question_id),
        )
        conn.execute("DELETE FROM citations WHERE question_id = ?", (question_id,))
        for c in citations or []:
            conn.execute(
                """INSERT INTO citations (question_id, source_path, chunk_text, score)
                   VALUES (?, ?, ?, ?)""",
                (question_id, c.get("source_path", ""), c.get("chunk_text"), c.get("score")),
            )


def save_final_answer(question_id: int, final_answer: str, status: str = "approved") -> None:
    if status not in QUESTION_STATUSES:
        raise ValueError(f"Invalid question status: {status}")
    with _get_connection() as conn:
        conn.execute(
            """UPDATE questions
                  SET final_answer = ?, status = ?, updated_at = CURRENT_TIMESTAMP
                WHERE id = ?""",
            (final_answer, status, question_id),
        )
    # New/changed approved answer → past-answers cache must rebuild on next query.
    _invalidate_past_answers_cache()


def get_citations(question_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM citations WHERE question_id = ? ORDER BY score DESC NULLS LAST",
            (question_id,),
        ).fetchall()
        return [dict(r) for r in rows]


# ------------------------------------------------------------------ Uploads

def save_upload(
    request_id: int,
    file: FileStorage,
    kind: str = "inbound_questionnaire",
) -> Optional[Dict[str, Any]]:
    """Save a file under uploads/<request_id>/ and record it in the uploads table."""
    if not file or not file.filename:
        return None
    name = secure_filename(file.filename)
    if not name:
        return None
    ext = Path(name).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        logger.warning(f"[customer_assurance] rejected file type: {name}")
        return None
    content = file.read()
    if len(content) > MAX_FILE_SIZE:
        logger.warning(f"[customer_assurance] rejected oversized file: {name} ({len(content)} bytes)")
        return None

    sub_dir = UPLOADS_DIR / str(request_id)
    sub_dir.mkdir(parents=True, exist_ok=True)
    dest = sub_dir / name
    dest.write_bytes(content)

    with _get_connection() as conn:
        cur = conn.execute(
            """INSERT INTO uploads (request_id, filename, path, kind)
               VALUES (?, ?, ?, ?)""",
            (request_id, name, str(dest), kind),
        )
        upload_id = cur.lastrowid

    return {"id": upload_id, "filename": name, "path": str(dest), "kind": kind}


def list_uploads(request_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM uploads WHERE request_id = ? ORDER BY uploaded_at ASC",
            (request_id,),
        ).fetchall()
        return [dict(r) for r in rows]


# ------------------------------------------------------------------ Questionnaire extraction

# Header keywords used to locate the question / section columns inside an
# uploaded vendor security questionnaire .xlsx. Vendor templates vary widely
# (HCL VRA, DOL Cyber, the company TPRM, Farmers Group CRQ all look different) so
# detection is keyword-based, not positional. We use word-boundary regex so
# "question" doesn't substring-match "questionnaire" in banner/title rows.
import re as _re

_Q_HEADER_RE = _re.compile(
    r"\b(?:question text|assessment question|questions?)\b", _re.IGNORECASE
)
_SECTION_HEADER_RE = _re.compile(
    r"\b(?:sub[- ]?domain|control category|privacy domain|category|domain)\b",
    _re.IGNORECASE,
)
# Response-column keywords. Only used to back into the question column on
# sheets that lack a labeled question header (DOL CHECKLIST, Farmers Default
# Questions): the question is whatever's immediately to the left of the first
# response column.
_RESPONSE_HEADER_RE = _re.compile(
    r"\b(?:vendor response|vendor answers?|yes/no|responses?|answers?)\b",
    _re.IGNORECASE,
)


def _scan_for_xlsx_header(ws, max_scan: int = 10) -> Optional[Dict[str, Any]]:
    """Find the header row + question/section column indices for one sheet.

    Strategy:
      1. Scan rows 1..max_scan.
      2. If any cell's header text contains a question keyword, that's the
         question column.
      3. Else, if any cell contains a response keyword, the question column is
         the cell immediately to its left (vendors who skip a "Question" label
         still consistently put a "Response" column right after the question).
      4. The section column is any cell whose header contains a section
         keyword. Sometimes a sheet has a generic "Section" header that
         actually holds row numbers — we filter those out post-hoc by
         checking that section values look like text.

    Returns a dict {row, q_col, section_col, headers} or None.
    """
    rows = []
    try:
        for r_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=max_scan, values_only=True), start=1):
            cells = ["" if c is None else str(c).strip() for c in row]
            rows.append((r_idx, cells))
    except Exception:
        return None

    def _is_label_cell(c: str) -> bool:
        # Headers are short labels. Real keyword-bearing headers across all
        # four sample questionnaires (HCL VRA, DOL Cyber, the company TPRM,
        # Farmers CRQ) max out at ~32 chars / 4 words ("YES/NO       In
        # Process  Partial" on the DOL sheet). Tighten beyond that to keep
        # narrative cells like "Full assessment (High Default Questions &
        # CGRX)" — found inside risk-tier reference tables — from hijacking
        # detection just because they happen to contain the word "questions".
        if not c or len(c) > 35:
            return False
        return len(c.split()) <= 4

    def _find_section_col(cells):
        for i, c in enumerate(cells):
            if not _is_label_cell(c):
                continue
            if _SECTION_HEADER_RE.search(c):
                return i
        return None

    # Real header rows have multiple labeled columns. Skip banner / title rows
    # like "Default Questions (Low IR)" sitting alone in row 1 — they hijack
    # Pass 1 because they happen to contain a question keyword.
    def _looks_like_header_row(cells) -> bool:
        return sum(1 for c in cells if c) >= 3

    def _find_response_col(cells, after_col: int) -> Optional[int]:
        """Locate the response column (Vendor Response / Yes-No / Answers /
        Responses) somewhere to the right of `after_col`. Used by Pass 1 so
        round-trip export knows which cell to write the final answer into."""
        for i, c in enumerate(cells):
            if i <= after_col or not _is_label_cell(c):
                continue
            if _RESPONSE_HEADER_RE.search(c):
                return i
        return None

    # Pass 1: explicit question header
    for r_idx, cells in rows:
        if not _looks_like_header_row(cells):
            continue
        for i, c in enumerate(cells):
            if not _is_label_cell(c):
                continue
            if _Q_HEADER_RE.search(c):
                resp_col = _find_response_col(cells, after_col=i)
                # Fallback: most templates put the response right after the
                # question — q_col + 1 is a safe default if no labeled column
                # is found.
                if resp_col is None:
                    resp_col = i + 1
                return {
                    "row": r_idx,
                    "q_col": i,
                    "response_col": resp_col,
                    "section_col": _find_section_col(cells),
                    "headers": cells,
                }

    # Pass 2: derive question column from a response column header
    for r_idx, cells in rows:
        if not _looks_like_header_row(cells):
            continue
        resp_col = None
        for i, c in enumerate(cells):
            if not _is_label_cell(c):
                continue
            if _RESPONSE_HEADER_RE.search(c):
                resp_col = i
                break
        if resp_col is None or resp_col == 0:
            continue
        return {
            "row": r_idx,
            "q_col": resp_col - 1,
            "response_col": resp_col,
            "section_col": _find_section_col(cells),
            "headers": cells,
        }

    return None


def _is_questionable_text(text: str) -> bool:
    """Return True if `text` looks like a real question (not a number, junk, or yes/no)."""
    if not text:
        return False
    t = text.strip()
    if len(t) < 5:
        return False
    # Pure numbering / lettering tokens (e.g. "1)", "i.1", "1.2.3")
    if all(ch.isdigit() or ch in ".)(- " for ch in t):
        return False
    junk = {"yes", "no", "n/a", "na", "tbd", "true", "false", "pass", "gap", "the company"}
    if t.lower() in junk:
        return False
    return True


def extract_questions_from_xlsx(path: str) -> List[Dict[str, Any]]:
    """Extract questions from an uploaded vendor security questionnaire .xlsx.

    Returns a list of {section, question} dicts in spreadsheet order, ready to
    pass into `add_questions`. Sheets whose layout we can't decipher are
    skipped silently — the analyst can still paste raw text or edit afterward.

    Sheet section values carry forward across rows so questions under a merged
    section cell still get tagged correctly.
    """
    from openpyxl import load_workbook

    try:
        wb = load_workbook(path, data_only=True, read_only=True)
    except Exception as e:
        logger.warning(f"[customer_assurance] openpyxl failed on {path}: {e}")
        return []

    out: List[Dict[str, Any]] = []
    seen: set = set()

    try:
        for sheet_name in wb.sheetnames:
            try:
                ws = wb[sheet_name]
            except Exception:
                continue

            header = _scan_for_xlsx_header(ws)
            if not header:
                continue

            q_col = header["q_col"]
            section_col = header["section_col"]
            response_col = header["response_col"]
            start_row = header["row"] + 1
            last_section = None

            try:
                row_iter = ws.iter_rows(min_row=start_row, values_only=True)
            except Exception:
                continue

            for r_idx, row in enumerate(row_iter, start=start_row):
                cells = ["" if c is None else str(c).strip() for c in (row or [])]
                if q_col >= len(cells):
                    continue
                q_text = cells[q_col]
                if not _is_questionable_text(q_text):
                    continue

                section_val = None
                if section_col is not None and section_col < len(cells):
                    raw = cells[section_col]
                    # Reject section values that look like row numbers / bullets
                    if raw and not all(ch.isdigit() or ch in ".)(- " for ch in raw):
                        section_val = raw
                if section_val:
                    last_section = section_val
                composed_section = sheet_name
                if last_section:
                    composed_section = f"{sheet_name} / {last_section}"

                key = q_text.strip().lower()
                if key in seen:
                    continue
                seen.add(key)

                out.append({
                    "section": composed_section,
                    "question": q_text.strip(),
                    "source_sheet": sheet_name,
                    "source_row": r_idx,
                    "source_response_col": response_col,
                })
    finally:
        try:
            wb.close()
        except Exception:
            pass

    return out


# ------------------------------------------------------------------ Audit log

def log_audit(request_id: int, action: str, detail: str = "", actor: str = "system") -> None:
    with _get_connection() as conn:
        conn.execute(
            "INSERT INTO audit_log (request_id, actor, action, detail) VALUES (?, ?, ?, ?)",
            (request_id, actor, action, detail),
        )


def get_audit_log(request_id: int) -> List[Dict[str, Any]]:
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM audit_log WHERE request_id = ? ORDER BY at DESC",
            (request_id,),
        ).fetchall()
        return [dict(r) for r in rows]


# ------------------------------------------------------------------ KB retrieval

# Lazy singleton — Chroma client + collection are expensive to create.
_kb_state: Dict[str, Any] = {
    "client": None,
    "collection": None,
    "retriever": None,
    "embedding_fn": None,
    "bm25_docs": None,
}


def _get_kb_collection():
    """Return the Chroma collection for the customer_assurance_kb vector store."""
    if _kb_state["collection"] is not None:
        return _kb_state["collection"]
    import chromadb
    client = chromadb.PersistentClient(path=str(CHROMA_KB_PATH))
    collection = client.get_or_create_collection(
        name=KB_COLLECTION_NAME,
        metadata={"description": "Customer Assurance KB: policies, SOC2, standard responses"},
    )
    _kb_state["client"] = client
    _kb_state["collection"] = collection
    return collection


def _get_embedding_fn():
    if _kb_state["embedding_fn"] is None:
        from my_bot.document.document_processor import OllamaEmbeddingFunction
        _kb_state["embedding_fn"] = OllamaEmbeddingFunction()
    return _kb_state["embedding_fn"]


# ------------------------------------------------------------------ Past-answer retrieval
#
# Analyst-approved (question, final_answer) pairs from prior requests are the
# highest-priority retrieval source: they've already been vetted by a human.
# We cache embeddings of the past questions in memory and cosine-sim against
# incoming questions. The cache is invalidated whenever an answer changes
# (save_final_answer / delete_request).

_past_answers_state: Dict[str, Any] = {
    "loaded": False,
    "entries": [],  # list of {question_id, request_id, seq, question, final_answer, customer_name, assigned_to, embedding}
}


def _invalidate_past_answers_cache() -> None:
    _past_answers_state["loaded"] = False
    _past_answers_state["entries"] = []


def _cosine(a: List[float], b: List[float]) -> float:
    dot = sum(x * y for x, y in zip(a, b))
    na = math.sqrt(sum(x * x for x in a))
    nb = math.sqrt(sum(x * x for x in b))
    if na == 0.0 or nb == 0.0:
        return 0.0
    return dot / (na * nb)


def _rebuild_past_answers_cache() -> None:
    """Load all approved (question, final_answer) pairs and embed the questions."""
    with _get_connection() as conn:
        rows = conn.execute("""
            SELECT q.id AS question_id, q.request_id, q.seq, q.question, q.final_answer,
                   r.customer_name, r.assigned_to
              FROM questions q
              JOIN requests r ON r.id = q.request_id
             WHERE q.final_answer IS NOT NULL AND TRIM(q.final_answer) != ''
        """).fetchall()

    if not rows:
        _past_answers_state["entries"] = []
        _past_answers_state["loaded"] = True
        return

    texts = [r["question"] for r in rows]
    try:
        embeddings = _get_embedding_fn()(texts)
    except Exception as e:
        logger.warning(f"[customer_assurance] past-answers embedding failed, cache empty: {e}")
        _past_answers_state["entries"] = []
        _past_answers_state["loaded"] = True
        return

    entries = []
    for row, emb in zip(rows, embeddings):
        entries.append({
            "question_id": row["question_id"],
            "request_id": row["request_id"],
            "seq": row["seq"],
            "question": row["question"],
            "final_answer": row["final_answer"],
            "customer_name": row["customer_name"],
            "assigned_to": row["assigned_to"],
            "embedding": emb,
        })
    _past_answers_state["entries"] = entries
    _past_answers_state["loaded"] = True
    logger.info(f"[customer_assurance] past-answers cache rebuilt: {len(entries)} pairs")


def get_past_answers(
    question_text: str,
    exclude_demo: bool = False,
    k: int = 3,
    exclude_question_id: Optional[int] = None,
) -> List[Dict[str, Any]]:
    """Return the top-k past approved answers most similar to `question_text`.

    Each entry is a dict with question_id, request_id, seq, question, final_answer,
    customer_name, assigned_to, and a float `score` in [0, 1]. Entries are sorted
    by score descending. Demo-seeded pairs (assigned_to='demo') can be filtered
    via `exclude_demo=True`.
    """
    if not _past_answers_state["loaded"]:
        _rebuild_past_answers_cache()

    entries = _past_answers_state["entries"]
    if exclude_question_id is not None:
        entries = [e for e in entries if e["question_id"] != exclude_question_id]
    if exclude_demo:
        entries = [e for e in entries if (e["assigned_to"] or "").lower() != "demo"]
    if not entries:
        return []

    try:
        query_emb = _get_embedding_fn()([question_text])[0]
    except Exception as e:
        logger.warning(f"[customer_assurance] past-answers query embedding failed: {e}")
        return []

    scored = [(_cosine(query_emb, e["embedding"]), e) for e in entries]
    scored.sort(key=lambda x: x[0], reverse=True)

    out = []
    for score, e in scored[:k]:
        out.append({
            "question_id": e["question_id"],
            "request_id": e["request_id"],
            "seq": e["seq"],
            "question": e["question"],
            "final_answer": e["final_answer"],
            "customer_name": e["customer_name"],
            "assigned_to": e["assigned_to"],
            "score": score,
        })
    return out


def get_kb_retriever(k: int = 5):
    """Return a LangChain retriever over the customer_assurance_kb collection.

    Uses ChromaRetriever from document_processor; if there are documents cached
    in the collection, wraps it in an EnsembleRetriever with BM25 (70/30).
    Returns None if the collection is empty.
    """
    collection = _get_kb_collection()
    try:
        count = collection.count()
    except Exception as e:
        logger.error(f"[customer_assurance] chroma count failed: {e}")
        return None
    if count == 0:
        logger.warning("[customer_assurance] KB collection is empty — retriever unavailable")
        return None

    from my_bot.document.document_processor import ChromaRetriever
    from langchain_core.documents import Document

    vector_retriever = ChromaRetriever(
        collection=collection,
        embedding_fn=_get_embedding_fn(),
        k=k,
    )

    # BM25 hybrid — load docs from chroma cache once
    try:
        if _kb_state["bm25_docs"] is None:
            result = collection.get()
            if result and result.get("documents"):
                metadatas = result.get("metadatas") or [{}] * len(result["documents"])
                _kb_state["bm25_docs"] = [
                    Document(page_content=doc, metadata=meta or {})
                    for doc, meta in zip(result["documents"], metadatas)
                ]
        if _kb_state["bm25_docs"]:
            from langchain_community.retrievers import BM25Retriever
            from langchain_classic.retrievers.ensemble import EnsembleRetriever
            bm25 = BM25Retriever.from_documents(_kb_state["bm25_docs"])
            bm25.k = k
            retriever = EnsembleRetriever(retrievers=[vector_retriever, bm25], weights=[0.7, 0.3])
            _kb_state["retriever"] = retriever
            return retriever
    except Exception as e:
        logger.warning(f"[customer_assurance] BM25 hybrid setup failed, vector-only: {e}")

    _kb_state["retriever"] = vector_retriever
    return vector_retriever


def _kb_vector_scores(query: str, k: int = 5) -> List[Dict[str, Any]]:
    """Direct Chroma vector query returning hits with cosine similarity scores.

    Used for the confidence gate in `draft_question`. Bypasses the LangChain
    EnsembleRetriever because RRF fusion loses the underlying distance signal,
    and we need that signal to decide whether retrieval is strong enough to
    draft at all. The ensemble path still runs after the gate for the actual
    document fetch (so BM25 hybrid recall is preserved).

    Returns list of {content, metadata, score} sorted by score descending.
    Empty list if the collection is empty or the query fails.
    """
    collection = _get_kb_collection()
    try:
        count = collection.count()
    except Exception as e:
        logger.warning(f"[customer_assurance] chroma count failed: {e}")
        return []
    if count == 0:
        return []

    try:
        query_emb = _get_embedding_fn()([query])[0]
    except Exception as e:
        logger.warning(f"[customer_assurance] kb score query embedding failed: {e}")
        return []

    try:
        results = collection.query(
            query_embeddings=[query_emb],
            n_results=min(k, count),
            include=["documents", "metadatas", "distances"],
        )
    except Exception as e:
        logger.warning(f"[customer_assurance] chroma query failed: {e}")
        return []

    hits: List[Dict[str, Any]] = []
    if results.get("ids") and results["ids"][0]:
        docs = results.get("documents") or [[]]
        metas = results.get("metadatas") or [[]]
        dists = results.get("distances") or [[]]
        for i in range(len(results["ids"][0])):
            dist = dists[0][i] if dists and dists[0] and i < len(dists[0]) else None
            # Chroma default is L2 on normalized vectors, which maps to a
            # distance in [0, 2]. Similarity = 1 - (distance / 2) keeps the
            # score in [0, 1]. For collections using cosine directly the
            # distance is already in [0, 2] so the formula is the same.
            score = None
            if dist is not None:
                score = max(0.0, min(1.0, 1.0 - (dist / 2.0)))
            hits.append({
                "content": docs[0][i] if docs and docs[0] and i < len(docs[0]) else "",
                "metadata": (metas[0][i] if metas and metas[0] and i < len(metas[0]) else {}) or {},
                "score": score,
            })
    return hits


def _kb_doc_id(content: str, source: str) -> str:
    return hashlib.md5(f"{source}:{content[:200]}".encode()).hexdigest()


# Structured-loader helpers (`load_word_structured`, `load_excel_structured`)
# and the shared reranker client (`rerank_documents`) live in
# `my_bot.document.document_processor` so every RAG pipeline in IR can use
# them. Imported at call sites in `ingest_kb_source` and `draft_question`.


def ingest_kb_source(reset: bool = False) -> Dict[str, Any]:
    """Index all files in KB_SOURCE_DIR into the customer_assurance_kb collection.

    Dispatch by extension:
      - .docx → `_load_word_structured`: paragraphs routed through the splitter;
        table rows stay atomic (one chunk per row with column headers embedded).
      - .xlsx → `_load_excel_structured`: one atomic chunk per row.
      - .pdf  → `PyPDFDirectoryLoader` (unchanged — PDF table extraction is a
        separate problem we haven't tackled).
      - .doc / .xls (legacy) → fall back to the LangChain unstructured loaders;
        python-docx / openpyxl don't read them.

    Returns a summary dict: {files, chunks_added, chunks_skipped, errors, by_type}.
    """
    from langchain_community.document_loaders import (
        PyPDFDirectoryLoader,
        UnstructuredExcelLoader,
        UnstructuredWordDocumentLoader,
    )
    from langchain_core.documents import Document as LCDocument
    try:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    except ImportError:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

    summary: Dict[str, Any] = {
        "files": 0,
        "chunks_added": 0,
        "chunks_skipped": 0,
        "errors": [],
        "by_type": {"paragraph": 0, "table_row": 0, "xlsx_row": 0, "pdf_page": 0, "legacy": 0},
    }

    if reset:
        import chromadb
        client_obj = chromadb.PersistentClient(path=str(CHROMA_KB_PATH))
        try:
            client_obj.delete_collection(KB_COLLECTION_NAME)
        except Exception:
            pass
        _kb_state.update({"client": None, "collection": None, "retriever": None, "bm25_docs": None})

    collection = _get_kb_collection()

    files = [
        f for f in os.listdir(KB_SOURCE_DIR)
        if os.path.isfile(KB_SOURCE_DIR / f)
        and Path(f).suffix.lower() in {".pdf", ".doc", ".docx", ".xlsx", ".xls"}
    ]
    if not files:
        logger.warning(f"[customer_assurance] no source files in {KB_SOURCE_DIR}")
        return summary

    summary["files"] = len(files)

    # Items that SHOULD NOT be split (table rows — they're already atomic units)
    atomic_items: List[Dict[str, Any]] = []
    # Documents that SHOULD go through the text splitter (narrative text)
    splittable_docs: List[LCDocument] = []

    pdf_loaded = False
    for fname in files:
        fpath = str(KB_SOURCE_DIR / fname)
        ext = Path(fname).suffix.lower()
        try:
            if ext == ".pdf":
                # Single directory load covers all PDFs at once.
                if not pdf_loaded:
                    splittable_docs.extend(PyPDFDirectoryLoader(str(KB_SOURCE_DIR)).load())
                    pdf_loaded = True
            elif ext == ".docx":
                from my_bot.document.document_processor import load_word_structured
                word_items = load_word_structured(fpath)
                for item in word_items:
                    if item["metadata"].get("element_type") == "table_row":
                        atomic_items.append(item)
                    else:
                        splittable_docs.append(LCDocument(
                            page_content=item["content"],
                            metadata=item["metadata"],
                        ))
            elif ext == ".doc":
                # Legacy binary .doc — python-docx won't read it, fall back.
                loaded = UnstructuredWordDocumentLoader(fpath).load()
                for d in loaded:
                    d.metadata = {**(d.metadata or {}), "source": fname, "element_type": "legacy_doc"}
                splittable_docs.extend(loaded)
                summary["by_type"]["legacy"] += len(loaded)
            elif ext == ".xlsx":
                from my_bot.document.document_processor import load_excel_structured
                atomic_items.extend(load_excel_structured(fpath))
            elif ext == ".xls":
                # Legacy binary .xls — openpyxl won't read it, fall back.
                loaded = UnstructuredExcelLoader(fpath).load()
                for d in loaded:
                    d.metadata = {**(d.metadata or {}), "source": fname, "element_type": "legacy_xls"}
                splittable_docs.extend(loaded)
                summary["by_type"]["legacy"] += len(loaded)
        except Exception as e:
            summary["errors"].append(f"{fname}: {e}")
            logger.error(f"[customer_assurance] failed to load {fname}: {e}")

    if not splittable_docs and not atomic_items:
        return summary

    # Split narrative content; table rows bypass the splitter so they stay
    # semantically atomic.
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1500,
        chunk_overlap=300,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
    )
    split_chunks = splitter.split_documents(splittable_docs) if splittable_docs else []

    # Unified list of {content, metadata} tuples for upsert
    unified: List[Dict[str, Any]] = []
    for chunk in split_chunks:
        meta = dict(chunk.metadata or {})
        etype = meta.get("element_type")
        if not etype:
            # PDF pages come through without element_type — tag them.
            etype = "pdf_page"
            meta["element_type"] = etype
        unified.append({"content": chunk.page_content, "metadata": meta})
    unified.extend(atomic_items)

    # Tally by type for the summary.
    for item in unified:
        etype = (item.get("metadata") or {}).get("element_type") or "unknown"
        if etype in summary["by_type"]:
            summary["by_type"][etype] += 1

    existing_ids: set = set()
    try:
        result = collection.get()
        existing_ids = set(result.get("ids") or [])
    except Exception as e:
        logger.warning(f"[customer_assurance] could not read existing ids: {e}")

    ids, contents, metadatas = [], [], []
    for item in unified:
        content = item["content"]
        metadata = item["metadata"]
        cid = _kb_doc_id(content, metadata.get("source", ""))
        if cid in existing_ids:
            summary["chunks_skipped"] += 1
            continue
        ids.append(cid)
        contents.append(content)
        # Chroma metadata values must be str/int/float/bool — coerce anything else.
        clean_meta = {
            k: (v if isinstance(v, (str, int, float, bool)) else str(v))
            for k, v in metadata.items()
            if v is not None
        }
        metadatas.append(clean_meta)

    if not ids:
        return summary

    try:
        embedding_fn = _get_embedding_fn()
        embeddings = embedding_fn(contents)
        collection.upsert(ids=ids, documents=contents, metadatas=metadatas, embeddings=embeddings)
        summary["chunks_added"] = len(ids)
        # Invalidate BM25 cache so next retriever build picks up new docs.
        _kb_state["bm25_docs"] = None
        _kb_state["retriever"] = None
    except Exception as e:
        summary["errors"].append(f"upsert: {e}")
        logger.error(f"[customer_assurance] upsert failed: {e}", exc_info=True)

    logger.info(f"[customer_assurance] KB ingest: {summary}")
    return summary


def kb_stats() -> Dict[str, Any]:
    """Return counts for the KB collection — used by the KB admin page."""
    collection = _get_kb_collection()
    try:
        count = collection.count()
    except Exception:
        count = 0
    source_files = sorted(
        f for f in os.listdir(KB_SOURCE_DIR)
        if os.path.isfile(KB_SOURCE_DIR / f)
    )
    return {
        "collection": KB_COLLECTION_NAME,
        "chunk_count": count,
        "source_files": source_files,
        "source_dir": str(KB_SOURCE_DIR),
    }


# ------------------------------------------------------------------ Drafting

def draft_question(question_id: int, force_demo: bool = False) -> Dict[str, Any]:
    """Generate a draft answer for a question.

    Retrieval priority:
      Tier 0 — Past approved answers. If a prior analyst-approved answer for a
        semantically similar question exists (cosine >= PAST_ANSWER_HIGH_THRESHOLD),
        reuse its text verbatim and skip the LLM. Medium matches are folded into
        the LLM prompt as a priority context block.
      Tier 1 — KB chunks via hybrid Chroma+BM25 retrieval, drafted by the LLM.
      Tier 2 — Canned demo drafts (KB empty / the LLM unreachable / force_demo).

    Returns the saved draft dict (answer, confidence, citations, source).
    """
    from src.components.web.customer_assurance_demo import generate_demo_draft

    question = get_question(question_id)
    if not question:
        raise ValueError(f"Question {question_id} not found")

    # Tier 0: past approved answers. In Live mode (force_demo=False) we still
    # exclude demo-seeded pairs so real drafts never inherit canned content.
    past_matches: List[Dict[str, Any]] = []
    if not force_demo:
        past_matches = get_past_answers(
            question["question"],
            exclude_demo=True,
            k=3,
            exclude_question_id=question_id,
        )

    top_past = past_matches[0] if past_matches else None

    if top_past and top_past["score"] >= PAST_ANSWER_HIGH_THRESHOLD:
        answer = top_past["final_answer"]
        citation = {
            "source_path": f"Past answer — {top_past['customer_name']} (Q{top_past['seq']})",
            "chunk_text": f"Q: {top_past['question']}\n\nA: {answer}",
            "score": top_past["score"],
        }
        save_draft(question_id, answer, confidence=0.95, citations=[citation])
        log_audit(
            question["request_id"],
            "drafted",
            f"q#{question['seq']} (past-answer reuse, sim={top_past['score']:.2f})",
        )
        return {
            "answer": answer,
            "confidence": 0.95,
            "citations": [citation],
            "source": "past_answer",
        }

    # Medium-confidence past answers become a priority block in the LLM prompt.
    past_context = [m for m in past_matches if m["score"] >= PAST_ANSWER_MED_THRESHOLD]

    # Side-channel Chroma query to recover a real retrieval-confidence signal.
    # LangChain's EnsembleRetriever loses the underlying distance through RRF
    # fusion, so we run a direct vector query here. Cheap (tens of ms). Used
    # for both the Tier 1 SME gate below and the confidence value saved on
    # successful LLM drafts so analysts see retrieval strength in the UI.
    kb_score_hits: List[Dict[str, Any]] = []
    top_kb_score: Optional[float] = None
    if not force_demo:
        kb_score_hits = _kb_vector_scores(question["question"], k=5)
        if kb_score_hits and kb_score_hits[0].get("score") is not None:
            top_kb_score = kb_score_hits[0]["score"]

    # Tier 1 confidence gate: KB has docs but none match strongly AND no past
    # answer rescued us → withhold the draft and route to SME. Protects external
    # customers from confidently-wrong drafts built on thin evidence. An empty
    # KB is a separate case handled further down by `retriever is None`.
    if (
        not force_demo
        and not past_context
        and kb_score_hits
        and top_kb_score is not None
        and top_kb_score < KB_RETRIEVAL_MIN_SCORE
    ):
        gate_msg = (
            "Automated draft withheld: the knowledge base does not contain "
            f"sufficiently relevant material (top match similarity "
            f"{top_kb_score:.2f}, below the {KB_RETRIEVAL_MIN_SCORE:.2f} "
            "threshold). This question has been routed to a subject-matter "
            "expert for a manual answer."
        )
        gate_citations = [
            {
                "source_path": h["metadata"].get("source", f"weak_hit_{i+1}"),
                "chunk_text": (h.get("content") or "")[:500],
                "score": h.get("score"),
            }
            for i, h in enumerate(kb_score_hits[:3])
        ]
        save_draft(
            question_id,
            draft_answer=gate_msg,
            confidence=top_kb_score,
            citations=gate_citations,
        )
        with _get_connection() as conn:
            conn.execute(
                "UPDATE questions SET status = 'needs_sme' WHERE id = ?",
                (question_id,),
            )
        log_audit(
            question["request_id"],
            "drafted",
            f"q#{question['seq']} (needs_sme — weak KB, top={top_kb_score:.2f})",
        )
        return {
            "answer": gate_msg,
            "confidence": top_kb_score,
            "citations": gate_citations,
            "source": "needs_sme_gate",
        }

    # Retrieve wide (k=20) so the reranker has real candidates to re-score;
    # the downstream `rerank_documents` call will precision-filter to top_k=5.
    retriever = None if force_demo else get_kb_retriever(k=20)

    if retriever is None:
        result = generate_demo_draft(question["question"])
        save_draft(
            question_id,
            draft_answer=result["answer"],
            confidence=result["confidence"],
            citations=result["citations"],
        )
        if result.get("needs_sme"):
            with _get_connection() as conn:
                conn.execute(
                    "UPDATE questions SET status = 'needs_sme' WHERE id = ?",
                    (question_id,),
                )
        log_audit(question["request_id"], "drafted",
                  f"q#{question['seq']} (demo mode)")
        return result

    # Real path — retrieve from KB + rerank + draft via the LLM
    try:
        docs = retriever.invoke(question["question"])
    except Exception as e:
        logger.error(f"[customer_assurance] retriever failed: {e}")
        result = generate_demo_draft(question["question"])
        save_draft(question_id, result["answer"], result["confidence"], result["citations"])
        log_audit(question["request_id"], "drafted",
                  f"q#{question['seq']} (demo fallback — retriever error)")
        return result

    # Cross-encoder rerank: trades wide hybrid recall (k=20) for precision in
    # the 5 chunks the LLM actually sees. Gracefully returns input unchanged if
    # the reranker endpoint is down.
    try:
        from my_bot.document.document_processor import rerank_documents
        docs = rerank_documents(question["question"], docs, top_k=5)
    except Exception as e:
        logger.warning(f"[customer_assurance] rerank failed, using raw retrieval: {e}")
        docs = docs[:5]

    try:
        from my_bot.utils.llm_factory import create_llm
        from langchain_core.messages import HumanMessage
    except Exception as e:
        logger.warning(f"[customer_assurance] the LLM client unavailable, using demo: {e}")
        result = generate_demo_draft(question["question"])
        save_draft(question_id, result["answer"], result["confidence"], result["citations"])
        log_audit(question["request_id"], "drafted",
                  f"q#{question['seq']} (demo fallback — the LLM unavailable)")
        return result

    context_blocks = []
    citation_rows = []
    for i, doc in enumerate(docs[:5]):
        src = doc.metadata.get("source", f"doc{i}")
        context_blocks.append(f"[{i+1}] Source: {src}\n{doc.page_content}")
        citation_rows.append({
            "source_path": src,
            "chunk_text": doc.page_content[:500],
            # Prefer the cross-encoder rerank score (calibrated [0,1] relevance
            # probability) over the upstream retriever score. LangChain's
            # EnsembleRetriever returns None here, so reranking is what
            # populates the per-chunk score column in the workspace UI.
            "score": doc.metadata.get("rerank_score") or doc.metadata.get("score"),
        })

    # Surface medium-confidence past answers in the citations panel so analysts
    # can see which prior approved answers informed the draft.
    for pa in past_context:
        citation_rows.insert(0, {
            "source_path": f"Past answer — {pa['customer_name']} (Q{pa['seq']})",
            "chunk_text": f"Q: {pa['question']}\n\nA: {pa['final_answer']}",
            "score": pa["score"],
        })

    prompt = _build_drafting_prompt(
        question=question["question"],
        section=question.get("section") or "",
        context="\n\n".join(context_blocks),
        past_answers=past_context,
    )

    try:
        llm = create_llm(timeout=90)
        resp = llm.invoke([HumanMessage(content=prompt)])
        answer = (resp.content or "").strip()
    except Exception as e:
        logger.error(f"[customer_assurance] the LLM chat failed: {e}")
        result = generate_demo_draft(question["question"])
        save_draft(question_id, result["answer"], result["confidence"], result["citations"])
        log_audit(question["request_id"], "drafted",
                  f"q#{question['seq']} (demo fallback — the LLM error)")
        return result

    if not answer:
        result = generate_demo_draft(question["question"])
        save_draft(question_id, result["answer"], result["confidence"], result["citations"])
        log_audit(question["request_id"], "drafted",
                  f"q#{question['seq']} (demo fallback — empty the LLM response)")
        return result

    source_tag = "llm+past_context" if past_context else "llm"
    # Pick whichever retrieval signal is strongest for the analyst-facing
    # confidence percentage: a past-answer match beats a KB match beats nothing.
    effective_confidence: Optional[float] = None
    if past_context:
        effective_confidence = past_context[0]["score"]
    elif top_kb_score is not None:
        effective_confidence = top_kb_score
    save_draft(question_id, answer, confidence=effective_confidence, citations=citation_rows)
    log_audit(
        question["request_id"],
        "drafted",
        f"q#{question['seq']} ({source_tag})"
        + (f" sim={past_context[0]['score']:.2f}" if past_context else "")
        + (f" kb={top_kb_score:.2f}" if top_kb_score is not None else ""),
    )
    return {
        "answer": answer,
        "confidence": effective_confidence,
        "citations": citation_rows,
        "source": source_tag,
    }


def _build_drafting_prompt(
    question: str,
    section: str,
    context: str,
    past_answers: Optional[List[Dict[str, Any]]] = None,
) -> str:
    past_block = ""
    if past_answers:
        lines = [
            "PAST APPROVED ANSWERS (highest priority — already reviewed and approved "
            "by an analyst for a semantically similar past question):",
            "",
        ]
        for i, pa in enumerate(past_answers, 1):
            lines.append(f"[P{i}] From {pa['customer_name']} (similarity {pa['score']:.2f}):")
            lines.append(f"    Past question: {pa['question']}")
            lines.append(f"    Approved answer: {pa['final_answer']}")
            lines.append("")
        past_block = "\n".join(lines) + "\n"

    past_rule = (
        "- If PAST APPROVED ANSWERS are provided and address the question, prefer "
        "their wording. Adapt only where the current question asks something different.\n"
        if past_answers
        else ""
    )

    return f"""You are drafting a response to a customer security questionnaire on behalf of \
the company's Customer Assurance team. Your answer will be reviewed by an analyst \
before release.

RULES:
- Use ONLY the provided context and past approved answers. Do not invent facts.
{past_rule}- If the context does not contain enough information, say so clearly and \
recommend routing to an SME.
- Keep the answer factual, professional, and 2–5 sentences.
- If the question touches contracts, breach notification, or legal commitments, \
add a line recommending Legal review.

WRONG — never do this:
  "We have industry-leading security and are fully compliant with all standards."
  (vague, unsupported, no grounding)

RIGHT:
  "Yes. Customer data at rest is encrypted with AES-256. Keys are managed via \
our enterprise KMS with annual rotation [Source: InfoSec-Std v4.2]."

{past_block}Section: {section}
Question: {question}

Context:
{context}

Draft answer:"""


# ------------------------------------------------------------------ Export

def export_request_xlsx(request_id: int) -> Optional[Path]:
    """Round-trip export: write final answers back into the customer's
    original .xlsx, in the same row + response column the question was
    extracted from. Returns the saved path, or None if round-trip isn't
    possible (no source coords on questions, or no inbound .xlsx upload).

    Empty cells are skipped — if a question has no final/draft answer, the
    customer's original cell is left as-is rather than being overwritten with
    "[No response drafted]".
    """
    from openpyxl import load_workbook
    from datetime import datetime
    import shutil

    req = get_request(request_id)
    if not req:
        return None

    questions = list_questions(request_id)
    coord_qs = [q for q in questions
                if q.get("source_sheet") and q.get("source_row")
                and q.get("source_response_col") is not None]
    if not coord_qs:
        return None

    # Find the original inbound xlsx upload.
    inbound = [u for u in list_uploads(request_id)
               if (u.get("kind") == "inbound_questionnaire"
                   and u.get("filename", "").lower().endswith(".xlsx"))]
    if not inbound:
        return None
    src_path = Path(inbound[-1]["path"])
    if not src_path.exists():
        return None

    exports_dir = UPLOADS_DIR / str(request_id)
    exports_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = req["customer_name"].replace(" ", "_").replace("/", "_")
    out_path = exports_dir / f"response_{safe_name}_{stamp}.xlsx"

    try:
        shutil.copy2(src_path, out_path)
        wb = load_workbook(out_path)
    except Exception as e:
        logger.error(f"[customer_assurance] xlsx export copy/open failed: {e}", exc_info=True)
        try:
            out_path.unlink(missing_ok=True)
        except Exception:
            pass
        return None

    written = 0
    skipped_no_answer = 0
    skipped_missing_sheet = 0
    try:
        for q in coord_qs:
            sheet = q["source_sheet"]
            row = q["source_row"]
            # Stored as 0-indexed (the position inside the cells list);
            # openpyxl cell coordinates are 1-indexed.
            resp_col = q["source_response_col"] + 1
            answer = (q.get("final_answer") or q.get("draft_answer") or "").strip()
            if not answer:
                skipped_no_answer += 1
                continue
            if sheet not in wb.sheetnames:
                skipped_missing_sheet += 1
                continue
            try:
                wb[sheet].cell(row=row, column=resp_col, value=answer)
                written += 1
            except Exception as e:
                logger.warning(f"[customer_assurance] failed to write {sheet}!R{row}C{resp_col}: {e}")

        wb.save(str(out_path))
    finally:
        try:
            wb.close()
        except Exception:
            pass

    with _get_connection() as conn:
        conn.execute(
            """INSERT INTO uploads (request_id, filename, path, kind)
               VALUES (?, ?, ?, ?)""",
            (request_id, out_path.name, str(out_path), "response_export"),
        )
    detail = (f"{out_path.name} (round-trip xlsx, {written} answer(s) written"
              + (f", {skipped_no_answer} skipped no-answer" if skipped_no_answer else "")
              + (f", {skipped_missing_sheet} skipped missing-sheet" if skipped_missing_sheet else "")
              + ")")
    log_audit(request_id, "exported", detail)
    return out_path


def export_request_docx(request_id: int) -> Optional[Path]:
    """Export approved/drafted answers to a .docx file. Returns the path."""
    try:
        from docx import Document
    except ImportError:
        logger.error("[customer_assurance] python-docx not available")
        return None

    req = get_request(request_id)
    if not req:
        return None
    questions = list_questions(request_id)

    doc = Document()
    doc.add_heading(f"Security Questionnaire Response — {req['customer_name']}", level=0)
    doc.add_paragraph(f"Request: {req['title']}")
    doc.add_paragraph(f"Type: {req['request_type']} | Segment: {req['customer_segment']}")
    doc.add_paragraph(f"Prepared by: the company Customer Assurance Team")
    doc.add_paragraph("")

    current_section = None
    for q in questions:
        if q.get("section") and q["section"] != current_section:
            current_section = q["section"]
            doc.add_heading(current_section, level=1)
        doc.add_heading(f"Q{q['seq']}. {q['question']}", level=2)
        answer = q.get("final_answer") or q.get("draft_answer") or "[No response drafted]"
        doc.add_paragraph(answer)
        if q.get("status") == "needs_sme":
            p = doc.add_paragraph()
            run = p.add_run("⚠ Pending SME review before release")
            run.italic = True
        doc.add_paragraph("")

    exports_dir = UPLOADS_DIR / str(request_id)
    exports_dir.mkdir(parents=True, exist_ok=True)
    from datetime import datetime
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    path = exports_dir / f"response_{req['customer_name'].replace(' ', '_')}_{stamp}.docx"
    doc.save(str(path))

    with _get_connection() as conn:
        conn.execute(
            """INSERT INTO uploads (request_id, filename, path, kind)
               VALUES (?, ?, ?, ?)""",
            (request_id, path.name, str(path), "response_export"),
        )
    log_audit(request_id, "exported", path.name)
    return path


# ------------------------------------------------------------------ Demo seeding

def seed_demo_data(wipe: bool = False) -> Dict[str, Any]:
    """Populate the DB with the sample requests from customer_assurance_demo.

    Idempotent: if wipe=True, deletes all existing demo data first. Otherwise
    skips any sample whose (customer_name, title) already exists.
    """
    from src.components.web.customer_assurance_demo import SAMPLE_REQUESTS

    if wipe:
        with _get_connection() as conn:
            conn.execute("DELETE FROM requests")
            conn.execute("DELETE FROM sqlite_sequence WHERE name='requests'")

    created = 0
    skipped = 0
    with _get_connection() as conn:
        for sample in SAMPLE_REQUESTS:
            existing = conn.execute(
                "SELECT id FROM requests WHERE customer_name = ? AND title = ?",
                (sample["customer_name"], sample["title"]),
            ).fetchone()
            if existing:
                skipped += 1
                continue
            cur = conn.execute(
                """INSERT INTO requests
                   (customer_name, customer_segment, account_team_contact, request_type,
                    source_format, due_date, priority, title, raw_text, notes,
                    assigned_to, status, needs_legal_review, legal_note)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    sample["customer_name"], sample["customer_segment"],
                    sample.get("account_team_contact"), sample["request_type"],
                    sample.get("source_format"), sample.get("due_date"),
                    sample.get("priority"), sample["title"],
                    sample.get("raw_text"), sample.get("notes"),
                    sample.get("assigned_to"), sample.get("status", "new"),
                    sample.get("needs_legal_review", 0), sample.get("legal_note"),
                ),
            )
            rid = cur.lastrowid
            for seq, q in enumerate(sample.get("questions", []), start=1):
                conn.execute(
                    """INSERT INTO questions (request_id, seq, section, question)
                       VALUES (?, ?, ?, ?)""",
                    (rid, seq, q.get("section"), q["question"]),
                )
            conn.execute(
                "INSERT INTO audit_log (request_id, actor, action, detail) VALUES (?, ?, ?, ?)",
                (rid, "demo", "created", f"Seeded demo request: {sample['title']}"),
            )
            created += 1

    # After seeding, auto-draft any 'drafting' or 'ready' or 'needs_legal' requests so the
    # workspace has content when business users open it.
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT id, status FROM requests WHERE status IN ('drafting','ready','needs_legal')"
        ).fetchall()
    drafted = 0
    for row in rows:
        qs = list_questions(row["id"])
        for q in qs:
            if not q.get("draft_answer"):
                try:
                    draft_question(q["id"], force_demo=True)
                    drafted += 1
                except Exception as e:
                    logger.warning(f"[customer_assurance] demo auto-draft failed q#{q['id']}: {e}")

    # Mark the 'ready' request's questions as approved so the export flow shows final answers
    with _get_connection() as conn:
        ready = conn.execute("SELECT id FROM requests WHERE status = 'ready'").fetchall()
        for r in ready:
            conn.execute(
                """UPDATE questions
                      SET final_answer = COALESCE(final_answer, draft_answer),
                          status = 'approved'
                    WHERE request_id = ? AND draft_answer IS NOT NULL""",
                (r["id"],),
            )

    return {"created": created, "skipped": skipped, "auto_drafted": drafted}
