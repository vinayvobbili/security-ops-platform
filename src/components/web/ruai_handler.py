"""RUAI Screening Review System Handler.

Backend for the Responsible Use of AI screening form, LLM-powered first-pass
review, and reviewer dashboard.
"""

import json
import logging
import sqlite3
import uuid
from concurrent.futures import ThreadPoolExecutor, TimeoutError as FuturesTimeoutError
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

import openpyxl
from pydantic import BaseModel, Field
from werkzeug.datastructures import FileStorage
from werkzeug.utils import secure_filename
from webexpythonsdk import WebexAPI

from my_config import get_config
from src.utils.webex_messaging import safe_send_message

try:
    import docx as _docx_mod
except ImportError:
    _docx_mod = None

try:
    import pdfplumber as _pdfplumber_mod
except ImportError:
    _pdfplumber_mod = None

logger = logging.getLogger(__name__)
CONFIG = get_config()

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
DATA_DIR = Path(__file__).resolve().parent.parent.parent.parent / "data" / "ruai_screening"
DB_PATH = DATA_DIR / "ruai_screening.db"
UPLOADS_DIR = Path(__file__).resolve().parent.parent.parent.parent / "data" / "transient" / "ruai_uploads"
DATA_DIR.mkdir(parents=True, exist_ok=True)
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXTENSIONS = {
    '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.pptx',
    '.txt', '.csv', '.png', '.jpg', '.jpeg', '.gif', '.svg', '.vsdx',
}

LLM_TIMEOUT_SECONDS = 300

# In-memory review phase tracker (submission_id → phase string)
# Phases: queued → loading_submission → connecting_llm → analyzing → saving_results → done | error
_review_phases: dict[int, str] = {}
_review_stats: dict[int, dict] = {}  # submission_id → {duration_s, input_tokens, output_tokens, model}

# ---------------------------------------------------------------------------
# Database
# ---------------------------------------------------------------------------

@contextmanager
def _get_connection():
    """Context manager for SQLite connections with auto-commit/rollback."""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


def init_db():
    """Create tables if they don't exist."""
    with _get_connection() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS submissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                use_case_name TEXT NOT NULL,
                use_case_id TEXT,
                use_case_description TEXT NOT NULL,
                use_case_stage TEXT NOT NULL,
                owner_name TEXT NOT NULL,
                owner_email TEXT NOT NULL,
                lob TEXT NOT NULL,
                region TEXT,
                vendor_or_inhouse TEXT NOT NULL,
                vendor_name TEXT,
                responses JSON NOT NULL,
                documents TEXT,
                status TEXT NOT NULL DEFAULT 'submitted',
                submitted_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS ai_reviews (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                submission_id INTEGER NOT NULL REFERENCES submissions(id),
                overall_risk_score TEXT,
                completeness_issues TEXT,
                risk_flags TEXT,
                threat_boundary_analysis TEXT,
                clarifying_questions TEXT,
                preliminary_risk_assessment TEXT,
                review_summary_md TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        # Migrations: add columns if upgrading from earlier schema
        ai_cols = [row[1] for row in conn.execute("PRAGMA table_info(ai_reviews)").fetchall()]
        if 'overall_risk_score' not in ai_cols:
            conn.execute("ALTER TABLE ai_reviews ADD COLUMN overall_risk_score TEXT")
        if 'ai_threat_surface' not in ai_cols:
            conn.execute("ALTER TABLE ai_reviews ADD COLUMN ai_threat_surface TEXT")
        if 'control_coverage_map' not in ai_cols:
            conn.execute("ALTER TABLE ai_reviews ADD COLUMN control_coverage_map TEXT")
        if 'review_meta' not in ai_cols:
            conn.execute("ALTER TABLE ai_reviews ADD COLUMN review_meta TEXT")
        conn.execute("""
            CREATE TABLE IF NOT EXISTS reviewer_actions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                submission_id INTEGER NOT NULL REFERENCES submissions(id),
                reviewer_name TEXT NOT NULL,
                action TEXT NOT NULL,
                notes TEXT,
                created_at DATETIME DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS reviewer_checklists (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                submission_id INTEGER NOT NULL REFERENCES submissions(id),
                reviewer_name TEXT NOT NULL,
                checklist_items JSON NOT NULL DEFAULT '{}',
                updated_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                UNIQUE(submission_id, reviewer_name)
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS reviewer_assignments (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                submission_id INTEGER NOT NULL REFERENCES submissions(id),
                reviewer_name TEXT NOT NULL,
                assigned_by TEXT NOT NULL,
                assigned_at DATETIME DEFAULT CURRENT_TIMESTAMP,
                signed_off INTEGER NOT NULL DEFAULT 0,
                signed_off_at DATETIME,
                UNIQUE(submission_id, reviewer_name)
            )
        """)
    logger.info(f"RUAI screening database initialized at {DB_PATH}")


# Auto-init on import
init_db()


# ---------------------------------------------------------------------------
# Screening form configuration (rules engine)
# ---------------------------------------------------------------------------

SCREENING_SECTIONS = [
    {
        "id": "use_case_info",
        "title": "Use Case Information",
        "always_show": True,
        "fields": [
            {"name": "use_case_name", "label": "Use Case Name", "type": "text", "required": True},
            {"name": "use_case_description", "label": "Use Case Description", "type": "textarea", "required": True,
             "placeholder": "Describe the AI use case, its goals, and expected behavior"},
            {"name": "use_case_stage", "label": "Use Case Stage", "type": "select", "required": True,
             "options": ["PoC", "Pilot", "Implementation"]},
            {"name": "owner_name", "label": "Business Owner", "type": "text", "required": True},
            {"name": "owner_email", "label": "Owner Email", "type": "email", "required": True},
            {"name": "lob", "label": "LOB / Function", "type": "text", "required": True,
             "placeholder": "e.g. Claims, Underwriting, IT, HR"},
            {"name": "region", "label": "Region / Country", "type": "text",
             "placeholder": "e.g. US, EMEA, APAC"},
        ],
    },
    {
        "id": "prohibited_uses",
        "title": "Prohibited Use Cases",
        "always_show": True,
        "fields": [
            {"name": "eu_uk_based", "label": "Is the AI system going to be based in or used within the EU or UK?",
             "type": "yesno", "required": True},
            {"name": "prohibited_techniques", "label": "Does the AI exhibit any of the following prohibited techniques?",
             "type": "checklist", "required": True,
             "options": [
                 "Subliminal/manipulative techniques",
                 "Exploiting group vulnerabilities",
                 "Biometric categorization for sensitive characteristics",
                 "Social scoring",
                 "Untargeted facial image scraping",
                 "Real-time biometric ID in public spaces",
                 "Emotion recognition in workplace",
                 "Predicting crime risk",
                 "None of the above",
             ]},
        ],
    },
    {
        "id": "business_purpose",
        "title": "AI Solution and Business Purpose",
        "always_show": True,
        "fields": [
            {"name": "what_developing", "label": "What are you developing or deploying?", "type": "textarea", "required": True},
            {"name": "how_changes_things", "label": "How will this change existing processes or outcomes?", "type": "textarea", "required": True},
            {"name": "business_need", "label": "What is the business need or problem this addresses?", "type": "textarea", "required": True},
            {"name": "key_features", "label": "Key features of the AI solution", "type": "textarea", "required": True},
            {"name": "intended_users", "label": "Who are the intended users?", "type": "textarea", "required": True},
        ],
    },
    {
        "id": "decision_making",
        "title": "Use of AI in Decision-Making",
        "always_show": True,
        "fields": [
            {"name": "ai_drives_decisions", "label": "Does your solution use AI to drive decisions or outcomes?",
             "type": "yesno", "required": True},
            {"name": "ai_techniques", "label": "What AI techniques and methods are used?", "type": "textarea",
             "depends_on": {"ai_drives_decisions": "yes"},
             "placeholder": "e.g. ML classification, NLP, GenAI, rule-based engine"},
            {"name": "decision_making_role", "label": "Describe the role of AI in decision-making", "type": "textarea",
             "depends_on": {"ai_drives_decisions": "yes"},
             "placeholder": "How are decisions made? Fully automated, human-in-the-loop, advisory only?"},
        ],
    },
    {
        "id": "critical_decisions",
        "title": "Critical Decision Making",
        "depends_on": {"ai_drives_decisions": "yes"},
        "fields": [
            {"name": "critical_areas", "label": "Which critical areas does the AI impact?",
             "type": "checklist", "required": True,
             "options": [
                 "Claims processing/adjudication",
                 "Underwriting decisions",
                 "Pricing/rating",
                 "Hiring/employment decisions",
                 "Fraud detection",
                 "Biometric identification",
                 "Customer service routing/decisions",
                 "Marketing/targeting",
                 "Investment decisions",
                 "None of the above",
             ]},
            {"name": "critical_decisions_detail", "label": "Describe how AI is used in these critical areas", "type": "textarea"},
        ],
    },
    {
        "id": "customer_impact",
        "title": "Impact on Customers",
        "depends_on": {"ai_drives_decisions": "yes"},
        "fields": [
            {"name": "affects_customers", "label": "Does the AI output directly affect customers?",
             "type": "yesno", "required": True},
            {"name": "financial_impact", "label": "Could it impact customers financially (premiums, claims, coverage)?",
             "type": "yesno", "depends_on": {"affects_customers": "yes"}},
            {"name": "insurance_impact", "label": "Could it impact insurance availability or terms?",
             "type": "yesno", "depends_on": {"affects_customers": "yes"}},
            {"name": "customer_impact_detail", "label": "Describe the potential impact on customers", "type": "textarea",
             "depends_on": {"affects_customers": "yes"}},
        ],
    },
    {
        "id": "nontraditional_factors",
        "title": "Non-Traditional or Suspect Factors",
        "depends_on": {"ai_drives_decisions": "yes"},
        "fields": [
            {"name": "uses_lifestyle_data", "label": "Does the AI use lifestyle or behavioral data?",
             "type": "yesno", "required": True},
            {"name": "lifestyle_data_detail", "label": "What lifestyle/behavioral data is used and how?",
             "type": "textarea", "depends_on": {"uses_lifestyle_data": "yes"}},
            {"name": "infers_sensitive_info", "label": "Could the AI infer sensitive information (race, religion, health, etc.)?",
             "type": "yesno", "required": True},
            {"name": "sensitive_info_detail", "label": "Describe what sensitive information could be inferred",
             "type": "textarea", "depends_on": {"infers_sensitive_info": "yes"}},
            {"name": "lifestyle_categorization", "label": "Does the AI categorize individuals based on lifestyle factors?",
             "type": "yesno", "required": True},
        ],
    },
    {
        "id": "human_oversight",
        "title": "Human Oversight",
        "always_show": True,
        "fields": [
            {"name": "human_review_types", "label": "What types of AI decisions are subject to human review?", "type": "textarea", "required": True},
            {"name": "who_reviews", "label": "Who reviews AI-generated decisions?", "type": "textarea", "required": True},
            {"name": "review_criteria", "label": "What criteria are used for human review?", "type": "textarea"},
            {"name": "can_override", "label": "Can humans override AI decisions?", "type": "yesno", "required": True},
            {"name": "override_process", "label": "Describe the override process", "type": "textarea",
             "depends_on": {"can_override": "yes"}},
        ],
    },
    {
        "id": "vendor_info",
        "title": "Vendor / In-house",
        "always_show": True,
        "fields": [
            {"name": "vendor_or_inhouse", "label": "Development approach", "type": "select", "required": True,
             "options": ["In-house", "Third Party Vendor", "Not yet decided"]},
            {"name": "vendor_name", "label": "Vendor name", "type": "text",
             "depends_on": {"vendor_or_inhouse": "Third Party Vendor"}},
            {"name": "vendor_reason", "label": "Reason for selecting this vendor", "type": "textarea",
             "depends_on": {"vendor_or_inhouse": "Third Party Vendor"}},
            {"name": "vendor_ai_transparency", "label": "Has the vendor provided AI transparency documentation?",
             "type": "yesno", "depends_on": {"vendor_or_inhouse": "Third Party Vendor"}},
        ],
    },
    {
        "id": "data_governance",
        "title": "Data Governance",
        "always_show": True,
        "fields": [
            {"name": "processes_personal_data", "label": "Does this tool process personal data?",
             "type": "yesno", "required": True},
            {"name": "data_classification", "label": "Is there a data classification process in place?",
             "type": "textarea", "depends_on": {"processes_personal_data": "yes"}},
            {"name": "data_retention", "label": "What is the data retention policy?",
             "type": "textarea", "depends_on": {"processes_personal_data": "yes"}},
            {"name": "data_access_controls", "label": "What access controls are in place for personal data?",
             "type": "textarea", "depends_on": {"processes_personal_data": "yes"}},
            {"name": "ingests_the company_data", "label": "Does this tool ingest, create, update, or delete company data?",
             "type": "yesno", "required": True},
            {"name": "data_domains", "label": "Which data domains are relevant?",
             "type": "textarea", "depends_on": {"ingests_the company_data": "yes"},
             "placeholder": "e.g. Customer, Claims, Policy, Financial, HR"},
            {"name": "data_lineage", "label": "Is data lineage tracked?",
             "type": "yesno", "depends_on": {"ingests_the company_data": "yes"}},
        ],
    },
    {
        "id": "additional_info",
        "title": "Additional Information",
        "always_show": True,
        "fields": [
            {"name": "pra_status", "label": "PRA (Privacy Risk Assessment) status", "type": "select",
             "options": ["Not started", "In progress", "Completed", "Not applicable"]},
            {"name": "garb_checklist", "label": "Has the GARB checklist been completed?", "type": "yesno"},
            {"name": "feedback_mechanisms", "label": "What feedback mechanisms are in place for end users?", "type": "textarea"},
            {"name": "training_docs", "label": "Is training/documentation available for the AI tool?", "type": "yesno"},
            {"name": "additional_notes", "label": "Any additional information or context", "type": "textarea"},
        ],
    },
]


def get_form_config() -> List[Dict[str, Any]]:
    """Return the screening form configuration for the frontend."""
    return SCREENING_SECTIONS


# ---------------------------------------------------------------------------
# File uploads
# ---------------------------------------------------------------------------

def _save_uploads(files: List[FileStorage], use_case_name: str) -> List[str]:
    """Save uploaded files to a named subdirectory. Returns list of saved filenames."""
    if not files:
        return []

    folder_name = secure_filename(use_case_name) or "unnamed"
    sub_dir = UPLOADS_DIR / folder_name
    sub_dir.mkdir(parents=True, exist_ok=True)

    saved = []
    for f in files:
        if not f or not f.filename:
            continue
        name = secure_filename(f.filename)
        if not name:
            continue
        ext = Path(name).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            logger.warning(f"RUAI upload: skipped disallowed file type: {name}")
            continue
        content = f.read()
        if len(content) > MAX_FILE_SIZE:
            logger.warning(f"RUAI upload: skipped oversized file: {name} ({len(content)} bytes)")
            continue
        dest = sub_dir / name
        dest.write_bytes(content)
        saved.append(name)
        logger.debug(f"RUAI upload saved: {dest}")

    return saved


# ---------------------------------------------------------------------------
# XLSX survey parsing & document text extraction (upload flow)
# ---------------------------------------------------------------------------

# Map XLSX question text (lowercase prefix) → DB structured field name
_XLSX_STRUCT_MATCHERS: list[tuple[str, str]] = [
    ('unique id associated with use case', '_external_id'),
    ('use case name', 'use_case_name'),
    ('use case description', 'use_case_description'),
    ('use case stage', 'use_case_stage'),
    ('use case business owner', 'owner_name'),
    ('lob', 'lob'),
    ('region', 'region'),
    ('name of the vendor', 'vendor_name'),
]


def _parse_survey_xlsx(file_path: Path) -> Dict[str, Any]:
    """Parse an RUAI screening survey XLSX into structured fields + full Q&A.

    Returns dict with keys: structured, responses, gt_questions, cp_questions.
    """
    wb = openpyxl.load_workbook(file_path, data_only=True)

    result: Dict[str, Any] = {
        'structured': {},
        'responses': {},       # All Q&A from main survey sheet
        'gt_questions': [],    # Governance & Technology follow-ups
        'cp_questions': [],    # Control Partner follow-ups
    }

    # --- Sheet 1: Main survey (first sheet) ---
    ws = wb.worksheets[0]
    current_section = ''
    for row in ws.iter_rows(min_row=2, max_row=ws.max_row, values_only=False):
        a_val = str(row[0].value).strip() if row[0].value else ''
        b_val = str(row[1].value).strip() if len(row) > 1 and row[1].value else ''
        c_val = str(row[2].value).strip() if len(row) > 2 and row[2].value else ''

        if a_val:
            current_section = a_val

        if not b_val or not c_val:
            continue

        # Store raw Q&A in responses
        result['responses'][b_val] = c_val

        # Try to match to structured fields
        q_lower = b_val.lower()
        for prefix, field_name in _XLSX_STRUCT_MATCHERS:
            if q_lower.startswith(prefix):
                result['structured'][field_name] = c_val
                break

        # Special case: vendor_or_inhouse (question text varies)
        if 'in-house' in q_lower and ('third-party' in q_lower or 'third party' in q_lower):
            result['structured']['vendor_or_inhouse'] = c_val

    # --- Additional sheets: GT Questions, CP Questions ---
    for ws_name in wb.sheetnames[1:]:
        ws_extra = wb[ws_name]
        name_upper = ws_name.upper()
        is_gt = 'GT' in name_upper and 'QUESTION' in name_upper
        is_cp = 'CP' in name_upper and 'QUESTION' in name_upper

        if not is_gt and not is_cp:
            continue

        start_row = 4 if is_gt else 2
        target = result['gt_questions'] if is_gt else result['cp_questions']

        for row in ws_extra.iter_rows(min_row=start_row, max_row=ws_extra.max_row, values_only=False):
            a_val = str(row[0].value).strip() if row[0].value else ''
            b_val = str(row[1].value).strip() if len(row) > 1 and row[1].value else ''
            c_val = str(row[2].value).strip() if len(row) > 2 and row[2].value else ''
            if not b_val or not c_val:
                continue

            entry: Dict[str, Any] = {'q': b_val, 'a': c_val}
            if is_cp and a_val:
                entry['area'] = a_val

            # Capture follow-up columns D, E, F if present
            follow_ups = []
            for col_idx in range(3, min(len(row), 6)):
                if row[col_idx].value:
                    follow_ups.append(str(row[col_idx].value).strip())
            if follow_ups:
                entry['follow_ups'] = follow_ups

            target.append(entry)

    return result


def _extract_text_from_file(file_path: Path) -> str:
    """Extract text content from a DOCX or PDF file for LLM review."""
    ext = file_path.suffix.lower()
    try:
        if ext == '.docx' and _docx_mod:
            doc = _docx_mod.Document(file_path)
            parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            return '\n'.join(parts)

        elif ext == '.pdf' and _pdfplumber_mod:
            with _pdfplumber_mod.open(file_path) as pdf:
                return '\n'.join(page.extract_text() or '' for page in pdf.pages)

        elif ext in ('.txt', '.csv'):
            return file_path.read_text(errors='replace')

    except Exception as exc:
        logger.warning(f"Could not extract text from {file_path.name}: {exc}")
    return ''


def handle_upload_submission(files: List[FileStorage]) -> Dict[str, Any]:
    """Create a new RUAI case from uploaded intake documents.

    Expects at least one XLSX survey file.  Any other files (PDF, DOCX, etc.)
    are treated as supporting documents — their text is extracted and fed to
    the LLM review.
    """
    if not files:
        return {'status': 'error', 'message': 'No files uploaded'}

    # Separate XLSX survey from supporting docs
    xlsx_file = None
    supporting_files: List[FileStorage] = []
    for f in files:
        if not f or not f.filename:
            continue
        ext = Path(f.filename).suffix.lower()
        if ext == '.xlsx' and xlsx_file is None:
            xlsx_file = f
        else:
            supporting_files.append(f)

    if not xlsx_file:
        return {'status': 'error', 'message': 'No survey XLSX file found in uploads'}

    # Save all files to a temporary upload directory
    temp_dir = UPLOADS_DIR / f"_upload_{uuid.uuid4().hex[:8]}"
    temp_dir.mkdir(parents=True, exist_ok=True)

    xlsx_name = secure_filename(xlsx_file.filename) or 'survey.xlsx'
    xlsx_path = temp_dir / xlsx_name
    xlsx_path.write_bytes(xlsx_file.read())

    try:
        parsed = _parse_survey_xlsx(xlsx_path)
    except Exception as exc:
        logger.error(f"Failed to parse survey XLSX: {exc}", exc_info=True)
        return {'status': 'error', 'message': f'Failed to parse survey file: {exc}'}

    structured = parsed['structured']
    use_case_name = structured.get('use_case_name', '').strip()
    if not use_case_name:
        return {'status': 'error', 'message': 'Could not extract Use Case Name from survey'}

    # Save supporting files & extract text for LLM
    saved_files = [xlsx_name]
    supporting_texts: Dict[str, str] = {}
    for f in supporting_files:
        name = secure_filename(f.filename)
        if not name:
            continue
        ext = Path(name).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            continue
        content = f.read()
        if len(content) > MAX_FILE_SIZE:
            continue
        dest = temp_dir / name
        dest.write_bytes(content)
        saved_files.append(name)
        text = _extract_text_from_file(dest)
        if text.strip():
            supporting_texts[name] = text

    # Rename upload dir to use case name
    final_name = secure_filename(use_case_name) or 'unnamed'
    final_dir = UPLOADS_DIR / final_name
    if final_dir.exists():
        final_dir = UPLOADS_DIR / f"{final_name}_{uuid.uuid4().hex[:6]}"
    temp_dir.rename(final_dir)

    # Build responses JSON with all extracted data
    responses = parsed['responses'].copy()
    responses['_source'] = 'upload'
    if parsed['gt_questions']:
        responses['_gt_questions'] = parsed['gt_questions']
    if parsed['cp_questions']:
        responses['_cp_questions'] = parsed['cp_questions']
    if supporting_texts:
        responses['_supporting_doc_texts'] = supporting_texts

    # Use external ID if available, else generate
    ext_id = structured.get('_external_id', '')
    use_case_id = f"RUAI-{ext_id}" if ext_id else f"RUAI-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"

    with _get_connection() as conn:
        cursor = conn.execute(
            """INSERT INTO submissions
               (use_case_name, use_case_id, use_case_description, use_case_stage,
                owner_name, owner_email, lob, region, vendor_or_inhouse, vendor_name,
                responses, documents, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending_review')""",
            (use_case_name,
             use_case_id,
             structured.get('use_case_description', ''),
             structured.get('use_case_stage', ''),
             structured.get('owner_name', ''),
             '',  # owner_email not in XLSX
             structured.get('lob', ''),
             structured.get('region', ''),
             structured.get('vendor_or_inhouse', ''),
             structured.get('vendor_name'),
             json.dumps(responses),
             ', '.join(saved_files))
        )
        submission_id = cursor.lastrowid

    logger.info(f"RUAI upload submission #{submission_id} ({use_case_id}) — {use_case_name}")

    # Trigger AI review; final status stays pending_review
    _trigger_ai_review(submission_id, final_status='pending_review')

    return {
        'status': 'success',
        'message': f'Case created from uploaded documents (ID: {use_case_id}). AI review is running.',
        'submission_id': submission_id,
        'use_case_id': use_case_id,
    }


# ---------------------------------------------------------------------------
# Submission CRUD
# ---------------------------------------------------------------------------

def handle_submission(form_data: Dict[str, Any], files: Optional[List[FileStorage]] = None) -> Dict[str, Any]:
    """Process a new RUAI screening form submission.

    Args:
        form_data: Parsed JSON body from the form
        files: Optional uploaded documents

    Returns:
        Dict with status and message
    """
    use_case_name = form_data.get('use_case_name', '').strip()
    use_case_description = form_data.get('use_case_description', '').strip()
    use_case_stage = form_data.get('use_case_stage', '').strip()
    owner_name = form_data.get('owner_name', '').strip()
    owner_email = form_data.get('owner_email', '').strip()
    lob = form_data.get('lob', '').strip()
    region = form_data.get('region', '').strip()
    vendor_or_inhouse = form_data.get('vendor_or_inhouse', '').strip()
    vendor_name = form_data.get('vendor_name', '').strip()

    # Validate required fields
    required = {
        'Use Case Name': use_case_name,
        'Use Case Description': use_case_description,
        'Use Case Stage': use_case_stage,
        'Business Owner': owner_name,
        'Owner Email': owner_email,
        'LOB / Function': lob,
        'Development Approach': vendor_or_inhouse,
    }
    missing = [k for k, v in required.items() if not v]
    if missing:
        return {'status': 'error', 'message': f"Missing required fields: {', '.join(missing)}"}

    # Build the full responses JSON (everything except the structured columns)
    structured_keys = {
        'use_case_name', 'use_case_description', 'use_case_stage',
        'owner_name', 'owner_email', 'lob', 'region',
        'vendor_or_inhouse', 'vendor_name',
    }
    responses = {k: v for k, v in form_data.items() if k not in structured_keys}

    # Save uploads
    saved_files = _save_uploads(files or [], use_case_name)

    # Generate a short use-case ID
    use_case_id = f"RUAI-{datetime.now().strftime('%Y%m%d')}-{uuid.uuid4().hex[:6].upper()}"

    with _get_connection() as conn:
        cursor = conn.execute(
            """INSERT INTO submissions
               (use_case_name, use_case_id, use_case_description, use_case_stage,
                owner_name, owner_email, lob, region, vendor_or_inhouse, vendor_name,
                responses, documents, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'submitted')""",
            (use_case_name, use_case_id, use_case_description, use_case_stage,
             owner_name, owner_email, lob, region, vendor_or_inhouse, vendor_name or None,
             json.dumps(responses), ", ".join(saved_files) if saved_files else None)
        )
        submission_id = cursor.lastrowid

    logger.info(f"RUAI screening submission #{submission_id} ({use_case_id}) from {owner_name}")

    # Trigger async LLM review — submitter will see AI feedback before it goes to reviewers
    _trigger_ai_review(submission_id)

    return {
        'status': 'success',
        'message': f'Screening submission received (ID: {use_case_id}). AI review is running — you\'ll be able to review the feedback and revise before sending to the RUAI team.',
        'submission_id': submission_id,
        'use_case_id': use_case_id,
    }


def get_all_submissions(status_filter: Optional[str] = None) -> List[Dict[str, Any]]:
    """Return all submissions with AI risk scores, optionally filtered by status."""
    with _get_connection() as conn:
        base_query = """
            SELECT s.*,
                   ar.overall_risk_score AS ai_overall_risk,
                   ar.preliminary_risk_assessment AS ai_risk_assessment,
                   (SELECT GROUP_CONCAT(ra.reviewer_name || ':' || ra.signed_off)
                    FROM reviewer_assignments ra WHERE ra.submission_id = s.id) AS assignments_raw
            FROM submissions s
            LEFT JOIN ai_reviews ar ON ar.id = (
                SELECT id FROM ai_reviews WHERE submission_id = s.id ORDER BY created_at DESC LIMIT 1
            )
        """
        if status_filter and status_filter != 'all':
            rows = conn.execute(
                base_query + " WHERE s.status = ? ORDER BY s.submitted_at DESC",
                (status_filter,)
            ).fetchall()
        else:
            rows = conn.execute(
                base_query + " ORDER BY s.submitted_at DESC"
            ).fetchall()

        results = []
        terminal_statuses = {'approved', 'rejected'}
        for row in rows:
            d = dict(row)
            # Parse per-area risk assessment for dashboard display
            if d.get('ai_risk_assessment'):
                try:
                    d['ai_risk_assessment'] = json.loads(d['ai_risk_assessment'])
                except (json.JSONDecodeError, TypeError):
                    d['ai_risk_assessment'] = None
            # F7: Compute aging hours for non-terminal statuses
            if d.get('status') not in terminal_statuses and d.get('submitted_at'):
                try:
                    submitted = datetime.strptime(d['submitted_at'][:19], '%Y-%m-%d %H:%M:%S')
                    d['aging_hours'] = round((datetime.now() - submitted).total_seconds() / 3600, 1)
                except (ValueError, TypeError):
                    d['aging_hours'] = None
            else:
                d['aging_hours'] = None
            # F10: Parse reviewer assignments
            assignments_raw = d.pop('assignments_raw', None)
            if assignments_raw:
                d['assignments'] = []
                for part in assignments_raw.split(','):
                    if ':' in part:
                        name, signed = part.rsplit(':', 1)
                        d['assignments'].append({'name': name, 'signed_off': signed == '1'})
            else:
                d['assignments'] = []
            results.append(d)
        return results


def get_submission(submission_id: int) -> Optional[Dict[str, Any]]:
    """Return a single submission with its AI review and reviewer actions."""
    with _get_connection() as conn:
        row = conn.execute(
            "SELECT * FROM submissions WHERE id = ?", (submission_id,)
        ).fetchone()
        if not row:
            return None

        submission = dict(row)

        # Parse JSON fields
        if submission.get('responses'):
            try:
                submission['responses'] = json.loads(submission['responses'])
            except (json.JSONDecodeError, TypeError):
                submission['responses'] = {}

        # Get AI review
        ai_row = conn.execute(
            "SELECT * FROM ai_reviews WHERE submission_id = ? ORDER BY created_at DESC LIMIT 1",
            (submission_id,)
        ).fetchone()
        if ai_row:
            ai_review = dict(ai_row)
            for json_field in ('completeness_issues', 'risk_flags', 'threat_boundary_analysis',
                               'ai_threat_surface', 'control_coverage_map',
                               'clarifying_questions', 'preliminary_risk_assessment',
                               'review_meta'):
                if ai_review.get(json_field):
                    try:
                        ai_review[json_field] = json.loads(ai_review[json_field])
                    except (json.JSONDecodeError, TypeError):
                        pass
            submission['ai_review'] = ai_review
        else:
            submission['ai_review'] = None

        # F4: Count total AI reviews for comparison feature
        ai_count = conn.execute(
            "SELECT COUNT(*) FROM ai_reviews WHERE submission_id = ?",
            (submission_id,)
        ).fetchone()[0]
        submission['ai_review_count'] = ai_count

        # Get reviewer actions
        actions = conn.execute(
            "SELECT * FROM reviewer_actions WHERE submission_id = ? ORDER BY created_at ASC",
            (submission_id,)
        ).fetchall()
        submission['reviewer_actions'] = [dict(a) for a in actions]

        # F10: Get reviewer assignments
        assignments = conn.execute(
            "SELECT * FROM reviewer_assignments WHERE submission_id = ? ORDER BY assigned_at ASC",
            (submission_id,)
        ).fetchall()
        submission['assignments'] = [dict(a) for a in assignments]

        return submission


def add_reviewer_action(submission_id: int, reviewer_name: str, action: str, notes: str = "") -> Dict[str, Any]:
    """Record a reviewer action and update submission status.

    Args:
        submission_id: Target submission
        reviewer_name: Name of the reviewer
        action: One of 'comment', 'request_changes', 'approve', 'reject'
        notes: Optional notes/comments

    Returns:
        Dict with status and message
    """
    valid_actions = {'comment', 'request_changes', 'approve', 'reject'}
    if action not in valid_actions:
        return {'status': 'error', 'message': f'Invalid action. Must be one of: {", ".join(valid_actions)}'}

    status_map = {
        'approve': 'approved',
        'reject': 'rejected',
        'request_changes': 'changes_requested',
    }

    with _get_connection() as conn:
        # Verify submission exists
        row = conn.execute("SELECT id FROM submissions WHERE id = ?", (submission_id,)).fetchone()
        if not row:
            return {'status': 'error', 'message': 'Submission not found'}

        conn.execute(
            "INSERT INTO reviewer_actions (submission_id, reviewer_name, action, notes) VALUES (?, ?, ?, ?)",
            (submission_id, reviewer_name, action, notes or None)
        )

        # F10: Multi-reviewer sign-off on approve
        if action == 'approve':
            assignments = conn.execute(
                "SELECT * FROM reviewer_assignments WHERE submission_id = ?",
                (submission_id,)
            ).fetchall()
            if assignments:
                conn.execute(
                    """UPDATE reviewer_assignments SET signed_off = 1, signed_off_at = CURRENT_TIMESTAMP
                       WHERE submission_id = ? AND reviewer_name = ?""",
                    (submission_id, reviewer_name)
                )
                unsigned = conn.execute(
                    "SELECT COUNT(*) FROM reviewer_assignments WHERE submission_id = ? AND signed_off = 0",
                    (submission_id,)
                ).fetchone()[0]
                if unsigned > 0:
                    logger.info(f"RUAI approve sign-off by {reviewer_name} on #{submission_id}, {unsigned} still pending")
                    return {'status': 'success', 'message': f'Sign-off recorded. {unsigned} reviewer(s) still pending.'}

        # Update submission status for non-comment actions
        if action in status_map:
            conn.execute(
                "UPDATE submissions SET status = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (status_map[action], submission_id)
            )

    logger.info(f"RUAI reviewer action: {reviewer_name} -> {action} on submission #{submission_id}")
    return {'status': 'success', 'message': f'Action "{action}" recorded successfully'}


def update_submission(submission_id: int, form_data: Dict[str, Any], files: Optional[List[FileStorage]] = None) -> Dict[str, Any]:
    """Update an existing submission with revised answers (submitter revising after AI feedback).

    Args:
        submission_id: Submission to update
        form_data: Full form data (replaces previous responses)
        files: Optional new file uploads (appended to existing)

    Returns:
        Dict with status and message
    """
    with _get_connection() as conn:
        row = conn.execute(
            "SELECT id, status, use_case_name, documents FROM submissions WHERE id = ?",
            (submission_id,)
        ).fetchone()
        if not row:
            return {'status': 'error', 'message': 'Submission not found'}

        current = dict(row)
        # Only allow edits before it's sent to reviewers
        editable_statuses = {'submitted', 'ai_reviewing', 'ai_reviewed', 'changes_requested'}
        if current['status'] not in editable_statuses:
            return {'status': 'error', 'message': f'Cannot edit a submission with status "{current["status"]}"'}

    use_case_name = form_data.get('use_case_name', '').strip()
    use_case_description = form_data.get('use_case_description', '').strip()
    use_case_stage = form_data.get('use_case_stage', '').strip()
    owner_name = form_data.get('owner_name', '').strip()
    owner_email = form_data.get('owner_email', '').strip()
    lob = form_data.get('lob', '').strip()
    region = form_data.get('region', '').strip()
    vendor_or_inhouse = form_data.get('vendor_or_inhouse', '').strip()
    vendor_name = form_data.get('vendor_name', '').strip()

    # Validate required fields
    required = {
        'Use Case Name': use_case_name,
        'Use Case Description': use_case_description,
        'Use Case Stage': use_case_stage,
        'Business Owner': owner_name,
        'Owner Email': owner_email,
        'LOB / Function': lob,
        'Development Approach': vendor_or_inhouse,
    }
    missing = [k for k, v in required.items() if not v]
    if missing:
        return {'status': 'error', 'message': f"Missing required fields: {', '.join(missing)}"}

    structured_keys = {
        'use_case_name', 'use_case_description', 'use_case_stage',
        'owner_name', 'owner_email', 'lob', 'region',
        'vendor_or_inhouse', 'vendor_name',
    }
    responses = {k: v for k, v in form_data.items() if k not in structured_keys}

    # Handle new file uploads (append to existing)
    new_files = _save_uploads(files or [], use_case_name or current['use_case_name'])
    existing_docs = current.get('documents') or ''
    all_docs = [d.strip() for d in existing_docs.split(',') if d.strip()] + new_files
    documents = ', '.join(all_docs) if all_docs else None

    with _get_connection() as conn:
        conn.execute(
            """UPDATE submissions SET
               use_case_name = ?, use_case_description = ?, use_case_stage = ?,
               owner_name = ?, owner_email = ?, lob = ?, region = ?,
               vendor_or_inhouse = ?, vendor_name = ?, responses = ?,
               documents = ?, updated_at = CURRENT_TIMESTAMP
               WHERE id = ?""",
            (use_case_name, use_case_description, use_case_stage,
             owner_name, owner_email, lob, region,
             vendor_or_inhouse, vendor_name or None, json.dumps(responses),
             documents, submission_id)
        )

    logger.info(f"RUAI submission #{submission_id} updated by submitter")

    # Re-run AI review on the revised submission
    _trigger_ai_review(submission_id)

    return {
        'status': 'success',
        'message': 'Submission updated. AI review is re-running on your revised answers.',
    }


def submit_for_review(submission_id: int) -> Dict[str, Any]:
    """Submitter confirms they're done revising — move to reviewer queue and notify.

    Args:
        submission_id: Submission to promote

    Returns:
        Dict with status and message
    """
    with _get_connection() as conn:
        row = conn.execute(
            "SELECT * FROM submissions WHERE id = ?", (submission_id,)
        ).fetchone()
        if not row:
            return {'status': 'error', 'message': 'Submission not found'}

        submission = dict(row)
        promotable = {'submitted', 'ai_reviewed', 'changes_requested'}
        if submission['status'] not in promotable:
            return {'status': 'error', 'message': f'Cannot send for review from status "{submission["status"]}"'}

        conn.execute(
            "UPDATE submissions SET status = 'pending_review', updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (submission_id,)
        )

    # Now notify the RUAI reviewer team
    _send_webex_notification(
        submission_id,
        submission['use_case_id'] or str(submission_id),
        submission['use_case_name'],
        submission['owner_name'],
        submission['lob'],
        submission['use_case_stage'],
    )

    logger.info(f"RUAI submission #{submission_id} sent to reviewer queue")
    return {
        'status': 'success',
        'message': 'Your submission has been sent to the RUAI review team.',
    }


# ---------------------------------------------------------------------------
# F2: Reviewer Checklist
# ---------------------------------------------------------------------------

CHECKLIST_ITEMS = [
    'input_validation',
    'prompt_hardening',
    'output_sanitization',
    'logging_controls',
    'model_isolation',
    'rate_limiting',
    'least_privilege',
    'tool_scoping',
    'data_validation',
    'network_segmentation',
    'access_control_iam',
    'human_oversight',
]

CHECKLIST_LABELS = {
    'input_validation': 'Input Validation & Prompt Injection Defenses',
    'prompt_hardening': 'Prompt Hardening & Manipulation Prevention',
    'output_sanitization': 'Output Sanitization & Hallucination Safeguards',
    'logging_controls': 'Logging Controls & Monitoring',
    'model_isolation': 'Model Isolation & Inversion/Poisoning Defenses',
    'rate_limiting': 'Rate Limiting & Abuse Prevention',
    'least_privilege': 'Least Privilege & Agent/Tool Scoping',
    'tool_scoping': 'Tool/Plugin Access Controls',
    'data_validation': 'Data Validation & Supply Chain Integrity',
    'network_segmentation': 'Network Segmentation & Firewall',
    'access_control_iam': 'IAM & Authentication Controls',
    'human_oversight': 'Human Oversight & Override Capability',
}


def get_checklist(submission_id: int, reviewer_name: str) -> Dict[str, Any]:
    """Return the checklist for a given submission and reviewer."""
    with _get_connection() as conn:
        row = conn.execute(
            "SELECT checklist_items FROM reviewer_checklists WHERE submission_id = ? AND reviewer_name = ?",
            (submission_id, reviewer_name)
        ).fetchone()
        if row:
            try:
                items = json.loads(row['checklist_items'])
            except (json.JSONDecodeError, TypeError):
                items = {}
        else:
            items = {}
    return {
        'items': items,
        'labels': CHECKLIST_LABELS,
        'order': CHECKLIST_ITEMS,
    }


def save_checklist(submission_id: int, reviewer_name: str, items: Dict[str, bool]) -> Dict[str, Any]:
    """Save checklist items for a given submission and reviewer."""
    valid_items = {k: bool(v) for k, v in items.items() if k in CHECKLIST_ITEMS}
    with _get_connection() as conn:
        conn.execute(
            """INSERT INTO reviewer_checklists (submission_id, reviewer_name, checklist_items, updated_at)
               VALUES (?, ?, ?, CURRENT_TIMESTAMP)
               ON CONFLICT(submission_id, reviewer_name) DO UPDATE SET
               checklist_items = excluded.checklist_items, updated_at = excluded.updated_at""",
            (submission_id, reviewer_name, json.dumps(valid_items))
        )
    return {'status': 'success', 'items': valid_items}


# ---------------------------------------------------------------------------
# F4: All AI Reviews (for comparison)
# ---------------------------------------------------------------------------

def get_all_ai_reviews(submission_id: int) -> List[Dict[str, Any]]:
    """Return all AI reviews for a submission, ordered by created_at ASC."""
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM ai_reviews WHERE submission_id = ? ORDER BY created_at ASC",
            (submission_id,)
        ).fetchall()
        results = []
        for row in rows:
            review = dict(row)
            for json_field in ('completeness_issues', 'risk_flags', 'threat_boundary_analysis',
                               'ai_threat_surface', 'control_coverage_map',
                               'clarifying_questions', 'preliminary_risk_assessment',
                               'review_meta'):
                if review.get(json_field):
                    try:
                        review[json_field] = json.loads(review[json_field])
                    except (json.JSONDecodeError, TypeError):
                        pass
            results.append(review)
        return results


# ---------------------------------------------------------------------------
# F5: Similar Past Cases (BM25)
# ---------------------------------------------------------------------------

def find_similar_submissions(submission_id: int, top_n: int = 5) -> List[Dict[str, Any]]:
    """Find submissions similar to the given one using BM25 on use_case_name + description."""
    try:
        from rank_bm25 import BM25Okapi
    except ImportError:
        logger.warning("rank_bm25 not installed, similarity search unavailable")
        return []

    with _get_connection() as conn:
        target = conn.execute(
            "SELECT use_case_name, use_case_description FROM submissions WHERE id = ?",
            (submission_id,)
        ).fetchone()
        if not target:
            return []

        all_rows = conn.execute(
            """SELECT s.id, s.use_case_name, s.use_case_description, s.lob, s.status, s.submitted_at,
                      ar.overall_risk_score AS ai_overall_risk
               FROM submissions s
               LEFT JOIN ai_reviews ar ON ar.id = (
                   SELECT id FROM ai_reviews WHERE submission_id = s.id ORDER BY created_at DESC LIMIT 1
               )
               WHERE s.id != ?
               ORDER BY s.submitted_at DESC""",
            (submission_id,)
        ).fetchall()

    if not all_rows:
        return []

    query_text = f"{target['use_case_name']} {target['use_case_description']}"
    query_tokens = query_text.lower().split()

    corpus = []
    for row in all_rows:
        text = f"{row['use_case_name']} {row['use_case_description'] or ''}"
        corpus.append(text.lower().split())

    bm25 = BM25Okapi(corpus)
    scores = bm25.get_scores(query_tokens)

    scored = [(score, dict(row)) for score, row in zip(scores, all_rows)]
    scored.sort(key=lambda x: x[0], reverse=True)

    results = []
    max_score = scored[0][0] if scored else 1
    for score, row in scored[:top_n]:
        if score > 0:
            row['similarity_score'] = round(score / max_score * 100) if max_score > 0 else 0
            results.append(row)

    return results


# ---------------------------------------------------------------------------
# F9: Document file serving
# ---------------------------------------------------------------------------

def get_upload_file_path(submission_id: int, filename: str) -> Optional[Path]:
    """Resolve the path to an uploaded file for a given submission."""
    with _get_connection() as conn:
        row = conn.execute(
            "SELECT use_case_name, documents FROM submissions WHERE id = ?",
            (submission_id,)
        ).fetchone()
        if not row:
            return None

    docs_str = row['documents'] or ''
    doc_list = [d.strip() for d in docs_str.split(',') if d.strip()]
    safe_filename = secure_filename(filename)
    if safe_filename not in doc_list:
        return None

    folder_name = secure_filename(row['use_case_name']) or 'unnamed'

    # Check primary location
    file_path = UPLOADS_DIR / folder_name / safe_filename
    if file_path.is_file():
        return file_path

    # Check upload directories with UUID suffixes
    for d in UPLOADS_DIR.iterdir():
        if d.is_dir() and d.name.startswith(folder_name):
            candidate = d / safe_filename
            if candidate.is_file():
                return candidate

    # Check _upload_ prefixed directories
    for d in UPLOADS_DIR.iterdir():
        if d.is_dir() and d.name.startswith('_upload_'):
            candidate = d / safe_filename
            if candidate.is_file():
                return candidate

    return None


# ---------------------------------------------------------------------------
# F8: Dashboard Analytics
# ---------------------------------------------------------------------------

def get_dashboard_analytics() -> Dict[str, Any]:
    """Compute aggregate analytics for the RUAI dashboard."""
    with _get_connection() as conn:
        rows = conn.execute("""
            SELECT s.id, s.lob, s.status, s.submitted_at, s.updated_at,
                   ar.overall_risk_score AS ai_overall_risk
            FROM submissions s
            LEFT JOIN ai_reviews ar ON ar.id = (
                SELECT id FROM ai_reviews WHERE submission_id = s.id ORDER BY created_at DESC LIMIT 1
            )
        """).fetchall()

    submissions = [dict(r) for r in rows]
    total = len(submissions)

    lob_counts: Dict[str, int] = {}
    for s in submissions:
        lob = s.get('lob', 'Unknown') or 'Unknown'
        lob_counts[lob] = lob_counts.get(lob, 0) + 1

    risk_counts = {'Low': 0, 'Medium': 0, 'High': 0, 'Critical': 0}
    for s in submissions:
        risk = s.get('ai_overall_risk')
        if risk in risk_counts:
            risk_counts[risk] += 1

    approved = sum(1 for s in submissions if s.get('status') == 'approved')
    terminal = sum(1 for s in submissions if s.get('status') in ('approved', 'rejected'))
    approval_rate = round((approved / terminal * 100) if terminal > 0 else 0, 1)

    review_times: list[float] = []
    for s in submissions:
        if s.get('status') in ('approved', 'rejected') and s.get('submitted_at') and s.get('updated_at'):
            try:
                submitted = datetime.strptime(s['submitted_at'][:19], '%Y-%m-%d %H:%M:%S')
                updated = datetime.strptime(s['updated_at'][:19], '%Y-%m-%d %H:%M:%S')
                hours = (updated - submitted).total_seconds() / 3600
                if hours >= 0:
                    review_times.append(hours)
            except (ValueError, TypeError):
                pass
    avg_review_hours = round(sum(review_times) / len(review_times), 1) if review_times else 0

    monthly: Dict[str, int] = {}
    for s in submissions:
        if s.get('submitted_at'):
            month = s['submitted_at'][:7]
            monthly[month] = monthly.get(month, 0) + 1

    status_labels = {
        'submitted': 'Submitted', 'ai_reviewing': 'AI Reviewing',
        'ai_reviewed': 'AI Reviewed', 'pending_review': 'Pending Review',
        'approved': 'Approved', 'rejected': 'Rejected',
        'changes_requested': 'Changes Requested',
    }
    status_counts: Dict[str, int] = {}
    for s in submissions:
        raw = s.get('status', 'unknown')
        label = status_labels.get(raw, raw.replace('_', ' ').title())
        status_counts[label] = status_counts.get(label, 0) + 1

    return {
        'total': total,
        'lob_counts': lob_counts,
        'risk_counts': risk_counts,
        'status_counts': status_counts,
        'approval_rate': approval_rate,
        'avg_review_hours': avg_review_hours,
        'monthly_submissions': monthly,
    }


# ---------------------------------------------------------------------------
# F10: Reviewer Assignments
# ---------------------------------------------------------------------------

def assign_reviewer(submission_id: int, reviewer_name: str, assigned_by: str) -> Dict[str, Any]:
    """Assign a reviewer to a submission."""
    with _get_connection() as conn:
        row = conn.execute("SELECT id FROM submissions WHERE id = ?", (submission_id,)).fetchone()
        if not row:
            return {'status': 'error', 'message': 'Submission not found'}
        try:
            conn.execute(
                "INSERT INTO reviewer_assignments (submission_id, reviewer_name, assigned_by) VALUES (?, ?, ?)",
                (submission_id, reviewer_name, assigned_by)
            )
        except sqlite3.IntegrityError:
            return {'status': 'error', 'message': f'{reviewer_name} is already assigned'}
    return {'status': 'success', 'message': f'{reviewer_name} assigned'}


def remove_reviewer_assignment(submission_id: int, reviewer_name: str) -> Dict[str, Any]:
    """Remove a reviewer assignment."""
    with _get_connection() as conn:
        cursor = conn.execute(
            "DELETE FROM reviewer_assignments WHERE submission_id = ? AND reviewer_name = ?",
            (submission_id, reviewer_name)
        )
        if cursor.rowcount == 0:
            return {'status': 'error', 'message': 'Assignment not found'}
    return {'status': 'success', 'message': f'{reviewer_name} unassigned'}


def get_reviewer_assignments(submission_id: int) -> List[Dict[str, Any]]:
    """Get all reviewer assignments for a submission."""
    with _get_connection() as conn:
        rows = conn.execute(
            "SELECT * FROM reviewer_assignments WHERE submission_id = ? ORDER BY assigned_at ASC",
            (submission_id,)
        ).fetchall()
    return [dict(r) for r in rows]


def sign_off_assignment(submission_id: int, reviewer_name: str) -> Dict[str, Any]:
    """Mark a reviewer's assignment as signed off."""
    with _get_connection() as conn:
        cursor = conn.execute(
            """UPDATE reviewer_assignments SET signed_off = 1, signed_off_at = CURRENT_TIMESTAMP
               WHERE submission_id = ? AND reviewer_name = ?""",
            (submission_id, reviewer_name)
        )
        if cursor.rowcount == 0:
            return {'status': 'error', 'message': 'Assignment not found'}
    return {'status': 'success', 'message': f'{reviewer_name} signed off'}


# ---------------------------------------------------------------------------
# LLM Review
# ---------------------------------------------------------------------------

class RUAIReviewResponse(BaseModel):
    """Structured output model for the AI security screening review."""
    overall_risk_score: str = Field(
        default="Medium",
        description="Single overall risk rating for the entire use case: Low, Medium, High, or Critical"
    )
    completeness_issues: List[str] = Field(
        default_factory=list,
        description="List of questions that are unanswered, vague, or need more detail"
    )
    risk_flags: List[Dict[str, str]] = Field(
        default_factory=list,
        description="List of {area, level, reason} security risk flags identified"
    )
    threat_boundary_analysis: List[Dict[str, Any]] = Field(
        default_factory=list,
        description="Per-boundary assessment: {boundary, components, risks, controls_present, gaps}"
    )
    ai_threat_surface: List[Dict[str, str]] = Field(
        default_factory=list,
        description="AI-specific threats: {threat, likelihood, impact, details}"
    )
    control_coverage_map: List[Dict[str, str]] = Field(
        default_factory=list,
        description="Control mapping: {boundary, control, status, gap_detail}. status is Present, Partial, or Missing."
    )
    clarifying_questions: List[str] = Field(
        default_factory=list,
        description="Follow-up questions the reviewer should ask the submitting team"
    )
    preliminary_risk_assessment: Dict[str, str] = Field(
        default_factory=dict,
        description="Risk rating per security area. Each value must be Low, Medium, High, or Critical."
    )
    review_summary: str = Field(
        default="",
        description="2-3 paragraph draft review narrative for the RUAI reviewer"
    )


def _build_review_prompt(submission: Dict[str, Any]) -> str:
    """Build the LLM prompt for reviewing a submission.

    The static template is pulled from the active version managed via
    `ruai_prompts_handler`; dynamic submission data is substituted into
    `{{PLACEHOLDER}}` slots at render time.
    """
    from src.components.web import ruai_prompts_handler, ruai_docs_handler
    # Format responses into readable sections
    responses = submission.get('responses', {})
    sections_text = []

    if responses.get('_source') == 'upload':
        # Upload-based submission: responses are raw Q&A from XLSX
        survey_lines = ["## Survey Responses"]
        for q, a in responses.items():
            if q.startswith('_') or isinstance(a, (dict, list)):
                continue
            survey_lines.append(f"**{q}**: {a}")
        if len(survey_lines) > 1:
            sections_text.append("\n".join(survey_lines))
    else:
        # Standard form submission: format using SCREENING_SECTIONS
        for section in SCREENING_SECTIONS:
            section_lines = [f"## {section['title']}"]
            for field in section['fields']:
                name = field['name']
                value = responses.get(name) or submission.get(name, '')
                if value:
                    if isinstance(value, list):
                        value = ", ".join(value)
                    section_lines.append(f"**{field['label']}**: {value}")
            if len(section_lines) > 1:
                sections_text.append("\n".join(section_lines))

    # GT follow-up questions (upload submissions)
    gt_questions = responses.get('_gt_questions', [])
    if gt_questions:
        gt_lines = ["## Follow-up Questions (Governance & Technology Review)"]
        for item in gt_questions:
            gt_lines.append(f"**Q**: {item['q']}")
            gt_lines.append(f"**A**: {item['a']}")
            for fu in item.get('follow_ups', []):
                gt_lines.append(f"  *Follow-up*: {fu}")
            gt_lines.append("")
        sections_text.append("\n".join(gt_lines))

    # CP follow-up questions (upload submissions)
    cp_questions = responses.get('_cp_questions', [])
    if cp_questions:
        cp_lines = ["## Control Partner Questions"]
        for item in cp_questions:
            area = item.get('area', '')
            prefix = f"[{area}] " if area else ''
            cp_lines.append(f"**{prefix}Q**: {item['q']}")
            cp_lines.append(f"**A**: {item['a']}")
            for fu in item.get('follow_ups', []):
                cp_lines.append(f"  *Follow-up*: {fu}")
            cp_lines.append("")
        sections_text.append("\n".join(cp_lines))

    # Supporting document texts (upload submissions)
    for filename, text in responses.get('_supporting_doc_texts', {}).items():
        truncated = text[:8000]
        if len(text) > 8000:
            truncated += "\n[... truncated ...]"
        sections_text.append(f"## Supporting Document: {filename}\n\n{truncated}")

    submission_text = "\n\n".join(sections_text)

    try:
        ref_text = ruai_docs_handler.load_all_text()
    except Exception as exc:
        logger.warning("Failed to load RUAI reference docs for prompt stuffing: %s", exc)
        ref_text = ""
    if ref_text.strip():
        reference_block = "## Framework Reference Documents (Retrieved)\n\n" + ref_text
    else:
        reference_block = "_No framework reference documents uploaded._"

    return ruai_prompts_handler.render("review", {
        "USE_CASE_NAME": submission.get('use_case_name', 'N/A'),
        "STAGE": submission.get('use_case_stage', 'N/A'),
        "OWNER_NAME": submission.get('owner_name', 'N/A'),
        "OWNER_EMAIL": submission.get('owner_email', 'N/A'),
        "LOB": submission.get('lob', 'N/A'),
        "DEVELOPMENT": submission.get('vendor_or_inhouse', 'N/A'),
        "SUBMISSION_DATA": submission_text,
        "REFERENCE_DOCS": reference_block,
    })


def get_review_phase(submission_id: int) -> str:
    """Return the current review phase for a submission."""
    return _review_phases.get(submission_id, '')


def get_review_stats(submission_id: int) -> Optional[dict]:
    """Return in-memory LLM stats captured during the review, if available."""
    return _review_stats.get(submission_id)


def _run_ai_review(submission_id: int, final_status: str = 'ai_reviewed') -> None:
    """Execute the LLM review for a submission (runs in background thread)."""
    logger.info(f"Starting AI review for RUAI submission #{submission_id}")
    _review_phases[submission_id] = 'loading_submission'

    # Update status
    with _get_connection() as conn:
        conn.execute(
            "UPDATE submissions SET status = 'ai_reviewing', updated_at = CURRENT_TIMESTAMP WHERE id = ?",
            (submission_id,)
        )

    # Load full submission
    submission = get_submission(submission_id)
    if not submission:
        logger.error(f"RUAI AI review: submission #{submission_id} not found")
        _review_phases[submission_id] = 'error'
        return

    try:
        _review_phases[submission_id] = 'connecting_llm'
        import time as _time
        from my_bot.utils.llm_factory import create_llm, extract_token_metrics

        llm = create_llm(
            temperature=0.2,
            max_tokens=8192,
            model_kwargs={'response_format': {'type': 'json_object'}},
        )
        # Fallback to m1 (GLM-4.7-Flash) on the LLM failure. The JSON parser
        # below tolerates non-strict output from either model (strips
        # code fences, regex-finds JSON block, strips control chars).

        prompt = _build_review_prompt(submission)

        _review_phases[submission_id] = 'analyzing'
        _t0 = _time.monotonic()
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(llm.invoke, prompt)
            try:
                ai_message = future.result(timeout=LLM_TIMEOUT_SECONDS)
            except FuturesTimeoutError:
                future.cancel()
                raise RuntimeError("LLM did not respond within timeout")
        _duration_s = round(_time.monotonic() - _t0, 1)

        # Extract token metrics from LLM response
        _meta = getattr(ai_message, 'response_metadata', None) or {}
        _tok = extract_token_metrics(_meta)
        _model_name = _meta.get('model_name') or _meta.get('model') or ''
        _review_meta = {
            'duration_s': _duration_s,
            'input_tokens': _tok['input_tokens'],
            'output_tokens': _tok['output_tokens'],
            'total_tokens': _tok['input_tokens'] + _tok['output_tokens'],
            'model': _model_name,
        }

        # Parse JSON from response
        raw_text = ai_message.content.strip()
        logger.debug(f"RUAI raw LLM response length: {len(raw_text)}")

        # Handle </think> blocks from reasoning models
        if '</think>' in raw_text:
            raw_text = raw_text.split('</think>')[-1].strip()

        # Strip markdown code fences
        if raw_text.startswith('```'):
            raw_text = raw_text.split('\n', 1)[1] if '\n' in raw_text else raw_text[3:]
            raw_text = raw_text.rsplit('```', 1)[0].strip()

        # Try to find JSON object in the response
        import re
        if not raw_text.startswith('{'):
            json_match = re.search(r'\{[\s\S]*\}', raw_text)
            if json_match:
                raw_text = json_match.group(0)
            else:
                raise ValueError(f"No JSON object found in LLM response. Starts with: {raw_text[:200]}")

        # Strip ALL control chars — the LLM emits literal newlines/tabs
        # inside JSON string values which breaks parsers.  JSON structure
        # doesn't need literal newlines (they're just formatting whitespace),
        # and properly escaped \n sequences are two chars (\ + n), unaffected.
        raw_text = re.sub(r'[\x00-\x1f]', '', raw_text)

        result = RUAIReviewResponse.model_validate_json(raw_text)

        # Store the review
        _review_phases[submission_id] = 'saving_results'
        _review_stats[submission_id] = _review_meta
        with _get_connection() as conn:
            conn.execute(
                """INSERT INTO ai_reviews
                   (submission_id, overall_risk_score, completeness_issues, risk_flags,
                    threat_boundary_analysis, ai_threat_surface, control_coverage_map,
                    clarifying_questions, preliminary_risk_assessment, review_summary_md,
                    review_meta)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (submission_id,
                 result.overall_risk_score,
                 json.dumps(result.completeness_issues),
                 json.dumps([rf if isinstance(rf, dict) else dict(rf) for rf in result.risk_flags] if result.risk_flags else []),
                 json.dumps([tb if isinstance(tb, dict) else dict(tb) for tb in result.threat_boundary_analysis] if result.threat_boundary_analysis else []),
                 json.dumps([ts if isinstance(ts, dict) else dict(ts) for ts in result.ai_threat_surface] if result.ai_threat_surface else []),
                 json.dumps([cm if isinstance(cm, dict) else dict(cm) for cm in result.control_coverage_map] if result.control_coverage_map else []),
                 json.dumps(result.clarifying_questions),
                 json.dumps(result.preliminary_risk_assessment),
                 result.review_summary,
                 json.dumps(_review_meta))
            )
            conn.execute(
                "UPDATE submissions SET status = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (final_status, submission_id)
            )

        _review_phases[submission_id] = 'done'
        logger.info(f"AI review completed for RUAI submission #{submission_id}")

    except Exception as exc:
        logger.warning(f"RUAI AI review failed for submission #{submission_id}: {exc}")
        _review_phases[submission_id] = 'error'
        # Graceful degradation — revert so reviewer still sees it without AI pre-review
        fallback = final_status if final_status == 'pending_review' else 'submitted'
        with _get_connection() as conn:
            conn.execute(
                "UPDATE submissions SET status = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                (fallback, submission_id)
            )


def _trigger_ai_review(submission_id: int, final_status: str = 'ai_reviewed') -> None:
    """Trigger an async AI review in a background thread."""
    from threading import Thread
    thread = Thread(target=_run_ai_review, args=(submission_id, final_status), daemon=True)
    thread.start()
    logger.debug(f"RUAI AI review thread started for submission #{submission_id}")


def rerun_ai_review(submission_id: int) -> Dict[str, Any]:
    """Re-trigger the AI review for a submission."""
    with _get_connection() as conn:
        row = conn.execute("SELECT id, status, responses FROM submissions WHERE id = ?", (submission_id,)).fetchone()
        if not row:
            return {'status': 'error', 'message': 'Submission not found'}

    # Upload submissions should return to pending_review after AI review
    final_status = 'ai_reviewed'
    try:
        responses = json.loads(row['responses'] or '{}')
        if responses.get('_source') == 'upload':
            final_status = 'pending_review'
    except (json.JSONDecodeError, TypeError):
        pass

    _trigger_ai_review(submission_id, final_status=final_status)
    return {'status': 'success', 'message': 'AI review started.'}


# ---------------------------------------------------------------------------
# Webex notification
# ---------------------------------------------------------------------------

def _send_webex_notification(
    submission_id: int,
    use_case_id: str,
    use_case_name: str,
    owner_name: str,
    lob: str,
    stage: str,
) -> None:
    """Send a Webex notification about a new RUAI screening submission."""
    access_token = CONFIG.webex_bot_access_token_pokedex
    room_id = CONFIG.webex_room_id_dev_test_space

    if not access_token or not room_id:
        logger.warning("Webex credentials not configured — RUAI notification skipped")
        return

    base_url = CONFIG.web_server_url
    view_url = f"{base_url}/ruai-dashboard/{submission_id}"

    stage_emoji = {'PoC': '🧪', 'Pilot': '🚀', 'Implementation': '🏗️'}.get(stage, '📋')

    adaptive_card = {
        "$schema": "http://adaptivecards.io/schemas/adaptive-card.json",
        "type": "AdaptiveCard",
        "version": "1.3",
        "body": [
            {
                "type": "TextBlock",
                "text": "🤖 New RUAI Screening Submission",
                "weight": "Bolder",
                "size": "Large",
                "color": "Accent",
                "horizontalAlignment": "Center",
            },
            {
                "type": "Container",
                "style": "accent",
                "spacing": "Medium",
                "items": [{
                    "type": "TextBlock",
                    "text": f"📋 {use_case_name}",
                    "weight": "Bolder",
                    "size": "Medium",
                    "horizontalAlignment": "Center",
                    "wrap": True,
                }],
            },
            {
                "type": "ColumnSet",
                "spacing": "Medium",
                "columns": [
                    {
                        "type": "Column",
                        "width": "stretch",
                        "items": [
                            {"type": "TextBlock", "text": "👤 Owner", "weight": "Bolder", "color": "Accent", "size": "Small"},
                            {"type": "TextBlock", "text": owner_name, "spacing": "None", "wrap": True},
                        ],
                    },
                    {
                        "type": "Column",
                        "width": "stretch",
                        "items": [
                            {"type": "TextBlock", "text": "🏢 LOB", "weight": "Bolder", "color": "Accent", "size": "Small"},
                            {"type": "TextBlock", "text": lob, "spacing": "None", "wrap": True},
                        ],
                    },
                    {
                        "type": "Column",
                        "width": "stretch",
                        "items": [
                            {"type": "TextBlock", "text": f"{stage_emoji} Stage", "weight": "Bolder", "color": "Accent", "size": "Small"},
                            {"type": "TextBlock", "text": stage, "spacing": "None"},
                        ],
                    },
                ],
            },
            {
                "type": "TextBlock",
                "text": f"ID: {use_case_id}",
                "size": "Small",
                "isSubtle": True,
                "spacing": "Medium",
            },
        ],
        "actions": [
            {"type": "Action.OpenUrl", "title": "Review Submission", "url": view_url, "style": "positive"},
        ],
    }

    try:
        webex_api = WebexAPI(access_token=access_token, disable_ssl_verify=True)
        success = safe_send_message(
            webex_api, room_id,
            text=f"New RUAI Screening: {use_case_name} from {owner_name} ({lob})",
            attachments=[{
                "contentType": "application/vnd.microsoft.card.adaptive",
                "content": adaptive_card,
            }],
        )
        if not success:
            logger.warning("Webex message send returned False for RUAI notification")
    except Exception as exc:
        logger.warning(f"Could not send RUAI Webex notification: {exc}")
