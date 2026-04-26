# /services/document_processor.py
"""
Document Processing Module

This module handles RAG (Retrieval-Augmented Generation) document loading,
processing, and vector store management for the security operations bot.

Uses ChromaDB for persistent vector storage.
"""

import hashlib
import logging
import os
import time
from pathlib import Path
from typing import Any, List, Optional

import chromadb
import requests
from langchain_community.document_loaders import (
    PyPDFDirectoryLoader,
    UnstructuredExcelLoader,
    UnstructuredWordDocumentLoader,
)
from langchain_community.retrievers import BM25Retriever
from langchain_classic.retrievers.ensemble import EnsembleRetriever
from langchain_core.documents import Document
from langchain_core.retrievers import BaseRetriever

# Handle different LangChain versions (0.3.x vs 1.0.x)
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
except ImportError:
    from langchain_text_splitters import RecursiveCharacterTextSplitter


# Collection name for documents
DOCUMENTS_COLLECTION = "local_documents"


from my_bot.utils.embedding_function import OpenAIEmbeddingFunction

logger = logging.getLogger(__name__)


class OllamaEmbeddingFunction(OpenAIEmbeddingFunction):
    """Legacy alias — delegates to OpenAIEmbeddingFunction (OpenAI-compatible API)."""


# ------------------------------------------------------------------ Structured loaders
#
# The default LangChain `UnstructuredWordDocumentLoader` / `UnstructuredExcelLoader`
# flatten tables into prose, which destroys row→column relationships. Security
# policies, SOC 2 matrices, ticket exports, questionnaire banks, and rule
# catalogs all live in tables. The loaders below use python-docx and openpyxl
# directly to emit one atomic chunk per table row with column headers embedded
# in the content, plus heading-hierarchy `section` metadata for Word documents.
#
# These are app-agnostic — any RAG ingest pipeline in IR can call them. First
# consumer is the Customer Assurance KB; docs_library, rules catalog, etc. can
# opt in on their next re-index.


def _iter_docx_blocks(doc):
    from docx.oxml.ns import qn
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield DocxParagraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield DocxTable(child, doc)


def load_word_structured(path: str) -> List[dict]:
    """Extract structured (content, metadata) items from a .docx file.

    Paragraphs become items tagged with the running heading stack as `section`.
    Tables become one item per data row with column headers serialized as
    key-value lines inside the content, so embedding retrieval sees structured
    text (e.g. `Control ID: AC-2\\nOwner: Identity Team`) instead of a row blob.

    Returns a list of {"content": str, "metadata": dict}. Metadata includes
    `source`, `element_type` ("paragraph" or "table_row"), `section`, and for
    table rows also `table_headers` and `row_index`.
    """
    from docx import Document as DocxDocument
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph

    source = os.path.basename(path)
    items: List[dict] = []

    try:
        doc = DocxDocument(path)
    except Exception as e:
        logger.error(f"[document_processor] python-docx failed on {source}: {e}")
        return items

    heading_stack: List[str] = []

    for block in _iter_docx_blocks(doc):
        if isinstance(block, DocxParagraph):
            text = (block.text or "").strip()
            if not text:
                continue
            style_name = ""
            try:
                style_name = (block.style.name or "") if block.style else ""
            except Exception:
                style_name = ""

            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip() or "1")
                except ValueError:
                    level = 1
                heading_stack = heading_stack[: max(0, level - 1)]
                heading_stack.append(text)
                continue

            if len(text) < 20:
                continue

            section = " > ".join(heading_stack)
            content = f"[Section: {section}]\n\n{text}" if section else text
            items.append({
                "content": content,
                "metadata": {
                    "source": source,
                    "element_type": "paragraph",
                    "section": section,
                },
            })

        elif isinstance(block, DocxTable):
            rows = list(block.rows)
            if len(rows) < 2:
                continue

            headers = [(cell.text or "").strip() for cell in rows[0].cells]
            if not any(headers):
                continue

            section = " > ".join(heading_stack)
            headers_joined = " | ".join(h for h in headers if h)

            for row_idx, row in enumerate(rows[1:], start=1):
                cells = [(cell.text or "").strip() for cell in row.cells]
                if not any(cells):
                    continue

                kv_lines: List[str] = []
                for h, v in zip(headers, cells):
                    if h and v:
                        kv_lines.append(f"{h}: {v}")
                    elif v:
                        kv_lines.append(v)
                if not kv_lines:
                    continue

                header_prefix = f"[Table row in: {section}]" if section else "[Table row]"
                content = header_prefix + "\n" + "\n".join(kv_lines)
                items.append({
                    "content": content,
                    "metadata": {
                        "source": source,
                        "element_type": "table_row",
                        "section": section,
                        "table_headers": headers_joined,
                        "row_index": row_idx,
                    },
                })

    return items


def load_excel_structured(path: str) -> List[dict]:
    """Extract structured (content, metadata) items from an .xlsx file.

    Each sheet is treated as a separate table; row 1 is assumed to be headers;
    each subsequent non-empty row becomes one atomic chunk with key-value lines.
    Response banks, SOC 2 control matrices, and questionnaire libraries follow
    this shape almost universally.
    """
    from openpyxl import load_workbook

    source = os.path.basename(path)
    items: List[dict] = []

    try:
        wb = load_workbook(path, data_only=True, read_only=True)
    except Exception as e:
        logger.error(f"[document_processor] openpyxl failed on {source}: {e}")
        return items

    try:
        for sheet_name in wb.sheetnames:
            try:
                ws = wb[sheet_name]
                rows_iter = ws.iter_rows(values_only=True)
            except Exception as e:
                logger.warning(f"[document_processor] skipping sheet {sheet_name} in {source}: {e}")
                continue

            try:
                header_row = next(rows_iter)
            except StopIteration:
                continue

            headers = [(str(h).strip() if h is not None else "") for h in (header_row or [])]
            if not any(headers):
                continue

            headers_joined = " | ".join(h for h in headers if h)
            section = f"Sheet: {sheet_name}"

            for row_idx, row in enumerate(rows_iter, start=2):
                cells = [("" if v is None else str(v).strip()) for v in (row or [])]
                if not any(cells):
                    continue

                kv_lines: List[str] = []
                for h, v in zip(headers, cells):
                    if h and v:
                        kv_lines.append(f"{h}: {v}")
                    elif v:
                        kv_lines.append(v)
                if not kv_lines:
                    continue

                content = f"[{section} — row {row_idx}]\n" + "\n".join(kv_lines)
                items.append({
                    "content": content,
                    "metadata": {
                        "source": source,
                        "element_type": "xlsx_row",
                        "section": section,
                        "table_headers": headers_joined,
                        "row_index": row_idx,
                    },
                })
    finally:
        try:
            wb.close()
        except Exception:
            pass

    return items


# ------------------------------------------------------------------ Reranker client
#
# App-agnostic cross-encoder reranker client. Posts a (query, documents) batch
# to an HTTP reranker endpoint and returns the top_k documents with a
# calibrated relevance score attached to metadata under `rerank_score`.
#
# The endpoint defaults to http://localhost:8020 (local lab-vm1 CPU service,
# `deployment/reranker_server.py` behind `ir-reranker.service`). When the
# mac-m3 reranker with a larger model is stood up, flip `RERANKER_BASE_URL`
# in .env and restart — zero code change here or in callers.
#
# Graceful fallback: if the endpoint is unreachable or errors, the function
# returns the first top_k input documents unchanged, so upstream flows never
# break because the reranker is down.


def rerank_documents(
    query: str,
    documents: List[Document],
    top_k: int = 5,
    timeout: float = 10.0,
) -> List[Document]:
    if not documents:
        return []

    base_url = (os.environ.get("RERANKER_BASE_URL") or "http://localhost:8020").rstrip("/")

    try:
        resp = requests.post(
            f"{base_url}/rerank",
            json={
                "query": query,
                "documents": [d.page_content for d in documents],
                "top_k": top_k,
            },
            timeout=timeout,
        )
        resp.raise_for_status()
        results = (resp.json() or {}).get("results") or []
    except Exception as e:
        logger.warning(f"[rerank] endpoint unreachable, returning input unreranked: {e}")
        return documents[:top_k]

    reranked: List[Document] = []
    for r in results:
        idx = r.get("index")
        if idx is None or idx < 0 or idx >= len(documents):
            continue
        orig = documents[idx]
        new_meta = dict(orig.metadata or {})
        new_meta["rerank_score"] = r.get("score")
        reranked.append(Document(page_content=orig.page_content, metadata=new_meta))

    return reranked if reranked else documents[:top_k]


class ChromaRetriever(BaseRetriever):
    """Custom retriever for ChromaDB that implements LangChain's BaseRetriever interface."""

    collection: Any  # ChromaDB collection
    embedding_fn: OllamaEmbeddingFunction
    k: int = 5

    class Config:
        arbitrary_types_allowed = True

    def _get_relevant_documents(self, query: str) -> List[Document]:
        """Retrieve relevant documents from ChromaDB."""
        try:
            # Generate query embedding
            query_embedding = self.embedding_fn([query])[0]

            # Query ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=self.k,
                include=['documents', 'metadatas', 'distances']
            )

            # Convert to LangChain Documents
            documents = []
            if results['ids'] and results['ids'][0]:
                for i, doc_id in enumerate(results['ids'][0]):
                    content = results['documents'][0][i] if results['documents'] else ""
                    metadata = results['metadatas'][0][i] if results['metadatas'] else {}
                    documents.append(Document(page_content=content, metadata=metadata))

            return documents

        except Exception as e:
            logging.error(f"ChromaDB retrieval error: {e}")
            return []


class DocumentProcessor:
    """Handles document loading, processing, and vector store management using ChromaDB"""

    def __init__(self, pdf_directory: str, chroma_path: str = None):
        self.pdf_directory = pdf_directory

        # Set ChromaDB path (default to same location as old FAISS index)
        if chroma_path is None:
            # Convert old faiss_index_path to chroma path
            chroma_path = str(Path(pdf_directory).parent / "chroma_documents")
        self.chroma_path = chroma_path

        self._client: Optional[chromadb.PersistentClient] = None
        self._collection = None
        self._embedding_fn = OllamaEmbeddingFunction()

        self.retriever = None
        self.all_documents = []  # Store documents for BM25 retriever

        # Document processing configuration
        self.chunk_size = 1500
        self.chunk_overlap = 300
        self.retrieval_k = 5

    @property
    def client(self) -> chromadb.PersistentClient:
        """Lazy-load ChromaDB client."""
        if self._client is None:
            self._client = chromadb.PersistentClient(path=self.chroma_path)
        return self._client

    @property
    def collection(self):
        """Get or create the documents collection."""
        if self._collection is None:
            self._collection = self.client.get_or_create_collection(
                name=DOCUMENTS_COLLECTION,
                metadata={"description": "Local document embeddings for RAG"}
            )
        return self._collection

    @property
    def vector_store(self):
        """Compatibility property - returns collection count > 0."""
        try:
            return self.collection.count() > 0
        except Exception:
            return False

    def _generate_doc_id(self, content: str, source: str) -> str:
        """Generate a unique ID for a document chunk."""
        hash_input = f"{source}:{content[:200]}"
        return hashlib.md5(hash_input.encode()).hexdigest()

    def load_documents_from_folder(self) -> List:
        """Load documents from the specified folder"""
        documents = []
        pdf_loaded = False

        if not os.path.exists(self.pdf_directory):
            logging.warning(f"Folder does not exist: {self.pdf_directory}")
            return documents

        for fname in os.listdir(self.pdf_directory):
            fpath = os.path.join(self.pdf_directory, fname)
            if not os.path.isfile(fpath):
                continue

            ext = os.path.splitext(fname)[1].lower()

            try:
                if ext == ".pdf" and not pdf_loaded:
                    loader = PyPDFDirectoryLoader(self.pdf_directory)
                    documents.extend(loader.load())
                    pdf_loaded = True
                    logging.debug(f"Loaded PDF documents from {self.pdf_directory}")
                elif ext in [".doc", ".docx"]:
                    loader = UnstructuredWordDocumentLoader(fpath)
                    documents.extend(loader.load())
                    logging.debug(f"Loaded Word document: {fname}")
                elif ext in [".xlsx", ".xls"]:
                    loader = UnstructuredExcelLoader(fpath)
                    documents.extend(loader.load())
                    logging.debug(f"Loaded Excel document: {fname}")
            except Exception as e:
                logging.error(f"Failed to load {fname}: {e}")

        return documents

    def create_text_chunks(self, documents: List) -> List:
        """Split documents into optimized text chunks"""
        if not documents:
            return []

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
        )

        texts = text_splitter.split_documents(documents)
        logging.debug(f"Split into {len(texts)} text chunks.")
        return texts

    def sync_documents(self, embeddings=None, batch_size: int = 10) -> bool:
        """Add new documents to ChromaDB. Skips already-indexed chunks."""
        if not os.path.exists(self.pdf_directory):
            logging.warning(f"PDF directory '{self.pdf_directory}' does not exist.")
            return False

        files_in_dir = [f for f in os.listdir(self.pdf_directory)
                        if f.endswith(('.pdf', '.docx', '.doc', '.xlsx', '.xls'))]
        if not files_in_dir:
            logging.warning(f"PDF directory '{self.pdf_directory}' has no documents.")
            return False

        logging.debug(f"Loading documents from: {self.pdf_directory}")
        documents = self.load_documents_from_folder()

        if not documents:
            logging.warning(f"No documents could be loaded from '{self.pdf_directory}'.")
            return False

        logging.debug(f"Loaded {len(documents)} documents.")

        # Create text chunks
        texts = self.create_text_chunks(documents)
        self.all_documents = texts

        # Get existing IDs to avoid duplicates
        existing_ids = set()
        try:
            result = self.collection.get()
            existing_ids = set(result['ids']) if result['ids'] else set()
            logging.debug(f"ChromaDB has {len(existing_ids)} existing chunks")
        except Exception as e:
            logging.warning(f"Could not get existing IDs: {e}")

        # Prepare new chunks
        ids = []
        documents_content = []
        metadatas = []

        total_chunks = len(texts)
        print(f"   Processing {total_chunks} text chunks...", flush=True)

        for i, doc in enumerate(texts):
            doc_id = self._generate_doc_id(doc.page_content, doc.metadata.get('source', ''))

            # Skip if already indexed
            if doc_id in existing_ids:
                continue

            if (i + 1) % 20 == 0 or i == 0:
                print(f"   Processing chunk {i + 1}/{total_chunks}...", flush=True)

            ids.append(doc_id)
            documents_content.append(doc.page_content)
            metadatas.append(doc.metadata)

        if not ids:
            logging.info("No new chunks to add")
            return True

        # Generate embeddings and upsert to ChromaDB
        print(f"   Embedding and upserting {len(ids)} new chunks...", flush=True)
        try:
            embeddings_list = self._embedding_fn(documents_content)

            self.collection.upsert(
                ids=ids,
                documents=documents_content,
                metadatas=metadatas,
                embeddings=embeddings_list
            )

            logging.debug(f"Added {len(ids)} chunks to ChromaDB")
            return True

        except Exception as e:
            logging.error(f"Error adding to ChromaDB: {e}", exc_info=True)
            return False

    def load_vector_store(self, embeddings=None) -> bool:
        """Load existing vector store (ChromaDB loads automatically)."""
        try:
            count = self.collection.count()
            if count == 0:
                logging.warning("ChromaDB collection is empty")
                return False

            logging.debug(f"ChromaDB collection has {count} chunks")

            # Load documents for BM25 retriever from ChromaDB cache (fast)
            # instead of re-parsing files from disk (slow)
            if not self.all_documents:
                logging.debug("Loading documents for BM25 from ChromaDB cache...")
                result = self.collection.get()
                if result and result['documents']:
                    metadatas = result['metadatas'] or [{}] * len(result['documents'])
                    self.all_documents = [
                        Document(page_content=doc, metadata=meta or {})
                        for doc, meta in zip(result['documents'], metadatas)
                    ]
                    logging.debug(f"Loaded {len(self.all_documents)} chunks for BM25 from cache")

            return True

        except Exception as e:
            logging.error(f"Error loading ChromaDB: {e}", exc_info=True)
            return False

    def initialize_vector_store(self, embeddings=None) -> bool:
        """Load existing vector store. Use sync_documents() or rebuild_index() to update."""
        return self.load_vector_store(embeddings)

    def create_retriever(self):
        """Create hybrid retriever combining vector and keyword search."""
        try:
            count = self.collection.count()
            if count == 0:
                logging.error("ChromaDB collection is empty")
                return None
        except Exception as e:
            logging.error(f"ChromaDB not initialized: {e}")
            return None

        try:
            # Create ChromaDB vector retriever
            vector_retriever = ChromaRetriever(
                collection=self.collection,
                embedding_fn=self._embedding_fn,
                k=self.retrieval_k
            )

            # Create hybrid ensemble retriever (vector + BM25 keyword search)
            if self.all_documents:
                bm25_retriever = BM25Retriever.from_documents(self.all_documents)
                bm25_retriever.k = self.retrieval_k

                self.retriever = EnsembleRetriever(
                    retrievers=[vector_retriever, bm25_retriever],
                    weights=[0.7, 0.3]  # 70% vector similarity, 30% keyword matching
                )
                logging.info("Hybrid retriever created (70% vector + 30% BM25)")
            else:
                self.retriever = vector_retriever
                logging.info("Using vector-only retriever (no documents for BM25)")

            return self.retriever

        except Exception as e:
            logging.error(f"Failed to create retriever: {e}")
            return None

    def create_rag_tool(self):
        """Create RAG tool for document search with source attribution."""
        if not self.retriever:
            logging.error("Retriever not initialized")
            return None

        from langchain_core.tools import tool

        @tool
        def search_local_documents(query: str) -> str:
            """
            Searches and returns information from local PDF and Word documents with source attribution.
            Use this for questions about policies, reports, or specific documented information.
            Returns relevant content from documents with clear source references.
            """
            try:
                docs = self.retriever.invoke(query)

                if not docs:
                    return f"No relevant documents found for query: '{query}'"

                # Build response with source attribution
                sources_content = {}

                for doc in docs:
                    source_file = os.path.basename(doc.metadata.get('source', 'Unknown document'))
                    if source_file not in sources_content:
                        sources_content[source_file] = []
                    sources_content[source_file].append(doc.page_content.strip())

                response_parts = []
                for source_file, contents in sources_content.items():
                    combined_content = "\n\n".join(contents[:4])
                    response_parts.append(f"**From {source_file}:**\n{combined_content}")

                result = "\n\n" + "\n\n".join(response_parts[:10])

                source_list = list(sources_content.keys())
                if len(source_list) == 1:
                    result += f"\n\n**Source:** {source_list[0]}"
                else:
                    result += f"\n\n**Sources:** {', '.join(source_list[:10])}"
                    if len(source_list) > 10:
                        result += f" and {len(source_list) - 10} other documents"

                return result

            except Exception as e:
                logging.error(f"Error in document search: {e}")
                return f"Error searching documents: {str(e)}"

        return search_local_documents

    def test_document_search(self, query: str = "network access control") -> str:
        """Test the document search functionality directly."""
        if not self.retriever:
            return "Retriever not initialized"

        try:
            docs = self.retriever.invoke(query)
            if docs:
                result = f"Found {len(docs)} relevant documents for query '{query}':\n\n"
                for i, doc in enumerate(docs[:3], 1):
                    content = doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content
                    source = doc.metadata.get('source', 'Unknown')
                    result += f"{i}. Source: {source}\nContent: {content}\n\n"
                return result
            else:
                return f"No documents found for query '{query}'"
        except Exception as e:
            return f"Error testing document search: {e}"

    def rebuild_index(self, embeddings=None) -> bool:
        """Delete and rebuild the entire ChromaDB index from scratch."""
        logging.debug("Rebuilding vector store from scratch...")

        # Delete existing collection
        try:
            self.client.delete_collection(DOCUMENTS_COLLECTION)
            self._collection = None
            logging.debug("Deleted existing ChromaDB collection")
        except Exception as e:
            logging.warning(f"Could not delete collection: {e}")

        # Reset internal state
        self.retriever = None
        self.all_documents = []

        # Rebuild
        success = self.sync_documents(embeddings)
        if success:
            return self.initialize_vector_store(embeddings) and self.create_retriever() is not None
        return False

    def get_document_stats(self) -> dict:
        """Get statistics about loaded documents."""
        try:
            count = self.collection.count()
            return {
                "status": "initialized" if count > 0 else "empty",
                "total_chunks": count,
                "retrieval_k": self.retrieval_k,
                "chunk_size": self.chunk_size,
                "chunk_overlap": self.chunk_overlap,
                "chroma_path": self.chroma_path,
                "documents_directory": self.pdf_directory
            }
        except Exception as e:
            return {"status": "error", "error": str(e)}

    def update_retrieval_settings(self, k: int = None, chunk_size: int = None, chunk_overlap: int = None):
        """Update retrieval settings."""
        if k is not None:
            self.retrieval_k = k
            if self.retriever and hasattr(self.retriever, 'k'):
                self.retriever.k = k

        if chunk_size is not None:
            self.chunk_size = chunk_size

        if chunk_overlap is not None:
            self.chunk_overlap = chunk_overlap

        logging.debug(f"Updated settings: k={self.retrieval_k}, chunk_size={self.chunk_size}")

    # Compatibility properties for old faiss_index_path references
    @property
    def faiss_index_path(self) -> str:
        """Compatibility property - returns chroma_path."""
        return self.chroma_path

    @faiss_index_path.setter
    def faiss_index_path(self, value: str):
        """Compatibility setter - updates chroma_path."""
        self.chroma_path = value
